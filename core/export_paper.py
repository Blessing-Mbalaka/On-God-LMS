
from docx import Document
from core.models import ExamNode

def export_paper_to_docx(paper):
    doc = Document()
    doc.add_heading(f'Assessment Paper: {paper.title or str(paper.id)}', level=1)

    def render_node(node, doc, level=0):
        indent = " " * (level * 4)
        prefix = f"{indent}{node.number or ''} "

        if node.node_type == "question":
            doc.add_paragraph(f"{prefix}{node.text}")
            if node.marks:
                doc.add_paragraph(f"({node.marks} Marks)").italic = True

        elif node.node_type == "case_study":
            doc.add_paragraph(f"{prefix}Case Study: {node.text}")

        elif node.node_type == "table":
            content = node.content or []
            if content:
                table = doc.add_table(rows=1, cols=len(content[0]))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(content[0]):
                    hdr_cells[i].text = str(col)
                for row in content[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row):
                        row_cells[i].text = str(cell)

        for child in node.children.all().order_by("number"):
            render_node(child, doc, level + 1)

    root_nodes = ExamNode.objects.filter(paper=paper, parent__isnull=True).order_by("number")
    for node in root_nodes:
        render_node(node, doc)

    return doc
