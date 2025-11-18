import json
import base64
from docx import Document
from docx.shared import Inches
from io import BytesIO

def insert_block(doc, block):
    if block["type"] == "question":
        doc.add_heading(f"Question {block.get('number', '')} ({block.get('marks', '0')} Marks)", level=2)
        if block.get("text"):
            doc.add_paragraph(block["text"])
        for sub in block.get("content", []):
            insert_block(doc, sub)
        for child in block.get("children", []):
            insert_block(doc, child)

    elif block["type"] == "question_text":
        doc.add_paragraph(block["text"])
    elif block["type"] == "case_study":
        doc.add_paragraph("Case Study:", style='Intense Quote')
        doc.add_paragraph(block["text"])
    elif block["type"] == "table":
        rows = block["rows"]
        if rows:
            table = doc.add_table(rows=0, cols=len(rows[0]))
            for row in rows:
                cells = table.add_row().cells
                for i, cell in enumerate(row):
                    cells[i].text = str(cell)
    elif block["type"] == "figure":
        if "data_uri" in block:
            try:
                image_data = base64.b64decode(block["data_uri"].split(",")[-1])
                image_stream = BytesIO(image_data)
                doc.add_picture(image_stream, width=Inches(4))
            except Exception as e:
                doc.add_paragraph("Error loading image.")

def rebuild_doc(json_path, output_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    for block in data:
        insert_block(doc, block)

    doc.save(output_path)
    print(f"Saved: {output_path}")
