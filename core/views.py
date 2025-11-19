import json
import csv
import logging
import os
import random
import string
import re
import time
import traceback
import uuid
from copy import deepcopy
from decimal import Decimal, InvalidOperation
import google.generativeai as genai
from django.conf import settings
from django.contrib import admin, messages
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth import authenticate, get_user_model, login, logout
from django.contrib.auth.admin import UserAdmin
from django.contrib.auth.decorators import login_required
from django.core.files import File
from django.core.files.base import ContentFile
from urllib.parse import urlencode
from django.core.mail import send_mail
from django.db import models, transaction, IntegrityError
from django.db.models import Sum, Count, Q
from django.http import Http404, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils.timezone import now
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods, require_POST
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import JSONParser, MultiPartParser
from io import BytesIO

from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.dateparse import parse_date
from xhtml2pdf import pisa
from django.core.validators import validate_email
from django.core.exceptions import ValidationError
from docx import Document
from docx.shared import Inches
import os
from .models import OfflineStudent, Qualification, ExtractorBlockImage, ExamAnswer, ExtractorBlock, ExtractorPaper, ExtractorTestItem, ExtractorTestPaper,ExtractorUserBox
from .forms import StudentRegistrationForm

from robustexamextractor import extract_docx, save_robust_extraction_to_db
from utils_pro import populate_examnodes_from_structure_json, save_nodes_to_db
from utils_pro import copy_images_to_media_folder

from .forms import (
    AssessmentCentreForm,
    CustomUserForm,
    EmailRegistrationForm,
    QualificationForm,
)
from .models import (
    Assessment,
    AssessmentCentre,
    Batch,
    CaseStudy,
    CustomUser,
    ExamAnswer,
    ExamNode,
    ExamSubmission,
    Feedback,
    GeneratedQuestion,
    Paper,
    PaperBankEntry,
    Qualification,
    QuestionBankEntry,
    OfflineStudent,
)
from . import extractor_views
from .question_bank import QUESTION_BANK


# from .chieta_extractor.models import (
#     ExtractorBlock,
#     ExtractorBlockImage,
#     ExtractorPaper,
#     ExtractorUserBox,
#)

from utils.question_detect import annotate_paper_questions
from utils.extract_docx import extract_blocks_from_docx


# from .utils import (
#     annotate_paper_questions,
#     extract_blocks_from_docx,
# )
from .paper_utils import (
    build_node_tree,
    build_randomized_structure_from_pool,
    get_pool_summary,
    collect_randomization_pool,
    RandomizationPoolError,
    build_randomized_from_pool_only,
    calculate_pool_gaps,
)
from . import qualification_registry
from .randomization_config import (
    allowed_letters,
    cover_title,
    get_module_meta,
    randomization_status,
)


@require_POST
def set_theme(request):
    theme = request.POST.get("theme") or "light"
    request.session["site_theme"] = theme
    next_url = request.POST.get("next") or request.META.get("HTTP_REFERER") or "/"
    return redirect(next_url)


# core/views.py


def redirect_user_by_role(user):
    role = user.role
    if not role and user.is_staff:
        return redirect("admin_dashboard")
    if role == "default":
        return redirect("default")
    if role == "admin":
        return redirect("admin_dashboard")
    if role == "administrator":
        return redirect("admin_dashboard")
    elif role == "moderator":
        return redirect("moderator_developer")
    elif role == "internal_mod":
        return redirect("internal_moderator_dashboard")
    elif role == "assessor_marker":
        return redirect("assessor_maker_dashboard")
    elif role == "external_mod":
        return redirect("external_moderator_dashboard")
    elif role == "assessor_dev":
        return redirect("assessor_developer")
    elif role == "qcto":
        return redirect("qcto_dashboard")
    elif role == "etqa":
        return redirect("etqa_dashboard")
    elif role == "learner":
        return redirect("student_dashboard")
    elif role == "assessment_center":
        return redirect("assessment_center")
    else:
        return redirect("admin_dashboard") if user.is_staff else redirect("waiting_activation")


def register(request):
    if request.method == "POST":
        form = EmailRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.is_active = False
            user.role = user.role or "default"
            user.is_staff = False
            user.save()
            messages.success(
                request,
                "Your account request has been submitted. You'll be notified once it's activated.",
            )
            return redirect("waiting_activation")
    else:
        form = EmailRegistrationForm()

    return render(request, "core/login/login.html", {"form": form})


def custom_login(request):
    # we reuse the registration form so the template can render both sides
    list(messages.get_messages(request))
    reg_form = EmailRegistrationForm()

    if request.method == "POST":
        email = request.POST.get("email", "").strip().lower()
        password = request.POST.get("password", "")

        try:
            account = CustomUser.objects.get(email=email)
        except CustomUser.DoesNotExist:
            account = None
        else:
            if not account.is_active or account.role == "default":
                messages.info(
                    request,
                    "Your account is pending activation. Please wait for an administrator to approve it.",
                )
                return redirect("waiting_activation")

        user = authenticate(request, username=email, password=password)
        if user:
            login(request, user)
            return redirect_user_by_role(user)
        else:
            error_message = (
                "Invalid credentials"
                if account
                else "No account found with that email. Please sign up first."
            )
            return render(
                request,
                "core/login/login.html",
                {
                    "form": reg_form,
                    "error": error_message,
                },
            )

    return render(request, "core/login/login.html", {"form": reg_form})


# _______________________________________________________________________________________________________
# ******************************************************************************************************
# LOGIN LOGIC AND USER ACCESS CONTROL STARTS HERE********************************************************
# *******************************************************************************************************
# *******************************************************************************************************


CustomUser = get_user_model()


@login_required
# @staff_member_required
def user_management(request):
    if request.method == "POST":
        return handle_user_creation(request)

    users = CustomUser.objects.select_related("qualification").exclude(
        is_superuser=True
    )
    quals = Qualification.objects.all()
    assessment_centres = AssessmentCentre.objects.select_related(
        "qualification_assigned"
    ).all()
    # print("printing")
    # print(assessment_centers)
    return render(
        request,
        "core/administrator/user_management.html",
        {
            "users": users,
            "qualifications": quals,
            "assessment_centres": assessment_centres,
            "role_choices": CustomUser.ROLE_CHOICES,
        },
    )


def handle_user_creation(request):
    try:
        # Validate required fields
        name = request.POST.get("name", "").strip()
        email = request.POST.get("email", "").strip().lower()
        role = request.POST.get("role", "").strip()

        if not all([name, email, role]):
            messages.error(request, "Please fill in all required fields.")
            return redirect("user_management")

        # Special handling for assessment center role
        if role == "assessment_center":
            center_id = request.POST.get("assessment_centre", "").strip()
            if not center_id:
                messages.error(
                    request,
                    "Please select an assessment center for assessment center users.",
                )
                return redirect("user_management")

            # Get the assessment center and its qualification
            center = get_object_or_404(AssessmentCentre, pk=center_id)
            qualification = center.qualification_assigned
            if not qualification:
                messages.error(
                    request, "Selected assessment center has no qualification assigned."
                )
                return redirect("user_management")
        else:
            # For non-assessment center roles, require qualification
            qual_id = request.POST.get("qualification", "").strip()
            if not qual_id:
                messages.error(request, "Please select a qualification.")
                return redirect("user_management")
            qualification = get_object_or_404(Qualification, pk=qual_id)

        # Validate email format
        try:
            validate_email(email)
        except ValidationError:
            messages.error(request, "Please enter a valid email address.")
            return redirect("user_management")

        # Check for existing user
        if CustomUser.objects.filter(email=email).exists():
            messages.error(request, "A user with this email already exists.")
            return redirect("user_management")

        # Split name
        parts = name.split(maxsplit=1)
        first_name = parts[0]
        last_name = parts[1] if len(parts) > 1 else ""

        # Generate random password
        pwd = "".join(random.choices(string.ascii_letters + string.digits, k=12))

        # Create user
        user = CustomUser.objects.create_user(
            username=email,
            email=email,
            first_name=first_name,
            last_name=last_name,
            role=role,
            qualification=qualification,
            is_active=True,
            is_staff=(role != "learner"),
            activated_at=now(),
        )
        user.set_password(pwd)

        # Assign assessment center if applicable
        if role == "assessment_center":
            user.assessment_centre = center
            
        user.save()

        # Send email using new email module
        email_sent = send_account_creation_email(email, first_name, pwd)
        
        if email_sent:
            messages.success(
                request, f"User {email} created successfully. Password emailed."
            )
        else:
            messages.warning(
                request, 
                f"User created but email failed to send. Please check logs or contact the user directly."
            )

        return redirect("user_management")

        #user.save()

        # Send email
       # try:
          #  send_mail(
           #     "Your CHIETA LMS Account",
           #     f"Hello {first_name},\n\nYour account has been created.\n\n"
            #    f"Username: {email}\n"
             #   f"Password: {pwd}\n\n"
             #   "Please change your password after logging in.",
              #  "noreply@chieta.co.za",
             #   [email],
              #  fail_silently=False,
           # )
           ## print(
               ## f"[AccountEmail] Credentials sent to {email} with temporary password.",
               # flush=True,
          #  )
          #  messages.success(
                request, f"User {email} created successfully. Password emailed."
            )
      #  except Exception as e:
         #   messages.warning(
              #  request, f"User created but email failed to send: {str(e)}"
           # )
           # print(
             #   f"[AccountEmail] Failed to send credentials to {email}: {e}",
               # flush=True,
           # )

        #return redirect("user_management")

   # except Exception as e:
        #messages.error(request, f"An error occurred: {str(e)}")
       # return redirect("user_management")






########################################################################################################



# *******************************************************************************************************
# *******************************************************************************************************
# Role Management is done here
# ______________________________________________________________________________________________________


@require_POST
@login_required
def update_user_role(request, user_id):
    try:
        user = get_object_or_404(CustomUser, pk=user_id)
        user.role = request.POST.get("role")
        user.save()

        return JsonResponse(
            {
                "success": True,
                "message": f"Role updated for {user.get_full_name()}.",
                "new_role": user.role,
                "new_role_display": user.get_role_display(),
            }
        )

    except Exception as e:
        return JsonResponse({"success": False, "message": str(e)}, status=400)


# *******************************************************************************************************
# *******************************************************************************************************
# Role Management is done here
# ______________________________________________________________________________________________________


# @login_required
# # @staff_member_required
# def update_user_role(request, user_id):
#     user = get_object_or_404(CustomUser, pk=user_id)
#     if request.method == "POST":
#         user.role = request.POST["role"]
#         user.save()
#         messages.success(request, f"Role updated for {user.get_full_name()}.")
#     return redirect("user_management")


# _______________________________________________________________________________________________________
# _______________________________________________________________________________________________________
# User qualification
@login_required
# @staff_member_required
def update_user_qualification(request, user_id):
    user = get_object_or_404(CustomUser, pk=user_id)
    if request.method == "POST":
        qual = get_object_or_404(Qualification, pk=request.POST["qualification"])
        user.qualification = qual
        user.save()
        messages.success(request, f"Qualification updated for {user.get_full_name()}.")
    return redirect("user_management")


# User Status
@login_required
# @staff_member_required
def toggle_user_status(request, user_id):
    user = get_object_or_404(CustomUser, pk=user_id)
    user.is_active = not user.is_active
    if user.is_active:
        user.activated_at = now()
        user.deactivated_at = None
    else:
        user.deactivated_at = now()
    user.save()
    state = "activated" if user.is_active else "deactivated"
    messages.success(request, f"{user.get_full_name()} {state}.")
    return redirect("user_management")


# ********************************************************************************************************
# ********************************************************************************************************
# ASSESSMENT CENTRE VIEWS FOR ADDING ETC________________________________________________________________
def assessment_centres_view(request):
    centres = AssessmentCentre.objects.all()
    form = AssessmentCentreForm()

    # Handle form submission
    if request.method == "POST":
        form = AssessmentCentreForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Assessment centre added successfully.")
            return redirect("assessment_centres")
        else:
            messages.error(request, "Please correct the errors below.")

    return render(
        request,
        "core/administrator/centre.html",
        {
            "centres": centres,
            "form": form,
        },
    )


def edit_assessment_centre(request, centre_id):
    centre = get_object_or_404(AssessmentCentre, id=centre_id)

    if request.method == "POST":
        form = AssessmentCentreForm(request.POST, instance=centre)
        if form.is_valid():
            form.save()
            messages.success(request, "Assessment centre updated successfully!")
            return redirect("assessment_centres")
    else:
        form = AssessmentCentreForm(instance=centre)

    return render(
        request,
        "core/edit_centre.html",
        {
            "form": form,
            "centre": centre,
        },
    )


def delete_assessment_centre(request, centre_id):
    centre = get_object_or_404(AssessmentCentre, id=centre_id)
    centre.delete()
    messages.success(request, "Assessment centre removed successfully!")
    return redirect("assessment_centres")


# Helper Function
# --- robust helpers (safe defaults) ---------------------------------
MARK_RE = re.compile(r"(\d+)\s*(?:mark|marks)\b", re.I)


def extract_marks_from_text(text: str) -> int:
    if not text:
        return 0
    m = MARK_RE.findall(text)
    if m:
        # take the last number in the line like "(10 Marks)"
        try:
            return int(m[-1])
        except ValueError:
            return 0
    # bare numbers at the end e.g. "... [5]"
    tail = re.findall(r"(\d+)\s*$", text)
    return int(tail[-1]) if tail else 0


def extract_node_text_from_robust(node: dict) -> str:
    # prefer explicit text field
    if isinstance(node.get("text"), str):
        return node["text"]
    # try content array
    parts = []
    for item in node.get("content", []):
        if isinstance(item, dict):
            if "text" in item and isinstance(item["text"], str):
                parts.append(item["text"])
            # table text flatten (header + cells)
            if item.get("type") == "table":
                t = item.get("table", {})
                for row in t.get("rows", []):
                    for cell in row.get("cells", []):
                        ct = cell.get("text", "")
                        if ct:
                            parts.append(ct)
    return " ".join(p.strip() for p in parts if p)


def extract_marks_from_robust_data(node: dict) -> int:
    # explicit marks first
    val = node.get("marks")
    if isinstance(val, int):
        return val
    if isinstance(val, str):
        nums = re.findall(r"\d+", val)
        return int(nums[0]) if nums else 0
    # fall back to text/content
    return extract_marks_from_text(extract_node_text_from_robust(node))


def handle_robust_image_content(node: dict, paper_obj) -> list:
    """
    Placeholder: return lightweight descriptors only.
    If you later persist files, do it here and return stored paths.
    """
    meta = node.get("image") or {}
    # Avoid dumping heavy/binary data into JSONField
    cleaned = {
        k: v
        for k, v in meta.items()
        if isinstance(v, (str, int, float, bool)) or v is None
    }
    return [cleaned] if cleaned else []


# --------------------------------------------------------------------


def save_robust_manifest_to_db(nodes, paper_obj):
    """Save robust extraction nodes to database with better content handling"""
    try:
        # Track mapping from node numbers to DB records
        node_map = {}
        order_index = 0

        # First pass: create all nodes
        for node_data in nodes:
            # Extract basic node data
            node_type = node_data.get("type", "unknown")
            node_number = node_data.get("number", "")
            node_text = node_data.get("text", "")
            node_marks = node_data.get("marks", "")

            # Process content properly based on type
            content_items = []

            # Handle content items
            for item in node_data.get("content", []):
                if isinstance(item, dict):
                    # Handle table content
                    if item.get("type") == "table" and "rows" in item:
                        content_items.append({"type": "table", "rows": item["rows"]})
                    # Handle image content
                    elif item.get("type") == "figure":
                        image_data = {}
                        if "images" in item:
                            image_data["images"] = item["images"]
                        if "data_uri" in item:
                            image_data["data_uri"] = item["data_uri"]
                        content_items.append({"type": "figure", **image_data})
                    # Other content
                    else:
                        content_items.append(item)

            # Create the node
            node = ExamNode.objects.create(
                paper=paper_obj,
                node_type=node_type,
                number=node_number,
                text=node_text,
                marks=node_marks,
                content=content_items,  # Store as proper JSON
                order_index=order_index,
            )

            # Track in map for parent relationships
            if node_number:
                node_map[node_number] = node

            order_index += 1

        # Second pass: establish parent-child relationships
        for node_data in nodes:
            node_number = node_data.get("number", "")
            if not node_number or node_number not in node_map:
                continue

            # Find parent if exists
            parts = node_number.split(".")
            if len(parts) > 1:
                parent_number = ".".join(parts[:-1])
                if parent_number in node_map:
                    node = node_map[node_number]
                    node.parent = node_map[parent_number]
                    node.save()

        return True

    except Exception as e:
        print(traceback.format_exc())
        return False


def find_marks_in_following_tables(nodes, current_index):
    """Algorithm to find marks in following table nodes - handles your specific case"""
    print(f"[capture] Looking for marks in tables following node {current_index}")

    # Look at the next few nodes for tables
    for i in range(current_index + 1, min(current_index + 5, len(nodes))):
        next_node = nodes[i]

        # Stop if we hit another question (marks belong to previous question)
        if next_node.get("type") == "question":
            break

        # Check if this is a table with marks
        if next_node.get("type") == "table":
            table_marks = extract_marks_from_table_content(next_node)
            if table_marks > 0:
                print(
                    f"[capture] Found {table_marks} marks in following table at index {i}"
                )
                return table_marks

    print("[capture] No marks found in following tables")
    return 0


def extract_marks_from_table_content(node_data):
    """Extract marks from table content in node_data"""
    try:
        # Look through content array for tables
        for content_item in node_data.get("content", []):
            if content_item.get("type") == "table":
                marks = extract_marks_from_table(content_item)
                if marks > 0:
                    return marks

        # Fallback: look for marks in the raw text
        text = extract_node_text_from_robust(node_data)
        return extract_marks_from_text(text)

    except Exception as e:
        return 0


# Enhanced marks extraction with better table parsing
def extract_marks_from_table(table_item):
    """Extract marks from table structure - enhanced for your specific format."""
    try:
        table_data = table_item.get("table", {})
        rows = table_data.get("rows", [])

        # Strategy 1: look for a "Total" row that contains the overall marks
        for row in rows:
            cells = row.get("cells", [])
            if not cells:
                continue

            first_cell_text = (cells[0].get("text") or "").strip().lower()
            if "total" in first_cell_text:
                for cell in cells:
                    cell_text = (cell.get("text") or "").strip()
                    if not cell_text:
                        continue

                    if cell_text.isdigit():
                        marks = int(cell_text)
                        print(f"[capture] Found total-row marks: {marks}")
                        return marks

                    numbers = re.findall(r"\d+", cell_text)
                    if numbers:
                        marks = int(numbers[-1])
                        print(
                            f"[capture] Extracted total-row marks from '{cell_text}': {marks}"
                        )
                        return marks

        # Strategy 2: identify a marks column and sum the entries
        marks_column_index = None
        for row in rows:
            cells = row.get("cells", [])
            for col_idx, cell in enumerate(cells):
                cell_text = (cell.get("text") or "").strip().lower()
                if "marks" in cell_text or "mark" in cell_text:
                    marks_column_index = col_idx
                    break
            if marks_column_index is not None:
                break

        if marks_column_index is not None:
            total_marks = 0
            for row in rows[1:]:  # skip header row
                cells = row.get("cells", [])
                if len(cells) <= marks_column_index:
                    continue

                marks_cell = (cells[marks_column_index].get("text") or "").strip()
                first_cell = (
                    (cells[0].get("text") or "").strip().lower() if cells else ""
                )
                if "total" in first_cell:
                    continue

                if marks_cell.isdigit():
                    total_marks += int(marks_cell)
                    continue

                numbers = re.findall(r"\d+", marks_cell)
                if numbers:
                    total_marks += int(numbers[-1])

            if total_marks > 0:
                print(f"[capture] Calculated total marks from column: {total_marks}")
                return total_marks

        # Strategy 3: fall back to scanning all text within the table
        combined_text = " ".join(
            (cell.get("text") or "") for row in rows for cell in row.get("cells", [])
        )
        text_marks = extract_marks_from_text(combined_text)
        if text_marks > 0:
            print(f"[capture] Found marks in combined text: {text_marks}")
            return text_marks

    except Exception as exc:  # pragma: no cover - defensive logging
        print(f"[capture] Unable to extract marks from table: {exc}")

    return 0


# Enhanced admin dashboard calculation
def calculate_total_marks_from_manifest(nodes):
    """Enhanced marks calculation that looks at following tables"""
    total_marks = 0

    for i, node in enumerate(nodes):
        if node.get("type") == "question":
            # Try direct extraction first
            extracted_marks = extract_marks_from_robust_data(node)

            # If no marks found, look in following tables
            if extracted_marks == 0:
                extracted_marks = find_marks_in_following_tables(nodes, i)

            total_marks += extracted_marks

            if extracted_marks > 0:
                question_num = node.get("number", "Unknown")

    return total_marks


# core/views.py (near your other helpers)
def build_questions_tree_for_paper(paper):
    nodes = list(
        ExamNode.objects.filter(paper=paper)
        .select_related("parent")
        .order_by("order_index")
    )

    questions = []
    node_map = {}

    # pass 1: shape dicts
    for node in nodes:
        node_dict = {
            "id": str(node.id),
            "type": node.node_type,
            "number": node.number or "",
            "text": node.text or "",
            "marks": node.marks or "",
            "content": node.content or [],  # keep JSON as-is
            "children": [],
        }
        node_map[str(node.id)] = node_dict
        if node.node_type == "question":
            questions.append(node_dict)

    # pass 2: wire parents/children
    for node in nodes:
        if node.parent and str(node.parent.id) in node_map:
            parent_dict = node_map[str(node.parent.id)]
            child_dict = node_map[str(node.id)]
            parent_dict["children"].append(child_dict)

    return questions


# ===================================================================================


@login_required
def admin_dashboard(request):
    """Administrator dashboard view - handles file uploads and paper selection"""

    if request.method == "POST" and request.FILES.get("file_input"):
        file_obj = request.FILES["file_input"]
        memo_file_obj = request.FILES.get("memo_input")  # Get memo file if exists
        
        paper_number = request.POST.get("paper_number", "Unnamed Paper")
        qual_pk = (request.POST.get("qualification") or "").strip()

        # Generate EISA ID
        eisa_id = f"EISA-{uuid.uuid4().hex[:8].upper()}"

        if not qual_pk:
            messages.error(request, "Please select a qualification.")
            return redirect("admin_dashboard")

        try:
            qualification_obj = Qualification.objects.get(pk=int(qual_pk))
        except (ValueError, Qualification.DoesNotExist):
            messages.error(request, "Invalid qualification selected.")
            return redirect("admin_dashboard")

        # Create temp directory
        temp_dir = os.path.join(settings.MEDIA_ROOT, "temp_uploads")
        os.makedirs(temp_dir, exist_ok=True)

        # Create a unique filename for this upload
        unique_filename = f"upload_{uuid.uuid4().hex}.docx"
        temp_path = os.path.join(temp_dir, unique_filename)

        # DEBUG: Print file information
        print(f"[capture] Processing question paper: {file_obj.name}")
        if memo_file_obj:
            print(f"[capture] Processing memo: {memo_file_obj.name}")

        # Write question paper file to disk
        with open(temp_path, "wb+") as destination:
            for chunk in file_obj.chunks():
                destination.write(chunk)

        # After writing file to disk
        file_size = os.path.getsize(temp_path)
        if file_size == 0:
            print("[capture] ERROR: File was saved with 0 bytes!")
        else:
            print(f"[capture] Question paper saved with {file_size} bytes")

        # Process the document
        print("[capture] Starting robust extraction...")
        manifest = extract_docx(
            temp_path, out_dir=None, use_gemini=False, use_gemma=False
        )

        node_count = len(manifest.get("nodes", []))

        if not manifest or "nodes" not in manifest or not manifest["nodes"]:
            messages.error(request, "Extraction failed: No content found in document")
            return redirect("admin_dashboard")

        # Create paper object
        paper_obj = Paper.objects.create(
            name=paper_number,
            qualification=qualification_obj,
            created_by=request.user,
            is_randomized=False,
        )
        paper_obj.structure_json = manifest

        # Determine status from button clicked
        action = request.POST.get("action")  # "moderator" or "etqa"
        if action == "etqa":
            status = "Submitted to ETQA"
        elif action == "moderator":
            status = "Submitted to Moderator"
        else:
            status = "Pending"

        # Create Assessment record with BOTH files - use 'memo' field for admin uploads
        assessment = Assessment.objects.create(
            eisa_id=eisa_id,
            qualification=qualification_obj,
            paper=paper_number,
            file=file_obj,  # question paper file
            memo=memo_file_obj,  # CORRECTED: Use 'memo' field for memo file
            created_by=request.user,
            paper_link=paper_obj,
            status=status,
            paper_type='admin_upload',  # Explicitly set paper type
        )

        # Save extracted content to DB
        print("[capture] Saving nodes to database...")
        conversion_success = save_robust_manifest_to_db(manifest["nodes"], paper_obj)
        if not conversion_success:
            messages.error(request, "Failed to save extracted content to database")
            paper_obj.delete()
            return redirect("admin_dashboard")

        # Count extracted elements
        questions = sum(1 for n in manifest["nodes"] if n.get("type") == "question")
        tables = sum(
            1
            for n in manifest["nodes"]
            for c in n.get("content", [])
            if isinstance(c, dict) and c.get("type") == "table"
        )
        images = sum(
            1
            for n in manifest["nodes"]
            for c in n.get("content", [])
            if isinstance(c, dict) and c.get("type") == "figure"
        )
        total_marks = sum(
            int(node.get("marks", 0) or 0)
            for node in manifest["nodes"]
            if node.get("type") == "question" and node.get("marks")
        )

        # Update paper with marks
        paper_obj.total_marks = total_marks
        paper_obj.save()

        # Create success message
        if memo_file_obj:
            summary_message = (
                "Extraction complete. Both question paper and memo saved. "
                f"Questions: {questions}, Tables: {tables}, Images: {images}, "
                f"Total marks: {total_marks or 0}. "
                f"EISA ID: {eisa_id}. Status: {status}."
            )
        else:
            summary_message = (
                "Extraction complete. Question paper saved (no memo uploaded). "
                f"Questions: {questions}, Tables: {tables}, Images: {images}, "
                f"Total marks: {total_marks or 0}. "
                f"EISA ID: {eisa_id}. Status: {status}."
            )
            
        messages.success(request, summary_message)

        # Clean up temp file
        try:
            os.unlink(temp_path)
        except:
            pass

        return redirect("load_saved_paper", paper_pk=paper_obj.id)

    # Dashboard data
    tools = Assessment.objects.select_related("qualification", "created_by").order_by(
        "-created_at"
    )
    total_users = CustomUser.objects.filter(is_superuser=False).count()
    quals = Qualification.objects.all()

    return render(
        request,
        "core/administrator/admin_dashboard.html",
        {
            "tools": tools,
            "total_users": total_users,
            "qualifications": quals,
        },
    )

# _____________________________________________________________________________________________________


# _____________________________________________________________________________________________________
def qualification_management_view(request):
    qualifications = Qualification.objects.all()
    form = QualificationForm()
    registry_entries = qualification_registry.get_entries()
    registry_json = json.dumps(registry_entries)

    if request.method == "POST":
        form = QualificationForm(request.POST)
        registry_entries = qualification_registry.get_entries()
        registry_json = json.dumps(registry_entries)
        if form.is_valid():
            qualification = form.save()
            messages.success(request, "Qualification added successfully.")
            return redirect("manage_qualifications")
        else:
            messages.error(request, "Please correct the errors below.")

    return render(
        request,
        "core/administrator/qualifications.html",
        {
            "qualifications": qualifications,
            "form": form,
            "qualification_registry_entries": registry_entries,
            "qualification_registry_json": registry_json,
        },
    )


# 0) Databank View 2025/06/10 made to handle logic for the databank for the question generation.


def databank_view(request):
    # Get questions from ExamNode
    entries = (
        ExamNode.objects.filter(node_type="question")
        .select_related("paper__qualification")
        .order_by("-created_at")
    )

    return render(
        request,
        "core/administrator/databank.html",
        {
            "entries": entries,
            "qualifications": Qualification.objects.all(),
        },
    )


# -------------------------------------------
# 1) Add a new question to the Question Bank
# ---------------------------------------


@csrf_exempt
def add_question(request):
    if request.method == "POST":
        q_type = request.POST.get("question_type")
        qualification_id = request.POST.get("qualification")
        marks = request.POST.get("marks")
        text = request.POST.get("text")

        # Fetch qualification object
        try:
            qualification = get_object_or_404(Qualification, pk=qualification_id)
        except ValueError:
            messages.error(request, "Invalid qualification selected.")
            return redirect("databank")

        # Prepare case study if needed
        case_study_id = request.POST.get("case_study")
        case_study = None
        if q_type == "case_study" and case_study_id:
            try:
                case_study = CaseStudy.objects.get(id=case_study_id)
            except CaseStudy.DoesNotExist:
                messages.error(request, "Selected case study not found.")
                return redirect("assessor_developer")

        # Basic validation
        if not text or not marks:
            messages.error(request, "Please fill in all required fields.")
            return redirect("assessor_developer")

        # Create the question
        question = QuestionBankEntry.objects.create(
            qualification=qualification,
            question_type=q_type,
            text=text,
            marks=int(marks),
            case_study=case_study,
        )

        # Handle MCQ options
        if q_type == "mcq":
            has_correct = False
            for i in range(1, 5):
                opt_text = request.POST.get(f"opt_text_{i}")
                is_correct = request.POST.get(f"opt_correct_{i}") == "on"
                if opt_text:
                    MCQOption.objects.create(
                        question=question, text=opt_text, is_correct=is_correct
                    )
                    if is_correct:
                        has_correct = True
            if not has_correct:
                question.delete()
                messages.error(
                    request, "At least one MCQ option must be marked as correct."
                )
                return redirect("assessor_developer")

        messages.success(request, "Question added to the databank.")
        return redirect("databank")


@csrf_exempt
def add_case_study(request):
    if request.method == "POST":
        title = request.POST.get("cs_title")
        content = request.POST.get("cs_content")
        if title and content:
            CaseStudy.objects.create(title=title, content=content)
            messages.success(request, "Case study added successfully.")
    return redirect("databank")


############################################################################################
# Assessor developer pipeline (advanced extractor integration)
###########################################################################################
############################################################################################
###########################################################################################
def sync_assessment_paper_bank(assessment, force=False):
    """Ensure the latest uploaded file is kept in the assessor paper bank."""
    if not assessment.file:
        return None
    try:
        file_path = assessment.file.path
    except (ValueError, OSError):
        return None

    entry, _ = PaperBankEntry.objects.get_or_create(assessment=assessment)
    target_name = os.path.basename(assessment.file.name)
    needs_copy = (
        force
        or not entry.original_file
        or os.path.basename(entry.original_file.name) != target_name
    )

    if needs_copy:
        with open(file_path, "rb") as source:
            entry.original_file.save(target_name, File(source), save=True)
    return entry


def rebuild_extractor_from_bank(assessment, force=False):
    """Create or refresh the advanced extractor paper for an assessment."""
    entry = sync_assessment_paper_bank(assessment, force=force)
    if not entry or not entry.original_file:
        return None

    extractor = assessment.extractor_paper
    if extractor is None:
        extractor = ExtractorPaper.objects.create(
            title=assessment.paper or os.path.basename(entry.original_file.name),
            module_name=getattr(assessment, "module_name", "") or "",
            paper_number=assessment.paper,
            paper_letter=getattr(assessment, "module_number", "") or "",
        )
        assessment.extractor_paper = extractor
        assessment.save(update_fields=["extractor_paper"])
    else:
        extractor.title = assessment.paper or extractor.title
        if hasattr(assessment, "module_name"):
            extractor.module_name = assessment.module_name or extractor.module_name
        if hasattr(assessment, "module_number"):
            extractor.paper_letter = assessment.module_number or extractor.paper_letter

    if force or not extractor.original_file:
        with entry.original_file.open("rb") as source:
            extractor.original_file.save(
                os.path.basename(entry.original_file.name), File(source), save=False
            )
    extractor.save()

    try:
        docx_path = extractor.original_file.path
    except (ValueError, OSError):
        return extractor

    try:
        with transaction.atomic():
            extractor.user_boxes.all().delete()
            extractor.blocks.all().delete()
            payload = extract_blocks_from_docx(docx_path, paper=extractor)
            for idx, block in enumerate(payload):
                block_obj = ExtractorBlock.objects.create(
                    paper=extractor,
                    order_index=idx,
                    block_type=block.get("type", "paragraph"),
                    xml=block.get("xml", ""),
                    text=block.get("text", ""),
                )
                for image in block.get("images", []) or []:
                    ExtractorBlockImage.objects.create(block=block_obj, image=image)
        annotate_paper_questions(extractor)
    except Exception as exc:
        traceback.print_exc()
    return extractor


@login_required
def assessor_reports(request):
    data = [
        {
            "qualification": "Maintenance Planner",
            "toolsGenerated": 10,
            "toolsSubmitted": 8,
            "questionsAdded": 5,
        },
        {
            "qualification": "Quality Controller",
            "toolsGenerated": 15,
            "toolsSubmitted": 12,
            "questionsAdded": 9,
        },
    ]

    return render(
        request,
        "core/assessor-developer/assessor_reports.html",
        {
            "report_data": json.dumps(data),
            "report_list": data,
        },
    )


def assessment_archive(request):
    qs = Assessment.objects.all()

    qual = request.GET.get("qualification", "").strip()
    paper = request.GET.get("paper", "").strip()
    status = request.GET.get("status", "").strip()

    if qual:
        qs = qs.filter(qualification=qual)
    if paper:
        qs = qs.filter(paper__icontains=paper)
    if status:
        qs = qs.filter(status=status)

    all_quals = (
        Assessment.objects.order_by("qualification")
        .values_list("qualification", flat=True)
        .distinct()
    )
    all_statuses = [choice[0] for choice in Assessment.STATUS_CHOICES]

    if request.GET.get("export") == "csv":
        response = HttpResponse(content_type="text/csv")
        response["Content-Disposition"] = (
            'attachment; filename="assessment_archive.csv"'
        )
        writer = csv.writer(response)
        writer.writerow(["EISA ID", "Qualification", "Paper", "Status", "Date"])

        for assessment in qs:
            writer.writerow(
                [
                    assessment.eisa_id,
                    assessment.qualification,
                    assessment.paper,
                    assessment.status,
                    assessment.created_at.date().isoformat(),
                ]
            )

        return response

    return render(
        request,
        "core/assessor-developer/assessment_archive.html",
        {
            "assessments": qs,
            "filter_qualification": qual,
            "filter_paper": paper,
            "filter_status": status,
            "all_qualifications": all_quals,
            "all_statuses": all_statuses,
        },
    )


def assessor_dashboard(request):
    user = request.user
    qualification = getattr(user, "qualification", None)
    if qualification:
        assessments = Assessment.objects.filter(qualification=qualification).order_by(
            "-created_at"
        )
    else:
        assessments = Assessment.objects.none()
    return render(
        request,
        "core/assessor-developer/assessor_dashboard.html",
        {
            "assessments": assessments,
            "qualification": qualification,
            "user": user,
        },
    )


@login_required
@require_http_methods(["GET", "POST"])
def view_assessment(request, eisa_id):
    assessment = get_object_or_404(
        Assessment.objects.select_related(
            "qualification", "created_by"
        ).prefetch_related("generated_questions"),
        eisa_id=eisa_id,
    )

    if not request.user.is_staff:
        allowed_roles = {"admin", "moderator", "assessor_dev"}
        user_role = getattr(request.user, "role", "")
        if (
            assessment.created_by_id != request.user.id
            and user_role not in allowed_roles
        ):
            raise Http404("Assessment not available")

    if request.method == "POST":
        notes = (request.POST.get("moderator_notes") or "").strip()
        assessment.moderator_notes = notes
        assessment.status = "Submitted to Moderator"
        assessment.save(update_fields=["moderator_notes", "status"])
        messages.success(request, f"{assessment.eisa_id} forwarded to moderator.")
        return redirect("assessor_dashboard")

    questions = assessment.generated_questions.all()

    randomized_filter = (
        Q(paper_link__structure_json__randomization__base_assessment_id=assessment.id)
        | Q(
            paper_link__structure_json__randomization__previous_snapshot_assessment_id=assessment.id
        )
        | Q(paper_link__structure_json__randomization__source_snapshot_id=assessment.id)
    )
    if assessment.paper_link_id:
        randomized_filter |= Q(
            paper_link__structure_json__randomization__base_paper_id=assessment.paper_link_id
        )

    randomized_assessments = list(
        Assessment.objects.filter(
            paper_type="randomized",
            paper_link__isnull=False,
        )
        .filter(randomized_filter)
        .select_related("paper_link")
        .order_by("-created_at", "-id")
    )

    return render(
        request,
        "core/assessor-developer/view_assessment.html",
        {
            "assessment": assessment,
            "questions": questions,
            "randomized_assessments": randomized_assessments,
        },
    )


def assessor_developer(request):
    """Dashboard that manages assessor developer paper flow using the advanced extractor."""
    base_queryset = Assessment.objects.select_related(
        "qualification", "extractor_paper", "paper_bank_entry", "created_by"
    ).order_by("-created_at")

    user = request.user
    user_role = getattr(user, "role", "")
    user_qualification = getattr(user, "qualification", None)

    if user.is_superuser or user_role in {"admin", "moderator"}:
        assessments_qs = base_queryset
    elif user_qualification:
        assessments_qs = base_queryset.filter(qualification=user_qualification)
    else:
        assessments_qs = base_queryset.filter(created_by=user)

    if request.method == "POST":
        action = request.POST.get("action")
        assessment_id = request.POST.get("assessment_id")
        force_refresh = request.POST.get("refresh") == "1"
        if action and assessment_id:
            assessment = get_object_or_404(assessments_qs, pk=assessment_id)
            if action == "forward_etqa":
                assessment.status = "Submitted to ETQA"
                assessment.save(update_fields=["status"])
                messages.success(request, f"{assessment.eisa_id} forwarded to ETQA.")
            elif action == "forward_qcto":
                assessment.status = "Submitted to QCTO"
                assessment.save(update_fields=["status"])
                messages.success(request, f"{assessment.eisa_id} forwarded to QCTO.")
            elif action == "open_pipeline":
                extractor_paper = rebuild_extractor_from_bank(
                    assessment, force=force_refresh
                )
                if extractor_paper:
                    return redirect(
                        "assessor_developer_paper", paper_id=extractor_paper.id
                    )
                messages.error(
                    request, "Unable to prepare paper for advanced pipeline."
                )
        return redirect("assessor_developer")

    assessments_list = list(assessments_qs)
    awaiting_etqa_qs = assessments_qs.filter(status="Submitted to ETQA")
    awaiting_qcto_qs = assessments_qs.filter(status="Submitted to QCTO")
    randomized_snapshots_qs = (
        assessments_qs.filter(
            paper_type="randomized",
            paper_link__isnull=False,
        )
        .select_related("paper_link", "qualification")
        .order_by("-created_at", "-id")
    )

    context = {
        "assessments": assessments_list,
        "awaiting_etqa": list(awaiting_etqa_qs),
        "awaiting_qcto": list(awaiting_qcto_qs),
        "randomized_snapshots": list(randomized_snapshots_qs),
    }
    return render(request, "core/assessor-developer/assessor_developer.html", context)


@login_required
def assessor_pool_data(request):
    module_filter = (request.GET.get("module") or "").strip()
    letter_filter_raw = (request.GET.get("letter") or "").strip()
    letter_filter = "".join(ch for ch in letter_filter_raw if ch.isalpha()).lower()

    user = request.user
    user_role = getattr(user, "role", "")
    user_qualification = getattr(user, "qualification", None)
    restrict_to_qualification = None
    if user_qualification and not (
        user.is_superuser or user.is_staff or user_role in {"admin", "moderator"}
    ):
        restrict_to_qualification = user_qualification

    def normalize_letter(raw: str | None) -> str:
        raw = (raw or "").strip()
        letters = "".join(ch for ch in raw if ch.isalpha())
        return (letters or raw or "A").upper()

    def summarize_content(raw: str | None) -> str:
        if not raw:
            return ""
        try:
            data = json.loads(raw)
        except Exception:
            return (raw or "")[:160]
        if isinstance(data, dict):
            items = data.get("items") or []
        elif isinstance(data, list):
            items = data
        else:
            return str(data)[:160]
        snippets: list[str] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            item_type = (item.get("type") or "").lower()
            if item_type in {
                "text",
                "question_text",
                "instruction",
                "case_study",
                "heading",
            }:
                text_val = (item.get("text") or "").strip()
                if text_val:
                    snippets.append(text_val)
            elif item_type == "table":
                snippets.append("[table]")
            elif item_type == "image":
                images = item.get("images") or []
                if images:
                    snippets.append(f"[image x{len(images)}]")
            if len(" ".join(snippets)) > 160:
                break
        return " ".join(snippets)[:200]

    boxes_qs = (
        ExtractorUserBox.objects.select_related(
            "paper", "paper__assessment_record__qualification"
        )
        .filter(paper__module_name__isnull=False)
        .exclude(paper__module_name="")
    )

    if restrict_to_qualification:
        boxes_qs = boxes_qs.filter(
            paper__assessment_record__qualification=restrict_to_qualification
        )
    if module_filter:
        boxes_qs = boxes_qs.filter(paper__module_name__iexact=module_filter)

    modules_map: dict[str, dict[str, list[dict[str, object]]]] = {}

    def natural_key(value: str | None):
        text = value or ""
        return [
            int(part) if part.isdigit() else part.lower()
            for part in re.split(r"(\d+)", text)
        ]

    for box in boxes_qs:
        paper_ref = box.paper
        module_name = (paper_ref.module_name or "").strip() if paper_ref else ""
        if not module_name:
            module_name = "Unknown Module"
        letter_norm = normalize_letter(getattr(paper_ref, "paper_letter", ""))
        if letter_filter and letter_norm.lower() != letter_filter:
            continue

        entry_list = modules_map.setdefault(module_name, {}).setdefault(letter_norm, [])
        entry_list.append(
            {
                "id": box.id,
                "question_number": (box.question_number or "").strip(),
                "parent_number": (box.parent_number or "").strip(),
                "qtype": box.qtype or "",
                "marks": box.marks or "",
                "content_preview": summarize_content(box.content),
                "content": box.content or "",
                "paper_id": paper_ref.id if paper_ref else None,
                "paper_title": paper_ref.title if paper_ref else "",
                "created_at": box.created_at.isoformat() if box.created_at else None,
            }
        )

    modules_payload: list[dict[str, object]] = []
    for module_name, letters_map in sorted(
        modules_map.items(), key=lambda item: item[0].lower()
    ):
        letters_payload: list[dict[str, object]] = []
        for letter, entries in sorted(letters_map.items(), key=lambda item: item[0]):
            entries.sort(key=lambda entry: natural_key(entry.get("question_number")))
            letters_payload.append(
                {
                    "letter": letter,
                    "count": len(entries),
                    "entries": entries,
                }
            )
        modules_payload.append(
            {
                "name": module_name,
                "letters": letters_payload,
                "total": sum(letter["count"] for letter in letters_payload),
            }
        )

    return JsonResponse(
        {
            "ok": True,
            "modules": modules_payload,
        }
    )


@login_required
@require_POST
def assessor_pool_randomize(request):
    module_name = (request.POST.get("module") or "").strip()
    letter_raw = (request.POST.get("letter") or "").strip()
    module_letter = (
        "".join(ch for ch in letter_raw if ch.isalpha()) or letter_raw
    ).upper()
    mode = (request.POST.get("mode") or "pool").strip().lower()
    if mode not in {"pool", "base"}:
        mode = "pool"

    if not module_name or not module_letter:
        return JsonResponse(
            {"ok": False, "message": "Module and letter are required."}, status=400
        )

    user = request.user
    user_role = getattr(user, "role", "")
    user_qualification = getattr(user, "qualification", None)

    if (
        not (user.is_staff or user.is_superuser or user_role in {"admin", "moderator"})
        and not user_qualification
    ):
        return JsonResponse(
            {
                "ok": False,
                "message": "You do not have a qualification set in your profile.",
            },
            status=403,
        )

    assessment_base_qs = Assessment.objects.select_related(
        "paper_link", "extractor_paper", "qualification"
    ).filter(
        module_name__iexact=module_name,
        paper_link__isnull=False,
        paper_link__is_randomized=False,
    )

    if user_qualification and not (
        user.is_staff or user.is_superuser or user_role in {"admin", "moderator"}
    ):
        assessment_base_qs = assessment_base_qs.filter(qualification=user_qualification)

    base_assessment = (
        assessment_base_qs.filter(module_number__iexact=module_letter)
        .order_by("-paper_link__updated_at", "-paper_link__id")
        .first()
    ) or (
        assessment_base_qs.order_by(
            "-paper_link__updated_at", "-paper_link__id"
        ).first()
    )

    base_paper = (
        base_assessment.paper_link
        if (base_assessment and base_assessment.paper_link_id)
        else None
    )
    base_extractor = (
        base_assessment.extractor_paper
        if (base_assessment and base_assessment.extractor_paper_id)
        else None
    )
    try:
        pool_info = collect_randomization_pool(
            module_name, module_letter, base_extractor
        )
    except RandomizationPoolError as exc:
        return JsonResponse({"ok": False, "message": str(exc)}, status=400)

    # If using base blueprint, ensure pool coverage for its question numbers
    if mode == "base":
        if base_paper is None:
            return JsonResponse(
                {
                    "ok": False,
                    "message": "No base paper found for this module and letter. Switch to pool-only mode or capture the original paper first.",
                },
                status=400,
            )
        pool_by_number = pool_info["pool_by_number"]
        blueprint_nodes = ExamNode.objects.filter(
            paper=base_paper, node_type="question"
        ).order_by("order_index")
        missing_numbers: list[str] = []
        seen_numbers: set[str] = set()
        for node in blueprint_nodes:
            number = (node.number or "").strip()
            if not number or number in seen_numbers:
                continue
            seen_numbers.add(number)
            if (number, "question") not in pool_by_number:
                missing_numbers.append(number)

        if missing_numbers:
            return JsonResponse(
                {
                    "ok": False,
                    "message": "Not enough captured blocks to randomize this module. Capture the missing question numbers and try again.",
                    "missing": missing_numbers,
                },
                status=400,
            )
    elif mode == "pool":
        missing_numbers = calculate_pool_gaps(pool_info)
        if missing_numbers:
            return JsonResponse(
                {
                    "ok": False,
                    "message": "Snapshot pool has gaps. Capture the listed question numbers and try again.",
                    "missing": missing_numbers,
                },
                status=400,
            )

    qualification = (
        user_qualification
        or (base_assessment.qualification if base_assessment else None)
        or (base_paper.qualification if base_paper else None)
    )
    if qualification is None:
        return JsonResponse(
            {
                "ok": False,
                "message": "You do not have a qualification set in your profile.",
            },
            status=400,
        )

    created_by = request.user if request.user.is_authenticated else None
    timestamp_label = now().strftime("%Y%m%d-%H%M%S")
    new_paper_name = f"{module_name} {module_letter} Randomized {timestamp_label}"

    new_paper = Paper.objects.create(
        name=new_paper_name,
        qualification=qualification,
        created_by=created_by,
        is_randomized=True,
    )

    try:
        if mode == "base":
            result = build_randomized_structure_from_pool(
                base_paper,
                new_paper,
                module_name,
                module_letter,
                base_extractor=base_extractor,
            )
        else:
            result = build_randomized_from_pool_only(
                new_paper,
                module_name,
                module_letter,
                base_extractor=base_extractor,
            )
    except RandomizationPoolError as exc:
        new_paper.delete()
        return JsonResponse({"ok": False, "message": str(exc)}, status=400)

    total_marks = result.get("total_marks", 0)
    pool_size = result.get("pool_size", 0)
    pool_used = result.get("used_pool", False)
    selected_boxes = (
        result.get("selected_boxes", []) if isinstance(result, dict) else []
    )

    qual_name = qualification.name if qualification else None
    rand_status_value = (
        randomization_status(qual_name, module_letter) if qual_name else "unconfigured"
    )
    allowed_letters_value = (
        allowed_letters(qual_name, module_letter) if qual_name else []
    )
    cover_value = cover_title(qual_name, module_letter) if qual_name else None

    meta_payload = dict(new_paper.structure_json or {})
    meta_payload["randomization"] = {
        "module_name": module_name,
        "module_number": module_letter,
        "status": rand_status_value,
        "allowed_letters": allowed_letters_value,
        "base_paper_id": base_paper.id if base_paper else None,
        "source": "pool_modal",
        "mode": mode,
        "snapshot_pool_used": pool_used,
        "snapshot_pool_size": pool_size,
        "last_refreshed": now().isoformat(),
    }
    if selected_boxes:
        meta_payload["randomization"]["selected_boxes"] = selected_boxes
    if qual_name:
        meta_payload["randomization"]["qualification"] = qual_name
    if cover_value:
        meta_payload["randomization"]["cover_title"] = cover_value
    if base_assessment and base_assessment.id:
        meta_payload["randomization"]["base_assessment_id"] = base_assessment.id
        if base_assessment.extractor_paper_id:
            meta_payload["randomization"][
                "base_extractor_id"
            ] = base_assessment.extractor_paper_id

    new_paper.total_marks = total_marks
    new_paper.structure_json = meta_payload
    new_paper.save(update_fields=["total_marks", "structure_json"])

    eisa_id = f"EISA-{uuid.uuid4().hex[:8].upper()}"
    comment_parts = [
        "Randomized from pool view (preview)",
        f"module: {module_name}",
        f"letter: {module_letter}",
        f"pool size: {pool_size}",
    ]

    new_assessment = Assessment.objects.create(
        eisa_id=eisa_id,
        qualification=qualification,
        paper=new_paper.name,
        module_number=module_letter,
        module_name=module_name,
        saqa_id=(
            qualification.saqa_id
            if getattr(qualification, "saqa_id", None)
            else (
                base_assessment.saqa_id
                if base_assessment and base_assessment.saqa_id
                else ""
            )
        ),
        comment=" | ".join(comment_parts),
        created_by=created_by,
        paper_link=new_paper,
        status="draft",
        paper_type="randomized",
        forward_to_moderator=False,
    )

    source_for_files = base_assessment
    if source_for_files and source_for_files.file:
        with source_for_files.file.open("rb") as src_file:
            new_assessment.file.save(
                os.path.basename(source_for_files.file.name),
                File(src_file),
                save=False,
            )
    if source_for_files and source_for_files.memo:
        with source_for_files.memo.open("rb") as src_memo:
            new_assessment.memo.save(
                os.path.basename(source_for_files.memo.name),
                File(src_memo),
                save=False,
            )
    new_assessment.save()

    if new_assessment.file:
        sync_assessment_paper_bank(new_assessment, force=True)
        rebuild_extractor_from_bank(new_assessment, force=True)

    snapshot_params = urlencode(
        {
            "created": "1",
            "eisa": new_assessment.eisa_id,
            "paper": new_paper.id,
            "pool": pool_size,
            "marks": total_marks,
            "download": "1",
        }
    )
    snapshot_url = f"{reverse('assessor_randomized_snapshot', args=[new_assessment.id])}?{snapshot_params}"
    return JsonResponse(
        {
            "ok": True,
            "message": "Randomized paper ready.",
            "snapshot_url": snapshot_url,
            "download_url": reverse(
                "download_randomized_pdf", args=[new_assessment.id]
            ),
        }
    )


@login_required
def download_randomized_pdf(request, assessment_id):
    assessment = get_object_or_404(
        Assessment.objects.select_related("paper_link", "qualification"),
        id=assessment_id,
        paper_type="randomized",
    )
    paper = assessment.paper_link
    if not paper:
        raise Http404("Randomized paper not found.")

    user = request.user
    user_role = getattr(user, "role", "")
    user_qualification = getattr(user, "qualification", None)
    if not (
        user.is_staff
        or user.is_superuser
        or user_role in {"admin", "moderator", "etqa", "qcto", "assessment_center"}
    ):
        allowed = assessment.created_by_id == user.id
        if (
            not allowed
            and user_qualification
            and assessment.qualification_id == getattr(user_qualification, "id", None)
        ):
            allowed = True
        if not allowed:
            raise Http404("Assessment not available")

    node_tree, node_stats = build_node_tree(paper)
    context = {
        "assessment": assessment,
        "paper": paper,
        "node_tree": node_tree,
        "node_stats": node_stats,
        "generated_at": timezone.now(),
        "requested_by": request.user,
    }

    req_format = (request.GET.get("format") or "pdf").lower()

    if req_format == "docx":
        # Build a Word document from the node_tree
        doc = Document()
        doc.add_heading(f"{assessment.eisa_id} - {paper.name}", level=1)
        if getattr(assessment, "qualification", None):
            doc.add_paragraph(f"Qualification: {assessment.qualification.name}")
        doc.add_paragraph(f"Generated: {timezone.now().isoformat()}")

        from django.conf import settings as djsettings

        def _render_node_to_doc(node, level=2):
            ntype = (node.get("node_type") or "").lower()
            number = node.get("number") or ""
            if ntype == "cover_page":
                doc.add_heading("Cover Page", level=level)
            elif ntype == "instruction":
                doc.add_heading("Instruction", level=level)
            elif ntype == "question":
                doc.add_heading(f"Question {number}", level=level)

            for item in node.get("content") or []:
                try:
                    if isinstance(item, dict):
                        itype = (item.get("type") or "").lower()

                        # images/figures
                        if itype in {"image", "figure"}:
                            # Support a variety of image representations: data_uri, url, images list
                            import base64
                            import tempfile
                            import requests

                            def _save_and_add(img_path_or_bytes, is_bytes=False):
                                # img_path_or_bytes: filesystem path or bytes
                                try:
                                    if is_bytes:
                                        tf = tempfile.NamedTemporaryFile(
                                            delete=False, suffix=".png"
                                        )
                                        tf.write(img_path_or_bytes)
                                        tf.flush()
                                        tf.close()
                                        doc.add_picture(tf.name, width=Inches(5))
                                        os.unlink(tf.name)
                                    else:
                                        doc.add_picture(
                                            img_path_or_bytes, width=Inches(5)
                                        )
                                    return True
                                except Exception:
                                    return False

                            embedded = False
                            # 1) direct data_uri on item
                            data_uri = item.get("data_uri") or item.get("data")
                            if isinstance(data_uri, str) and data_uri.startswith(
                                "data:"
                            ):
                                try:
                                    header, b64 = data_uri.split(",", 1)
                                    bin_data = base64.b64decode(b64)
                                    embedded = _save_and_add(bin_data, is_bytes=True)
                                except Exception:
                                    embedded = False

                            # 2) direct url/file/src reference
                            if not embedded:
                                url = (
                                    item.get("url")
                                    or item.get("file")
                                    or item.get("src")
                                    or ""
                                )
                                if url:
                                    img_path = url
                                    # map MEDIA_URL to MEDIA_ROOT
                                    if url.startswith(djsettings.MEDIA_URL):
                                        img_path = os.path.join(
                                            djsettings.MEDIA_ROOT,
                                            url.replace(
                                                djsettings.MEDIA_URL, ""
                                            ).lstrip("/"),
                                        )
                                    # if looks like absolute http(s), try download
                                    if str(img_path).lower().startswith("http"):
                                        try:
                                            r = requests.get(img_path, timeout=8)
                                            if r.status_code == 200:
                                                embedded = _save_and_add(
                                                    r.content, is_bytes=True
                                                )
                                        except Exception:
                                            embedded = False
                                    else:
                                        embedded = _save_and_add(
                                            img_path, is_bytes=False
                                        )

                            # 3) images list
                            if not embedded:
                                images = item.get("images") or item.get("image") or []
                                if isinstance(images, (str, bytes)):
                                    images = [images]
                                for im in images:
                                    if embedded:
                                        break
                                    try:
                                        if isinstance(im, dict):
                                            # dict with data_uri or url or path
                                            diu = im.get("data_uri") or im.get("data")
                                            if isinstance(diu, str) and diu.startswith(
                                                "data:"
                                            ):
                                                header, b64 = diu.split(",", 1)
                                                bin_data = base64.b64decode(b64)
                                                embedded = _save_and_add(
                                                    bin_data, is_bytes=True
                                                )
                                                continue
                                            iurl = (
                                                im.get("url")
                                                or im.get("path")
                                                or im.get("filename")
                                                or im.get("name")
                                                or ""
                                            )
                                            if iurl:
                                                ipath = iurl
                                                if iurl.startswith(
                                                    djsettings.MEDIA_URL
                                                ):
                                                    ipath = os.path.join(
                                                        djsettings.MEDIA_ROOT,
                                                        iurl.replace(
                                                            djsettings.MEDIA_URL, ""
                                                        ).lstrip("/"),
                                                    )
                                                if (
                                                    str(ipath)
                                                    .lower()
                                                    .startswith("http")
                                                ):
                                                    try:
                                                        r = requests.get(
                                                            ipath, timeout=8
                                                        )
                                                        if r.status_code == 200:
                                                            embedded = _save_and_add(
                                                                r.content, is_bytes=True
                                                            )
                                                    except Exception:
                                                        embedded = False
                                                else:
                                                    embedded = _save_and_add(
                                                        ipath, is_bytes=False
                                                    )
                                        else:
                                            # string filename or url
                                            sval = (
                                                im.decode()
                                                if isinstance(im, bytes)
                                                else str(im)
                                            )
                                            spath = sval
                                            if sval.startswith(djsettings.MEDIA_URL):
                                                spath = os.path.join(
                                                    djsettings.MEDIA_ROOT,
                                                    sval.replace(
                                                        djsettings.MEDIA_URL, ""
                                                    ).lstrip("/"),
                                                )
                                            if str(spath).lower().startswith("http"):
                                                try:
                                                    r = requests.get(spath, timeout=8)
                                                    if r.status_code == 200:
                                                        embedded = _save_and_add(
                                                            r.content, is_bytes=True
                                                        )
                                                except Exception:
                                                    embedded = False
                                            else:
                                                embedded = _save_and_add(
                                                    spath, is_bytes=False
                                                )
                                    except Exception:
                                        continue

                            if not embedded:
                                # fallback: insert a reference to the image
                                ref = (
                                    item.get("url")
                                    or item.get("file")
                                    or item.get("src")
                                    or ""
                                )
                                doc.add_paragraph(f"[Image: {ref or 'unavailable'}]")

                        # text-like items
                        elif itype in {
                            "text",
                            "instruction",
                            "question_text",
                            "heading",
                        }:
                            txt = item.get("text") or ""
                            if txt:
                                doc.add_paragraph(txt)

                        # tables (support HTML table in item.get('html'))
                        elif itype == "table":
                            html_blob = (
                                item.get("html")
                                or item.get("data")
                                or item.get("text")
                                or ""
                            )
                            if (
                                isinstance(html_blob, str)
                                and "<table" in html_blob.lower()
                            ):
                                # simple HTML table -> native docx table
                                import re
                                from html import unescape

                                try:
                                    rows = re.findall(
                                        r"<tr[^>]*>(.*?)</tr>",
                                        html_blob,
                                        flags=re.I | re.S,
                                    )
                                    parsed = [
                                        re.findall(
                                            r"<t[dh][^>]*>(.*?)</t[dh]>",
                                            r,
                                            flags=re.I | re.S,
                                        )
                                        for r in rows
                                    ]
                                    if parsed:
                                        maxcols = max(len(r) for r in parsed)
                                        table = doc.add_table(
                                            rows=len(parsed), cols=maxcols
                                        )
                                        table.style = "Table Grid"
                                        for ri, r in enumerate(parsed):
                                            for ci in range(maxcols):
                                                cell_text = ""
                                                if ci < len(r):
                                                    txt = re.sub(
                                                        r"<[^>]+>", "", r[ci] or ""
                                                    )
                                                    cell_text = unescape(txt).strip()
                                                table.rows[ri].cells[
                                                    ci
                                                ].text = cell_text
                                        # small spacer after table
                                        doc.add_paragraph("")
                                        continue
                                except Exception:
                                    pass

                            # fallback: dump raw data/text
                            doc.add_paragraph(
                                str(item.get("data") or item.get("text") or "[table]")
                            )

                        else:
                            doc.add_paragraph(str(item))
                    else:
                        doc.add_paragraph(str(item))
                except Exception:
                    continue

            for child in node.get("children") or []:
                _render_node_to_doc(child, level=level + 1)

        for node in node_tree:
            _render_node_to_doc(node, level=2)

        out = BytesIO()
        doc.save(out)
        out.seek(0)
        filename = f"{assessment.eisa_id}_{paper.name.replace(' ', '_')}.docx"
        response = HttpResponse(
            out.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    html = render_to_string("core/assessor-developer/randomized_pdf.html", context)
    pdf_buffer = BytesIO()

    # link_callback helps xhtml2pdf resolve /static/ and /media/ URLs to filesystem paths
    def link_callback(uri, rel):
        # uri: the URL from the HTML
        # rel: relative path (ignored)
        from django.conf import settings
        import os

        if uri.startswith(settings.MEDIA_URL):
            path = uri.replace(settings.MEDIA_URL, "")
            return os.path.join(settings.MEDIA_ROOT, path)
        if uri.startswith(settings.STATIC_URL):
            path = uri.replace(settings.STATIC_URL, "")
            return os.path.join(settings.STATIC_ROOT, path)
        # if it's already an absolute filesystem path, return it
        if os.path.exists(uri):
            return uri
        return uri

    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer, link_callback=link_callback)
    if pisa_status.err:
        messages.error(request, "Unable to generate a printable version right now.")
        referer = request.META.get("HTTP_REFERER") or reverse(
            "assessor_randomized_snapshot", args=[assessment.id]
        )
        return redirect(referer)

    pdf_buffer.seek(0)
    filename = f"{assessment.eisa_id}_{paper.name.replace(' ', '_')}.pdf"
    response = HttpResponse(pdf_buffer.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@login_required
def assessor_randomized_snapshot(request, assessment_id):
    randomized_qs = Assessment.objects.select_related(
        "paper_link", "qualification", "created_by", "extractor_paper"
    )
    assessment = get_object_or_404(
        randomized_qs,
        id=assessment_id,
        paper_type="randomized",
    )

    user = request.user
    user_role = getattr(user, "role", "")
    user_qualification = getattr(user, "qualification", None)

    if not (
        user.is_staff
        or user.is_superuser
        or user_role in {"admin", "moderator", "etqa", "qcto", "assessment_center"}
    ):
        allowed = assessment.created_by_id == user.id
        if (
            not allowed
            and user_qualification
            and assessment.qualification_id == getattr(user_qualification, "id", None)
        ):
            allowed = True
        if not allowed:
            raise Http404("Assessment not available")
    paper = assessment.paper_link
    random_meta: dict = {}
    if paper and isinstance(paper.structure_json, dict):
        random_meta = paper.structure_json.get("randomization", {}) or {}

    base_paper = None
    base_assessment = None
    base_paper_id = None
    base_assessment_id = None
    if isinstance(random_meta, dict):
        base_paper_id = random_meta.get("base_paper_id")
        base_assessment_id = random_meta.get("base_assessment_id")
        if base_paper_id:
            base_paper = Paper.objects.filter(id=base_paper_id).first()
        if base_assessment_id:
            base_assessment = (
                Assessment.objects.select_related(
                    "extractor_paper", "qualification", "paper_link"
                )
                .filter(id=base_assessment_id)
                .first()
            )

    if base_paper is None and paper and not paper.is_randomized:
        base_paper = paper
    if base_paper is None and isinstance(random_meta, dict):
        previous_snapshot_id = random_meta.get("previous_snapshot_paper_id")
        if previous_snapshot_id:
            base_paper = Paper.objects.filter(id=previous_snapshot_id).first()
    if base_paper is None and paper:
        base_paper = getattr(paper, "base_paper", None)

    module_name_meta = (
        (random_meta.get("module_name") if isinstance(random_meta, dict) else None)
        or getattr(assessment, "module_name", "")
        or getattr(paper or None, "module_name", "")
        or ""
    )
    module_number_meta = (
        (random_meta.get("module_number") if isinstance(random_meta, dict) else None)
        or getattr(assessment, "module_number", "")
        or getattr(paper or None, "paper_letter", "")
        or ""
    )

    qual_name = None
    if assessment.qualification:
        qual_name = assessment.qualification.name
    elif isinstance(random_meta, dict):
        qual_name = random_meta.get("qualification")
    elif base_paper and base_paper.qualification:
        qual_name = base_paper.qualification.name

    allowed_letter_options = (
        allowed_letters(qual_name, module_number_meta) if qual_name else []
    )
    random_status_label = (
        randomization_status(qual_name, module_number_meta)
        if qual_name
        else "unconfigured"
    )
    cover_heading = None
    allowed_letters_display = ""
    letters_value = None
    if isinstance(random_meta, dict):
        letters_value = random_meta.get("allowed_letters")
    if letters_value:
        if isinstance(letters_value, (list, tuple, set)):
            allowed_letters_display = ", ".join(str(v) for v in letters_value if v)
        else:
            allowed_letters_display = str(letters_value)
    elif allowed_letter_options:
        allowed_letters_display = ", ".join(str(v) for v in allowed_letter_options if v)
    if isinstance(random_meta, dict):
        cover_heading = random_meta.get("cover_title")
    if not cover_heading and qual_name:
        cover_heading = cover_title(qual_name, module_number_meta)

    base_extractor = None
    if base_assessment and base_assessment.extractor_paper_id:
        base_extractor = base_assessment.extractor_paper
    elif assessment.extractor_paper_id:
        base_extractor = assessment.extractor_paper
    elif isinstance(random_meta, dict) and random_meta.get("base_extractor_id"):
        base_extractor = ExtractorPaper.objects.filter(
            id=random_meta["base_extractor_id"]
        ).first()

    # NEW: Check memo requirement for randomized papers
    can_forward_to_moderator = True
    memo_requirement_message = ""
    has_memo = bool(assessment.memo_file)
    
    if assessment.paper_type == 'randomized':
        if not assessment.memo_file:
            can_forward_to_moderator = False
            memo_requirement_message = "A memo must be uploaded before forwarding to moderator"
        else:
            can_forward_to_moderator = True

    if request.method == "POST":
        action = (request.POST.get("action") or "").strip().lower()
        refresh = request.POST.get("refresh") == "1"
        
        # NEW: Handle memo upload
        if action == "upload_memo":
            memo_file = request.FILES.get('memo_file')
            if memo_file:
                # Validate file type
                allowed_types = ['.pdf', '.doc', '.docx']
                file_ext = os.path.splitext(memo_file.name)[1].lower()
                
                if file_ext not in allowed_types:
                    messages.error(
                        request, 
                        f"Invalid file type. Allowed types: {', '.join(allowed_types)}"
                    )
                else:
                    assessment.memo_file = memo_file
                    assessment.save(update_fields=['memo_file'])
                    messages.success(request, "Memo uploaded successfully")
                    # Update the local variables
                    has_memo = True
                    can_forward_to_moderator = True
            else:
                messages.error(request, "Please select a memo file to upload")
            return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
        
        if action in {"forward_etqa", "forward_qcto"}:
            status_lookup = {
                "forward_etqa": "Submitted to ETQA",
                "forward_qcto": "Submitted to QCTO",
            }
            assessment.status = status_lookup[action]
            assessment.status_changed_at = now()
            if request.user.is_authenticated:
                assessment.status_changed_by = request.user
            assessment.save(
                update_fields=["status", "status_changed_at", "status_changed_by"]
            )
            messages.success(
                request, f"{assessment.eisa_id} status updated to {assessment.status}."
            )
            return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
        
        # UPDATED: Forward to moderator with memo check
        if action == "forward_moderator":
            # Check memo requirement for randomized papers
            if assessment.paper_type == 'randomized' and not assessment.memo_file:
                messages.error(
                    request, 
                    "Cannot forward to moderator: A memo must be uploaded for randomized papers"
                )
                return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
            
            if assessment.status == "Submitted to Moderator":
                messages.info(
                    request, f"{assessment.eisa_id} is already with the moderator."
                )
            else:
                assessment.status = "Submitted to Moderator"
                assessment.forward_to_moderator = True
                assessment.status_changed_at = now()
                if request.user.is_authenticated:
                    assessment.status_changed_by = request.user
                assessment.save(
                    update_fields=[
                        "status",
                        "forward_to_moderator",
                        "status_changed_at",
                        "status_changed_by",
                    ]
                )
                messages.success(
                    request, f"{assessment.eisa_id} forwarded to the moderator queue."
                )
            return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
        
        if action == "open_pipeline":
            extractor = rebuild_extractor_from_bank(assessment, force=refresh)
            if extractor:
                return redirect("assessor_developer_paper", paper_id=extractor.id)
            messages.error(
                request, "Unable to prepare extractor view for this snapshot."
            )
            return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
        
        if action == "release_students":
            assessment.status = "Released to students"
            assessment.save(update_fields=["status"])
            messages.success(request, f"{assessment.eisa_id} released to learners.")
            return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
        
        if action == "delete_node":
            node_id = request.POST.get("node_id")
            if paper and node_id:
                node = paper.nodes.filter(id=node_id).first()
                if node:
                    node.delete()
                    messages.success(request, f"Removed block {node_id}.")
                else:
                    messages.info(request, "Block not found; nothing removed.")
            else:
                messages.info(request, "No block selected to remove.")
            return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
        
        if action == "convert_node":
            node_id = request.POST.get("node_id")
            target_type = (request.POST.get("target_type") or "").strip().lower()
            allowed_types = {"question", "instruction", "text", "image", "table"}
            if paper and node_id and target_type in allowed_types:
                node = paper.nodes.filter(id=node_id).first()
                if node:
                    update_fields = ["node_type"]
                    node.node_type = target_type
                    if target_type == "instruction":
                        if node.parent_id is not None:
                            node.parent = None
                            update_fields.append("parent")
                        if node.number:
                            node.number = None
                            update_fields.append("number")
                        if node.marks:
                            node.marks = None
                            update_fields.append("marks")
                    node.save(update_fields=list(dict.fromkeys(update_fields)))
                    messages.success(request, f"Block converted to {target_type}.")
                else:
                    messages.info(request, "Block not found; nothing converted.")
            else:
                messages.info(
                    request, "Unable to convert block with the provided data."
                )
            return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
        
        if action == "update_snapshot":
            if not paper:
                messages.error(request, "Snapshot record missing; cannot refresh.")
                return redirect(
                    "assessor_randomized_snapshot", assessment_id=assessment.id
                )

            base_for_refresh = base_paper or paper
            try:
                result = build_randomized_structure_from_pool(
                    base_for_refresh,
                    paper,
                    module_name_meta,
                    module_number_meta,
                    base_extractor=base_extractor,
                )
            except RandomizationPoolError as exc:
                messages.error(request, str(exc))
                return redirect(
                    "assessor_randomized_snapshot", assessment_id=assessment.id
                )

            meta_payload = (
                paper.structure_json if isinstance(paper.structure_json, dict) else {}
            )
            meta_payload.setdefault("randomization", {})
            meta_payload["randomization"].update(
                {
                    "module_name": module_name_meta,
                    "module_number": module_number_meta,
                    "status": random_status_label,
                    "allowed_letters": allowed_letter_options,
                    "snapshot_pool_used": result.get("used_pool", False),
                    "snapshot_pool_size": result.get("pool_size", 0),
                    "source": "snapshot_pool" if result.get("used_pool") else "clone",
                    "last_refreshed": now().isoformat(),
                }
            )
            if base_paper_id:
                meta_payload["randomization"]["base_paper_id"] = base_paper_id
            if base_assessment_id:
                meta_payload["randomization"]["base_assessment_id"] = base_assessment_id
            if base_extractor and hasattr(base_extractor, "id"):
                meta_payload["randomization"]["base_extractor_id"] = base_extractor.id
            if qual_name:
                meta_payload["randomization"]["qualification"] = qual_name
            if cover_heading:
                meta_payload["randomization"]["cover_title"] = cover_heading

            if paper:
                paper.structure_json = meta_payload
                paper.total_marks = result.get("total_marks", paper.total_marks or 0)
                paper.save(update_fields=["structure_json", "total_marks"])

            random_meta = meta_payload.get("randomization", {})
            if result.get("used_pool"):
                messages.success(
                    request,
                    f"Snapshot refreshed from {result.get('pool_size', 0)} captured blocks.",
                )
            else:
                messages.info(
                    request,
                    "Snapshot rebuilt using original structure (no pool matches found).",
                )
            return redirect("assessor_randomized_snapshot", assessment_id=assessment.id)
        
        if action == "create_snapshot":
            base_source = base_paper or paper
            if base_source is None:
                messages.error(request, "No base paper available to create a snapshot.")
                return redirect(
                    "assessor_randomized_snapshot", assessment_id=assessment.id
                )

            base_random_meta = {}
            if isinstance(getattr(base_source, "structure_json", None), dict):
                base_random_meta = (
                    base_source.structure_json.get("randomization", {}) or {}
                )

            module_name_value = (
                module_name_meta
                or base_random_meta.get("module_name")
                or getattr(base_source, "module_name", "")
                or ""
            )
            module_number_value = (
                module_number_meta
                or base_random_meta.get("module_number")
                or getattr(base_source, "paper_letter", "")
                or ""
            )
            created_by = (
                request.user
                if request.user.is_authenticated
                else (assessment.created_by or getattr(base_source, "created_by", None))
            )
            timestamp_label = now().strftime("%Y%m%d-%H%M%S")
            new_paper_name = (
                f"{base_source.name} Snapshot {timestamp_label}"
                if getattr(base_source, "name", None)
                else f"Snapshot {timestamp_label}"
            )

            if not module_name_value:
                messages.error(
                    request, "Module name is required before creating a snapshot."
                )
                return redirect(
                    "assessor_randomized_snapshot", assessment_id=assessment.id
                )
            if not module_number_value:
                messages.error(
                    request, "Module letter is required before creating a snapshot."
                )
                return redirect(
                    "assessor_randomized_snapshot", assessment_id=assessment.id
                )

            new_paper = Paper.objects.create(
                name=new_paper_name,
                qualification=base_source.qualification,
                created_by=created_by,
                is_randomized=True,
            )
            try:
                result = build_randomized_structure_from_pool(
                    base_source,
                    new_paper,
                    module_name_value,
                    module_number_value,
                    base_extractor=base_extractor,
                )
            except RandomizationPoolError as exc:
                new_paper.delete()
                messages.error(request, str(exc))
                return redirect(
                    "assessor_randomized_snapshot", assessment_id=assessment.id
                )
            total_marks = result.get("total_marks", 0)
            used_pool = result.get("used_pool", False)
            pool_size = result.get("pool_size", 0)

            rand_status_value = (
                randomization_status(qual_name, module_number_value)
                if qual_name
                else "unconfigured"
            )
            allowed_letters_value = (
                allowed_letters(qual_name, module_number_value) if qual_name else []
            )
            cover_value = (
                cover_title(qual_name, module_number_value)
                if qual_name
                else cover_heading
            )

            new_meta = new_paper.structure_json or {}
            new_meta["randomization"] = {
                "module_name": module_name_value,
                "module_number": module_number_value,
                "status": rand_status_value,
                "allowed_letters": allowed_letters_value,
                "base_paper_id": (
                    base_source.id if getattr(base_source, "id", None) else None
                ),
                "source_snapshot_id": assessment.id,
                "source_snapshot_eisa_id": assessment.eisa_id,
                "snapshot_pool_used": used_pool,
                "snapshot_pool_size": pool_size,
                "source": "snapshot_pool" if used_pool else "clone",
                "last_refreshed": now().isoformat(),
            }
            if qual_name:
                new_meta["randomization"]["qualification"] = qual_name
            if cover_value:
                new_meta["randomization"]["cover_title"] = cover_value
            if base_assessment_id:
                new_meta["randomization"]["base_assessment_id"] = base_assessment_id
            if base_extractor and hasattr(base_extractor, "id"):
                new_meta["randomization"]["base_extractor_id"] = base_extractor.id
            if paper and getattr(paper, "id", None):
                new_meta["randomization"]["previous_snapshot_paper_id"] = paper.id
            if assessment.paper_link_id:
                new_meta["randomization"][
                    "previous_snapshot_assessment_id"
                ] = assessment.id

            new_paper.total_marks = total_marks
            new_paper.structure_json = new_meta
            new_paper.save(update_fields=["total_marks", "structure_json"])

            eisa_id = f"EISA-{uuid.uuid4().hex[:8].upper()}"
            qualification_for_snapshot = (
                assessment.qualification or base_source.qualification
            )
            comment_parts = [f"Randomized from {base_source.name}"]
            if used_pool:
                comment_parts.append("source: snapshot pool")
            else:
                comment_parts.append("source: direct clone")
            comment_parts.append(f"pool size: {pool_size}")
            random_comment = " | ".join(comment_parts)

            new_assessment = Assessment.objects.create(
                eisa_id=eisa_id,
                qualification=qualification_for_snapshot,
                paper=new_paper.name,
                module_number=module_number_value,
                module_name=module_name_value,
                saqa_id=getattr(assessment, "saqa_id", ""),
                comment=random_comment,
                created_by=created_by,
                paper_link=new_paper,
                status="draft",
                paper_type="randomized",
                forward_to_moderator=False,
            )

            source_for_files = base_assessment or assessment
            if source_for_files and source_for_files.file:
                with source_for_files.file.open("rb") as src_file:
                    new_assessment.file.save(
                        os.path.basename(source_for_files.file.name),
                        File(src_file),
                        save=False,
                    )
            if source_for_files and source_for_files.memo:
                with source_for_files.memo.open("rb") as src_memo:
                    new_assessment.memo.save(
                        os.path.basename(source_for_files.memo.name),
                        File(src_memo),
                        save=False,
                    )
            new_assessment.save()

            extractor = None
            if new_assessment.file:
                sync_assessment_paper_bank(new_assessment, force=True)
                extractor = rebuild_extractor_from_bank(new_assessment, force=True)

            print(
                f"\nSNAPSHOT SAVED: {eisa_id} | base_paper={getattr(base_source, 'id', 'n/a')} "
                f"-> new_paper={new_paper.id} | used_pool={used_pool} | pool_size={pool_size} | marks={total_marks}"
            )

            messages.success(
                request,
                f"Snapshot {eisa_id} saved with {total_marks} marks (pool size {pool_size}).",
            )
            if extractor:
                return redirect("assessor_developer_paper", paper_id=extractor.id)
            query_string = urlencode(
                {
                    "created": "1",
                    "eisa": eisa_id,
                    "paper": new_paper.id,
                    "pool": pool_size,
                    "marks": total_marks,
                }
            )
            snapshot_url = f"{reverse('assessor_randomized_snapshot', args=[new_assessment.id])}?{query_string}"
            return redirect(snapshot_url)

    base_for_pool = base_paper or paper
    if not module_name_meta:
        if base_for_pool and getattr(base_for_pool, "module_name", None):
            module_name_meta = base_for_pool.module_name
        elif base_extractor and getattr(base_extractor, "module_name", None):
            module_name_meta = base_extractor.module_name
        elif assessment.module_name:
            module_name_meta = assessment.module_name
    if not module_number_meta:
        if base_for_pool and getattr(base_for_pool, "paper_letter", None):
            module_number_meta = base_for_pool.paper_letter
        elif base_extractor and getattr(base_extractor, "paper_letter", None):
            module_number_meta = base_extractor.paper_letter
        elif assessment.module_number:
            module_number_meta = assessment.module_number

    if paper:
        raw_node_tree, node_stats = build_node_tree(paper)

        def filter_nodes(nodes, seen_numbers):
            filtered = []
            for node in nodes:
                node_type = (node.get("node_type") or "").lower()
                children = node.get("children") or []
                filtered_children = (
                    filter_nodes(children, seen_numbers) if children else []
                )
                include = True
                if node_type == "question":
                    number = (node.get("number") or "").strip()
                    if number:
                        if number in seen_numbers:
                            include = False
                        else:
                            seen_numbers.add(number)
                if include:
                    new_node = dict(node)
                    new_node["children"] = filtered_children
                    filtered.append(new_node)
            return filtered

        node_tree = filter_nodes(raw_node_tree, set())
    else:
        node_tree = []
        node_stats = {
            "total": 0,
            "questions": 0,
            "tables": 0,
            "images": 0,
            "instructions": 0,
        }

    cover_nodes = [
        node
        for node in node_tree
        if (node.get("node_type") or "").lower() != "question"
    ]
    question_nodes = [
        node
        for node in node_tree
        if (node.get("node_type") or "").lower() == "question"
    ]

    snapshots_qs = (
        Assessment.objects.filter(
            paper_type="randomized",
            paper_link__isnull=False,
        )
        .select_related("paper_link", "qualification")
        .order_by("-created_at")
    )
    if not (user.is_staff or user.is_superuser or user_role in {"admin", "moderator"}):
        if user_qualification:
            snapshots_qs = snapshots_qs.filter(qualification=user_qualification)
        else:
            snapshots_qs = snapshots_qs.filter(created_by=user)
    snapshots_total = snapshots_qs.count()
    created_flag = str(request.GET.get("created", "0")).lower() in {"1", "true", "yes"}
    created_snapshot_details = {}
    if created_flag:
        created_snapshot_details = {
            "eisa": request.GET.get("eisa") or assessment.eisa_id,
            "paper_id": request.GET.get("paper") or (paper.id if paper else ""),
            "pool": request.GET.get("pool"),
            "marks": request.GET.get("marks"),
        }

    context = {
        "assessment": assessment,
        "paper": paper,
        "node_tree": node_tree,
        "node_stats": node_stats,
        "can_open_pipeline": assessment.extractor_paper is not None,
        "randomized_snapshots": snapshots_qs,
        "random_meta": random_meta,
        "cover_nodes": cover_nodes,
        "question_nodes": question_nodes,
        "base_paper": base_paper,
        "module_name_meta": module_name_meta,
        "module_number_meta": module_number_meta,
        "random_status_label": random_status_label,
        "allowed_letters_meta": allowed_letter_options,
        "allowed_letters_display": allowed_letters_display,
        "qual_name_meta": qual_name,
        "cover_heading_meta": cover_heading,
        "current_snapshot_id": assessment.id,
        "created_snapshot": created_flag,
        "created_snapshot_details": created_snapshot_details,
        "snapshots_total": snapshots_total,
        # NEW: Memo-related context variables
        "can_forward_to_moderator": can_forward_to_moderator,
        "memo_requirement_message": memo_requirement_message,
        "has_memo": has_memo,
        "memo_filename": assessment.memo_file.name if assessment.memo_file else None,
    }
    return render(request, "core/assessor-developer/randomized_snapshot.html", context)


@login_required
def _moderator_dashboard_context(request):
    pending_statuses = [
        "Pending",
        "pending_moderation",
        "Submitted to Moderator",
    ]
    moderated_statuses = [
        "Approved by Moderator",
        "Submitted to QCTO",
        "pending_qcto",
        "qcto_approved",
        "Submitted to ETQA",
        "pending_etqa",
        "Approved by ETQA",
        "etqa_approved",
        "Moderated",
        "moderated",
        "Released to students",
        "Approved",
    ]
    
    # Statuses for achieved assessments (approved or returned by moderator)
    achieved_statuses = [
        "Submitted to QCTO",
        "Returned for Changes",
        "Approved by Moderator",
        "Moderated",
    ]

    pending_assessments = (
        Assessment.objects.select_related("qualification", "created_by")
        .filter(
            status__in=pending_statuses,
            paper__isnull=False,
        )
        .order_by("-created_at")
    )

    forwarded_assessments = (
        Assessment.objects.select_related("qualification", "created_by")
        .filter(
            status__in=["Submitted to ETQA", "pending_etqa"],
            paper__isnull=False,
        )
        .order_by("-created_at")
    )
    
    # New: Achieved assessments (moderated assessments)
    achieved_assessments = (
        Assessment.objects.select_related("qualification", "created_by")
        .filter(
            status__in=achieved_statuses,
            paper__isnull=False,
        )
        .order_by("-status_changed_at")
    )

    feedback_qs = Feedback.objects.select_related("assessment").order_by("-created_at")

    today = now().date()
    stats = {
        "pending_reviews": pending_assessments.count(),
        "moderated_today": Assessment.objects.filter(
            status__in=moderated_statuses,
            status_changed_at__date=today,
        ).count(),
        "total_moderated": Assessment.objects.filter(
            status__in=moderated_statuses
        ).count(),
        "feedback_given": feedback_qs.count(),
        "achieved_assessments": achieved_assessments.count(),  # New stat
    }

    return {
        "pending_assessments": pending_assessments,
        "forwarded_assessments": forwarded_assessments,
        "achieved_assessments": achieved_assessments,  # New context variable
        "recent_feedback": feedback_qs[:10],
        "feedback_entries": feedback_qs,
        "stats": stats,
        "today": now(),
    }



@login_required
def moderator_developer_dashboard(request):
    context = _moderator_dashboard_context(request)
    context["page_title"] = "Moderator Dashboard"
    return render(request, "core/moderator/moderator_developer.html", context)


@login_required
def moderator_review_list(request):
    context = _moderator_dashboard_context(request)
    context["page_title"] = "Review Assessments"
    context["show_actions"] = True
    return render(request, "core/moderator/moderator_developer.html", context)


@login_required
def moderator_approve_reject(request):
    context = _moderator_dashboard_context(request)
    context["page_title"] = "Approve or Reject"
    context["show_actions"] = True
    return render(request, "core/moderator/approve_reject.html", context)


@login_required
def moderator_feedback(request):
    context = _moderator_dashboard_context(request)
    context["page_title"] = "Feedback & Comments"
    return render(request, "core/moderator/feedback.html", context)


@login_required
def moderator_history(request):
    context = _moderator_dashboard_context(request)
    context["page_title"] = "Moderation History"
    return render(request, "core/moderator/history.html", context)


@login_required
def moderator_reports(request):
    context = _moderator_dashboard_context(request)
    context["page_title"] = "Quality Reports"
    return render(request, "core/moderator/reports.html", context)


@login_required
def moderator_statistics(request):
    context = _moderator_dashboard_context(request)
    context["page_title"] = "Statistics"
    return render(request, "core/moderator/statistics.html", context)


@login_required
def assessor_developer_paper(request, paper_id):
    extractor_paper = get_object_or_404(
        ExtractorPaper.objects.select_related("assessment_record"), id=paper_id
    )
    assessment = getattr(extractor_paper, "assessment_record", None)
    if assessment:
        user = request.user
        user_role = getattr(user, "role", "")
        user_qualification = getattr(user, "qualification", None)
        if not (
            user.is_staff or user.is_superuser or user_role in {"admin", "moderator"}
        ):
            allowed = assessment.created_by_id == user.id
            if (
                not allowed
                and user_qualification
                and assessment.qualification_id
                == getattr(user_qualification, "id", None)
            ):
                allowed = True
            if not allowed:
                raise Http404("Paper not available")
        refresh = request.GET.get("refresh") == "1"
        rebuild_extractor_from_bank(assessment, force=refresh)
    return extractor_views.paper_view(request, paper_id)


@require_http_methods(["GET", "POST"])
def upload_assessment(request):
    """Enhanced upload using robust extractor"""
    qualifications = Qualification.objects.all()
    submissions = Assessment.objects.select_related(
        "qualification", "created_by"
    ).order_by("-created_at")
    module_map = qualification_registry.module_map_by_pk()
    saqa_map = qualification_registry.saqa_map_by_pk()

    if request.method == "POST":
        action = (request.POST.get("action") or "save").strip().lower()
        eisa_id = f"EISA-{uuid.uuid4().hex[:8].upper()}"
        qual_id = request.POST.get("qualification")
        qualification_obj = get_object_or_404(Qualification, pk=qual_id)

        paper_number = (request.POST.get("paper_number") or "").strip()
        module_number = (request.POST.get("module_number") or "").strip()
        module_name = (request.POST.get("module_name") or "").strip()
        if not module_name:
            for option in module_map.get(str(qualification_obj.pk), []):
                if option.get("code") == module_number:
                    module_name = option.get("label") or module_number
                    break
        saqa = (
            request.POST.get("saqa_id") or saqa_map.get(str(qualification_obj.pk)) or ""
        ).strip()
        file = request.FILES.get("file_input")
        memo = request.FILES.get("memo_file")
        comment = (request.POST.get("comment_box") or "").strip()

        if not paper_number or not module_number or not module_name or not memo:
            messages.error(
                request, "Please complete all required fields and upload the memo."
            )
            return redirect("upload_assessment")

        if not file or not file.name.lower().endswith(".docx"):
            messages.error(request, "Please upload a .docx assessment file.")
            return redirect("upload_assessment")

        try:
            file.seek(0)
        except Exception:
            pass

        try:
            paper_obj = save_robust_extraction_to_db(
                docx_file=file,
                paper_name=paper_number,
                qualification=qualification_obj,
                user=request.user,
                use_gemini=False,
                use_gemma=False,
            )
        except Exception as exc:
            traceback.print_exc()
            messages.error(request, f"Extraction failed: {exc}")
            return redirect("upload_assessment")

        if not paper_obj:
            messages.error(request, "Extraction failed - please try again")
            return redirect("upload_assessment")

        if paper_obj and hasattr(paper_obj, "extract_dir"):
            media_dir = os.path.join(paper_obj.extract_dir, "media")
            copy_images_to_media_folder(media_dir)

        try:
            file.seek(0)
        except Exception:
            pass
        if memo:
            try:
                memo.seek(0)
            except Exception:
                pass

        status_lookup = {
            "etqa": "Submitted to ETQA",
            "assessor_dev": "Submitted to Moderator",
        }
        status = status_lookup.get(action, "draft")

        default_module_code = next(
            (
                opt.get("code")
                for opt in module_map.get(str(qualification_obj.pk), [])
                if opt.get("code")
            ),
            "1A",
        )
        assessment_obj = Assessment.objects.create(
            eisa_id=eisa_id,
            qualification=qualification_obj,
            paper=paper_number,
            module_number=module_number or default_module_code,
            module_name=module_name or qualification_obj.name,
            saqa_id=saqa,
            file=file,
            memo=memo,
            comment=comment,
            forward_to_moderator=(action == "assessor_dev"),
            created_by=request.user,
            paper_link=paper_obj,
            status=status,
        )

        sync_assessment_paper_bank(assessment_obj, force=True)

        nodes = ExamNode.objects.filter(paper=paper_obj)
        questions = nodes.filter(node_type="question").count()
        tables = nodes.filter(node_type="table").count()
        images = nodes.filter(node_type="image").count()

        if action == "assessor_dev":
            extractor = rebuild_extractor_from_bank(assessment_obj, force=True)
            if extractor:
                messages.success(
                    request,
                    (
                        "Assessment forwarded to Assessor Developer. "
                        f"Extraction saved ({questions} questions, {tables} tables, {images} images)."
                    ),
                )
            else:
                messages.warning(
                    request,
                    "Assessment queued for Assessor Developer, but the advanced view could not be prepared automatically.",
                )
            return redirect("upload_assessment")

        if action == "etqa":
            messages.success(
                request,
                (
                    "Assessment forwarded to ETQA. "
                    f"Extraction saved ({questions} questions, {tables} tables, {images} images)."
                ),
            )
            return redirect("upload_assessment")

        messages.success(
            request,
            (
                "Assessment uploaded and saved as draft. "
                f"Extraction captured ({questions} questions, {tables} tables, {images} images)."
            ),
        )
        return redirect("upload_assessment")

    return render(
        request,
        "core/assessor-developer/upload_assessment.html",
        {
            "submissions": submissions,
            "qualifications": qualifications,
            "module_map_json": json.dumps(module_map),
            "saqa_map_json": json.dumps(saqa_map),
        },
    )


@require_http_methods(["POST"])
@login_required
def moderate_assessment(request, eisa_id):
    # Fetch the assessment
    assessment = get_object_or_404(Assessment, eisa_id=eisa_id)

    # Get the action (approve or return)
    action = request.POST.get("action")

    if action == "approve":
        if "moderator_report" not in request.FILES:
            messages.error(request, "Please upload a Word report before approving.")
            return redirect("moderator_review_list")
        report_file = request.FILES["moderator_report"]
        allowed_extensions = [".doc", ".docx"]
        file_extension = os.path.splitext(report_file.name)[1].lower()
        if file_extension not in allowed_extensions:
            messages.error(request, "Please upload a Word document (.doc or .docx).")
            return redirect("moderator_review_list")
        assessment.moderator_report.save(
            f"moderator_report_{eisa_id}{file_extension}", report_file
        )
        assessment.status = "Submitted to QCTO"
        assessment.forward_to_moderator = False
        assessment.status_changed_at = now()
        assessment.status_changed_by = request.user
        assessment.save(
            update_fields=[
                "status",
                "forward_to_moderator",
                "status_changed_at",
                "status_changed_by",
                "moderator_report",
            ]
        )
        messages.success(
            request, f"{assessment.eisa_id} approved and forwarded to QCTO."
        )
        return redirect("moderator_review_list")

    if action == "return":
        assessment.status = "Returned for Changes"
        assessment.forward_to_moderator = False
        assessment.status_changed_at = now()
        assessment.status_changed_by = request.user
        assessment.save(
            update_fields=[
                "status",
                "forward_to_moderator",
                "status_changed_at",
                "status_changed_by",
            ]
        )
        messages.success(request, f"{assessment.eisa_id} returned for changes.")
        return redirect("moderator_review_list")

    messages.error(request, "Invalid action.")
    return redirect("moderator_review_list")


@login_required
def download_moderator_report(request, eisa_id):
    assessment = get_object_or_404(Assessment, eisa_id=eisa_id)
    
    if not assessment.moderator_report:
        messages.error(request, "No moderator report available for download.")
        return redirect("moderator_review_list")
    
    # You can add additional download logic here if needed
    return redirect("moderator_review_list")


@require_http_methods(["POST"])
def add_feedback(request, eisa_id):
    a = get_object_or_404(Assessment, eisa_id=eisa_id)
    to = request.POST.get("to_user", "").strip()
    msg = request.POST.get("message", "").strip()
    status = request.POST.get("status", "Pending")

    if not to or not msg:
        messages.error(request, "Both recipient and message are required.")
    else:
        Feedback.objects.create(assessment=a, to_user=to, message=msg, status=status)
        messages.success(request, "Feedback added.")
    return redirect("moderator_developer")


# checklist

from .models import ChecklistItem


def toggle_checklist_item(request, item_id):
    try:
        item = ChecklistItem.objects.get(pk=item_id)
        item.completed = not item.completed
        item.save()
        return JsonResponse({"status": "ok", "completed": item.completed})
    except ChecklistItem.DoesNotExist:
        raise Http404()


# not neccessary---not main
def checklist_stats(request):
    total = ChecklistItem.objects.count()
    done = ChecklistItem.objects.filter(completed=True).count()
    pending = total - done
    return JsonResponse({"total": total, "completed": done, "pending": pending})


# 1) QCTO Dashboard: list assessments submitted to ETQA
# QCTO Dashboard: list assessments submitted to ETQA
@require_http_methods(["GET"])
def qcto_dashboard(request):
    """
    QCTO dashboard: list assessments submitted by the moderator for QCTO review.
    Only those with status 'Submitted to QCTO' appear here.
    """
    pending_assessments = Assessment.objects.filter(
        status="Submitted to QCTO"
    ).order_by("-created_at")
    return render(
        request,
        "core/qcto/qcto_dashboard.html",
        {"pending_assessments": pending_assessments},
    )


# 2) QCTO Moderate Assessment: view + update status and notes


@require_http_methods(["GET", "POST"])
def qcto_moderate_assessment(request, eisa_id):
    """
    QCTO review step: only handles assessments with status 'Submitted to QCTO'.
    On 'approve', status becomes 'Submitted to ETQA'; on 'reject', status becomes 'Rejected'.
    """
    assessment = get_object_or_404(Assessment, eisa_id=eisa_id)

    if assessment.status != "Submitted to QCTO":

        messages.error(request, "This assessment is not pending QCTO review.")
        return redirect("qcto_dashboard")

    if request.method == "POST":
        notes = request.POST.get("qcto_notes", "").strip()
        decision = request.POST.get("decision")  # now 'approve' or 'reject'

        if decision == "approve":
            assessment.status = "Submitted to ETQA"
            assessment.status_changed_at = now()
            assessment.status_changed_by = (
                request.user if request.user.is_authenticated else None
            )
            messages.success(
                request, f"{assessment.eisa_id} approved and forwarded to ETQA."
            )
        elif decision == "reject":
            assessment.status = "Rejected"
            assessment.status_changed_at = now()
            assessment.status_changed_by = (
                request.user if request.user.is_authenticated else None
            )
            messages.success(request, f"{assessment.eisa_id} has been rejected.")
        else:
            messages.error(request, "Invalid decision.")
            return redirect("qcto_moderate_assessment", eisa_id=eisa_id)

        assessment.qcto_notes = notes
        assessment.save(
            update_fields=[
                "status",
                "status_changed_at",
                "status_changed_by",
                "qcto_notes",
            ]
        )
        return redirect("qcto_dashboard")

    # on GET, render the form
    return render(
        request,
        "core/qcto/qcto_moderate_assessment.html",
        {
            "assessment": assessment,
            "decision_choices": [
                ("approve", "Approve"),
                ("reject", "Reject"),
            ],
        },
    )


# 3) QCTO Reports: summary of QCTO-approved assessments
@require_http_methods(["GET"])
def qcto_reports(request):
    stats = (
        Assessment.objects.filter(status="Approved by ETQA")
        .values("qualification")
        .annotate(validated_count=Count("id"))
        .order_by("qualification")
    )
    return render(request, "core/qcto/qcto_reports.html", {"stats": stats})


# 4) QCTO Compliance & Reports: overview all assessments


@require_http_methods(["GET", "POST"])
def qcto_compliance(request):
    if request.method == "POST":
        eisa_id = request.POST.get("eisa_id")
        assessment = get_object_or_404(Assessment, eisa_id=eisa_id)

        # read the notes and decision from the form
        notes = request.POST.get("qcto_notes", "").strip()
        decision = request.POST.get("decision")

        # update status and save
        if decision == "approve":
            assessment.status = "Submitted to QCTO"
            messages.success(request, f"{eisa_id} sent to QCTO Final Review.")
        elif decision == "reject":
            assessment.status = "Non-compliant"
            messages.success(request, f"{eisa_id} marked non-compliant.")
        else:
            messages.error(request, "Invalid decision.")
            return redirect("qcto_compliance")

        assessment.qcto_notes = notes
        assessment.save()
        # Redirect to refresh the page so it is removed from the list
        return redirect("qcto_compliance")
        # Redirect to refresh the page so it is removed from the list
        return redirect("qcto_compliance")

    # GET: show only those not yet reviewed
    assessments = Assessment.objects.exclude(
        status__in=[
            "Submitted to QCTO",
            "Submitted to ETQA",
            "Approved by ETQA",
            "Non-compliant",
            "Rejected",
        ]
    ).order_by("-created_at")

    return render(
        request, "core/qcto/qcto_compliance.html", {"assessments": assessments}
    )


# 5) QCTO Final Assessment Review: list for QCTO decision
# views.py
@login_required
@require_http_methods(["GET", "POST"])
def qcto_assessment_review(request):
    if request.method == "POST":
        # Check if this is an advanced paper view request
        action = request.POST.get("action")
        assessment_id = request.POST.get("assessment_id")

        if action == "open_pipeline" and assessment_id:
            # Handle advanced paper view
            assessment = get_object_or_404(Assessment, pk=assessment_id)
            extractor_paper = rebuild_extractor_from_bank(assessment, force=False)
            if extractor_paper:
                return redirect("assessor_developer_paper", paper_id=extractor_paper.id)
            messages.error(request, "Unable to prepare paper for advanced pipeline.")
            return redirect("qcto_assessment_review")

        # Existing QCTO approval/rejection logic
        eisa_id = request.POST.get("eisa_id")
        assessment = get_object_or_404(Assessment, eisa_id=eisa_id)

        # Get the action (approve or reject)
        action = request.POST.get("action")

        # Check if a file was uploaded
        if "qcto_report" not in request.FILES:
            messages.error(request, "Please upload a QCTO report before taking action.")
            return redirect("qcto_assessment_review")

        report_file = request.FILES["qcto_report"]

        # Validate file type (Word documents)
        allowed_extensions = [".doc", ".docx"]
        file_extension = os.path.splitext(report_file.name)[1].lower()

        if file_extension not in allowed_extensions:
            messages.error(request, "Please upload a Word document (.doc or .docx).")
            return redirect("qcto_assessment_review")

        # Save the QCTO report
        assessment.qcto_report.save(
            f"qcto_report_{eisa_id}{file_extension}", report_file
        )

        # Update status based on action
        if action == "approve":
            assessment.status = "Submitted to ETQA"
            messages.success(request, f"{eisa_id} approved and forwarded to ETQA.")
        elif action == "reject":
            assessment.status = "Rejected"
            messages.success(request, f"{eisa_id} has been rejected.")
        else:
            messages.error(request, "Invalid decision.")
            return redirect("qcto_assessment_review")

        assessment.save()
        return redirect("qcto_assessment_review")

    # GET: only show those pending QCTO
    assessments = Assessment.objects.filter(status="Submitted to QCTO").order_by(
        "-created_at"
    )
    return render(
        request, "core/qcto/assessment_review.html", {"assessments": assessments}
    )


# 6) QCTO Archive: list archived QCTO decisions
@require_http_methods(["GET"])
def qcto_archive(request):
    archives = Assessment.objects.filter(
        status__in=["Approved by ETQA", "Rejected"]
    ).order_by("-created_at")
    return render(request, "core/qcto/qcto_archive.html", {"archives": archives})


# 7) QCTO View Single Assessment Details
def qcto_view_assessment(request, eisa_id):
    assessment = get_object_or_404(Assessment, eisa_id=eisa_id)
    generated_questions = GeneratedQuestion.objects.filter(assessment=assessment)
    return render(
        request,
        "core/qcto/qcto_view_assessment.html",
        {"assessment": assessment, "generated_questions": generated_questions},
    )


# 8) QCTO Assessment view logic
@require_http_methods(["GET"])
def qcto_latest_assessment_detail(request):
    latest = (
        Assessment.objects.filter(status="Approved by ETQA")
        .order_by("-created_at")
        .first()
    )

    questions = GeneratedQuestion.objects.filter(assessment=latest)

    return render(
        request,
        "core/qcto/qcto_view_assessment.html",
        {
            "assessment": latest,
            "generated_questions": questions,
        },
    )


# ****************************************************************************
# Default Page is added here...
def default_page(request):
    return render(request, "core/login/awaiting_activation.html")


# ****************************************************************************
# ****************************************************************************
#####Approved assessments view for the paper to be easily pulled and used for other uses.
def approved_assessments_view(request):
    assessments = Assessment.objects.filter(status="Approved by ETQA").prefetch_related(
        "generated_questions"
    )
    return render(
        request, "core/approved_assessments.html", {"assessments": assessments}
    )


#########################################################################################################################
# ASSessment Progress tracker view
@login_required
def assessment_progress_tracker(request):
    archived_statuses = [
        "Approved by ETQA",
        "Rejected",
        "Submitted to ETQA",
        "Approved by Moderator",
        "Submitted to Moderator",
        "Submitted to QCTO",
        "pending_qcto",
        "pending_etqa",
    ]

    assessments = Assessment.objects.filter(status__in=archived_statuses).order_by(
        "-created_at"
    )

    # Dynamically add the `currently_with` field to each assessment
    for a in assessments:
        a.currently_with = get_current_holder(a.status)

    return render(
        request,
        "core/paper_tracking/assessment_progress_tracker.html",
        {
            "assessments": assessments,
        },
    )


# View to get current holder of the paper
def get_current_holder(status):
    mapping = {
        "Pending": "Assessor/Developer",
        "Submitted to Moderator": "Moderator",
        "pending_moderation": "Moderator",
        "Returned for Changes": "Assessor/Developer",
        "Approved by Moderator": "QCTO",
        "Submitted to QCTO": "QCTO",
        "pending_qcto": "QCTO",
        "Submitted to ETQA": "ETQA",
        "pending_etqa": "ETQA",
        "Approved by ETQA": "Archived",
        "Rejected": "Archived",
    }
    return mapping.get(status, "Unknown")


# to summarise the tracker now tells us where the paper is and displays the approved papers so they can be pulled.


# Learner Assessment viewing view
# @login_required
def approved_assessments_for_learners(request):
    user_qualification = request.user.qualification
    assessments = Assessment.objects.filter(
        status="Approved by ETQA", qualification=user_qualification  # fetching from the
    ).prefetch_related("generated_questions", "case_study")

    return render(
        request,
        "learner/approved_assessments_by_qualification.html",
        {
            "assessments": assessments,
        },
    )


# _____________________________________________________________________________________


@login_required
def etqa_dashboard(request):
    centers = AssessmentCentre.objects.all()
    qualifications = Qualification.objects.all()
    assessments_for_etqa = Assessment.objects.filter(status="Submitted to ETQA")

    # 1) Figure out which qualification we're working with
    selected_qualification = (
        request.GET.get("qualification_id") or request.POST.get("qualification") or ""
    )

    # 2) Always load the APPROVED assessments for that qualification
    # approved_assessments = []
    # if selected_qualification:
    #     approved_assessments = Assessment.objects.filter(
    #         qualification_id=selected_qualification,
    #         status="Submitted to ETQA"
    #     )

    approved_assessments = []
    if selected_qualification:
        approved_assessments = Assessment.objects.filter(
            qualification_id=selected_qualification,
            status="Submitted to ETQA",
            is_selected_by_etqa=True,
        )

    created_batch = None

    if request.method == "POST":
        # 3) Simple presence check
        missing = [
            f
            for f in ("center", "qualification", "assessment", "date")
            if f not in request.POST
        ]
        if missing:
            return render(
                request,
                "core/etqa/etqa_dashboard.html",
                {
                    "centers": centers,
                    "qualifications": qualifications,
                    "selected_qualification": selected_qualification,
                    "approved_assessments": approved_assessments,
                    "assessments_for_etqa": assessments_for_etqa,
                    "error": f"Missing: {missing[0]}",
                },
            )

        # 4) Create the batch
        batch = Batch.objects.create(
            center_id=request.POST["center"],
            qualification_id=request.POST["qualification"],
            assessment_id=request.POST["assessment"],
            assessment_date=request.POST["date"],
            # number_of_learners  = request.POST['number_of_learners'],
        )
        created_batch = batch
        # Note: no redirect here; we simply fall through and re-render

    return render(
        request,
        "core/etqa/etqa_dashboard.html",
        {
            "centers": centers,
            "qualifications": qualifications,
            "selected_qualification": selected_qualification,
            "approved_assessments": approved_assessments,
            "assessments_for_etqa": assessments_for_etqa,
            "created_batch": created_batch,
        },
    )


@login_required
def toggle_selection_by_etqa(request, assessment_id):
    assessment = get_object_or_404(Assessment, id=assessment_id)
    attempt_number = int(request.POST.get("attempt_number", 1))
    assessment.is_selected_by_etqa = not assessment.is_selected_by_etqa
    assessment.save()
    return redirect("etqa_assessment_view")


@login_required
def etqa_assessment_view(request):
    submitted_assessments = Assessment.objects.filter(status="Submitted to ETQA")
    return render(
        request,
        "core/etqa/etqa_assessment.html",
        {"assessments": submitted_assessments},
    )

#New view############################################################
@login_required
def download_memo(request, assessment_id):
    """Download memo file for an assessment - UPDATED FOR RANDOMIZED PAPERS"""
    assessment = get_object_or_404(Assessment, id=assessment_id)
    
    # Check permissions - only allow moderators and authorized users
    user = request.user
    user_role = getattr(user, "role", "")
    if not (
        user.is_staff
        or user.is_superuser
        or user_role in {"admin", "moderator", "etqa", "qcto", "assessment_center"}
        or assessment.created_by_id == user.id
    ):
        raise Http404("Memo not available")
    
    # For randomized papers, use memo_file field
    if assessment.paper_type == 'randomized':
        if not assessment.memo_file:
            messages.error(request, "No memo file available for this assessment.")
            return redirect("moderator_review_list")
        
        try:
            # Get the file path
            file_path = assessment.memo_file.path
            
            # Check if file exists
            if not os.path.exists(file_path):
                messages.error(request, "Memo file not found on server.")
                return redirect("moderator_review_list")
            
            # Open the file for reading
            with open(file_path, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/octet-stream')
                
                # Get filename from the memo_file field
                filename = os.path.basename(assessment.memo_file.name)
                
                # Set content disposition for download
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                return response
                
        except Exception as e:
            messages.error(request, f"Error downloading memo: {str(e)}")
            return redirect("moderator_review_list")
    
    # For admin uploads, use memo field (existing logic)
    else:
        if not assessment.memo:
            messages.error(request, "No memo file available for this assessment.")
            return redirect("moderator_review_list")
        
        try:
            # Get the file path
            file_path = assessment.memo.path
            
            # Check if file exists
            if not os.path.exists(file_path):
                messages.error(request, "Memo file not found on server.")
                return redirect("moderator_review_list")
            
            # Open the file for reading
            with open(file_path, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/octet-stream')
                
                # Get filename from the memo field
                filename = os.path.basename(assessment.memo.name)
                
                # Set content disposition for download
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                return response
                
        except Exception as e:
            messages.error(request, f"Error downloading memo: {str(e)}")
            return redirect("moderator_review_list")



# NEW VIEW: Simple achieved page
@login_required
def achieved_assessments_view(request):
    # Show assessments that are selected by ETQA
    achieved_assessments = Assessment.objects.filter(is_selected_by_etqa=True)
    
    return render(
        request,
        "core/etqa/achieved_assessments.html",
        {"assessments": achieved_assessments},
    )

def assessment_center_view(request):
    batches = Batch.objects.filter(submitted_to_center=True)
    return render(request, "assessment_center.html", {"batches": batches})


def submit_to_center(request, batch_id):
    batch = Batch.objects.get(id=batch_id)
    batch.submitted_to_center = True
    batch.save()
    return redirect("etqa_dashboard")


@login_required
def release_assessment_to_students(request, assessment_id):
    if request.method == "POST":
        assessment = get_object_or_404(Assessment, id=assessment_id)
        if assessment.status == "Submitted to ETQA":
            assessment.status = "Released to students"
            assessment.save()
    return redirect("assessment_center")


def _preview_structure_for_paper(paper):
    """
    Return an ordered list of blocks for preview along with stats.
    Prefers the stored structure_json (original capture order) and falls back to ExamNodes.
    """
    raw = getattr(paper, "structure_json", None)
    candidate_blocks = None
    if isinstance(raw, list):
        candidate_blocks = raw
    elif isinstance(raw, dict):
        for key in ("blocks", "structure", "nodes", "sequence", "items"):
            maybe = raw.get(key)
            if isinstance(maybe, list):
                candidate_blocks = maybe
                break

    stats = {
        "total": 0,
        "questions": 0,
        "tables": 0,
        "images": 0,
        "instructions": 0,
    }

    def transform(block):
        data = deepcopy(block)
        node_type = (data.get("node_type") or data.get("type") or "").strip()
        data["node_type"] = node_type
        data.setdefault("text", data.get("text") or data.get("body") or "")
        content = data.get("content")
        if not isinstance(content, list):
            content = []
        data["content"] = content
        children = data.get("children")
        if isinstance(children, list):
            data["children"] = [
                transform(child) for child in children if isinstance(child, dict)
            ]
        else:
            data["children"] = []

        stats["total"] += 1
        lowered = node_type.lower()
        if lowered == "question":
            stats["questions"] += 1
        elif lowered == "table":
            stats["tables"] += 1
        elif lowered == "image":
            stats["images"] += 1
        else:
            stats["instructions"] += 1
        return data

    if candidate_blocks:
        tree = [
            transform(block) for block in candidate_blocks if isinstance(block, dict)
        ]
        return tree, stats

    fallback_tree, fallback_stats = build_node_tree(paper)
    if isinstance(fallback_stats, dict):
        stats.update({k: fallback_stats.get(k, stats[k]) for k in stats})
    return fallback_tree, stats


@login_required
def review_saved_selector(request):
    """View to select a saved paper for review"""

    filter_type = request.GET.get("type", "")

    # Get papers query
    all_papers = Paper.objects.order_by("-created_at")

    # Filter papers
    original_papers = all_papers.filter(is_randomized=False)
    randomized_papers = all_papers.filter(is_randomized=True)

    # Get qualifications for filtering
    qualifications = Qualification.objects.all()

    selected_paper = None
    node_sequence = []
    node_stats = {
        "total": 0,
        "questions": 0,
        "tables": 0,
        "images": 0,
        "instructions": 0,
    }

    if request.method == "POST":
        paper_id = request.POST.get("paper_id")
        if paper_id:
            selected_paper = Paper.objects.filter(pk=paper_id).first()
            if selected_paper:
                node_tree, extracted_stats = _preview_structure_for_paper(
                    selected_paper
                )
                node_sequence = node_tree
                if isinstance(extracted_stats, dict):
                    node_stats = extracted_stats
            return render(
                request,
                "core/administrator/review_saved_selector.html",
                {
                    "original_papers": original_papers,
                    "randomized_papers": randomized_papers,
                    "papers": all_papers,  # For dropdown
                    "qualifications": qualifications,
                    "filter_type": filter_type,
                    "selected_paper": selected_paper,
                    "node_sequence": node_sequence if selected_paper else [],
                    "node_stats": node_stats,
                },
            )

    return render(
        request,
        "core/administrator/review_saved_selector.html",
        {
            "original_papers": original_papers,
            "randomized_papers": randomized_papers,
            "papers": all_papers,  # For dropdown
            "qualifications": qualifications,
            "filter_type": filter_type,
            "selected_paper": selected_paper,
            "node_sequence": node_sequence,
            "node_stats": node_stats,
        },
    )


@login_required
def load_saved_paper_view(request, paper_pk):
    """Preview a saved paper (full flow-by-flow reading view), with optional editor toggle."""
    try:
        paper = get_object_or_404(
            Paper.objects.select_related("qualification"), id=paper_pk
        )

        # Prefer the captured structure_json for preview fidelity
        node_sequence_raw, node_stats = _preview_structure_for_paper(paper)

        assessment = Assessment.objects.filter(paper_link=paper).first()

        # Editor toggle (for advanced/manual block fixing if ever needed)
        use_editor = request.GET.get("editor") == "1"
        if use_editor:
            editor_tree, editor_stats = build_node_tree(paper)
            questions = [
                n
                for n in editor_tree
                if (n.get("node_type") or "").lower() == "question"
            ]
            context = {
                "paper": paper,
                "questions": questions,
                "assessment": assessment,
                "total_marks": paper.total_marks,
                "node_count": editor_stats.get("total", 0),
                "question_nodes": questions,
                "questions_count": len(questions),
                "tables_count": editor_stats.get("tables", 0),
                "images_count": editor_stats.get("images", 0),
            }
            return render(request, "core/administrator/review_paper.html", context)

        context = {
            "paper": paper,
            "assessment": assessment,
            "node_sequence": node_sequence_raw,
            "node_stats": node_stats,
        }
        return render(request, "core/administrator/review_paper_readonly.html", context)

    except Exception as e:
        print(traceback.format_exc())
        messages.error(request, f"Error loading paper: {str(e)}")
        return redirect("review_saved_selector")


@login_required
def student_dashboard(request):
    """Dashboard view for students/learners to see their available assessments"""

    user_qualification = request.user.qualification

    # Only include assessments with papers and nodes
    assessments = Assessment.objects.filter(
        status="Released to students",
        paper_link__isnull=False,
    )

    if user_qualification:
        assessments = assessments.filter(
            Q(qualification=user_qualification)
            | Q(
                qualification__isnull=True, paper_link__qualification=user_qualification
            )
        )

    assessments = assessments.select_related("qualification", "paper_link").order_by(
        "-created_at"
    )

    assessment_data = []
    for assessment in assessments:
        paper = assessment.paper_link

        question_count = 0
        if paper:
            preview_tree, preview_stats = _preview_structure_for_paper(paper)
            if isinstance(preview_stats, dict):
                question_count = preview_stats.get("questions", 0) or 0
            if not question_count and isinstance(preview_tree, list):
                question_count = sum(
                    1
                    for node in preview_tree
                    if (node.get("node_type") or "").lower() == "question"
                )

        if question_count <= 0:
            question_count = assessment.generated_questions.count()

        submission_count = ExamSubmission.objects.filter(
            assessment=assessment, student=request.user
        ).count()

        if submission_count > 0:
            continue

        assessment_data.append(
            {
                "assessment": assessment,
                "attempt_count": submission_count,
                "can_attempt": submission_count == 0,
                "question_count": question_count,
                "is_randomized": assessment.paper_type == "randomized",
            }
        )

    return render(
        request,
        "core/student/dashboard.html",
        {
            "assessment_data": assessment_data,
            "user": request.user,
            "qualification": user_qualification,
        },
    )


def custom_logout(request):
    logout(request)
    return redirect("custom_login")


@login_required
def randomize_paper_structure_view(request, paper_pk):
    """Create a randomized paper and push it through the advanced pipeline."""
    try:
        original_paper = get_object_or_404(Paper, pk=paper_pk)

        if original_paper.is_randomized:
            messages.error(request, "Cannot randomize an already randomized paper")
            return redirect("review_saved_selector")

        randomized_paper = Paper.objects.create(
            name=f"{original_paper.name} (Randomized)",
            qualification=original_paper.qualification,
            created_by=original_paper.created_by,
            is_randomized=True,
        )

        source_assessment = original_paper.assessments.order_by("-created_at").first()
        qualification = (
            source_assessment.qualification
            if source_assessment
            else original_paper.qualification
        )
        paper_meta = {}
        if isinstance(original_paper.structure_json, dict):
            paper_meta = original_paper.structure_json.get("randomization", {}) or {}

        module_catalog = qualification_registry.get_module_choices(
            qualification.name if qualification else ""
        )
        default_module_code = next(
            (mod.get("code") for mod in module_catalog if mod.get("code")), "1A"
        )
        module_number = (
            source_assessment.module_number
            if source_assessment and getattr(source_assessment, "module_number", None)
            else paper_meta.get("module_number") or default_module_code
        )
        module_name = (
            source_assessment.module_name
            if source_assessment and getattr(source_assessment, "module_name", None)
            else paper_meta.get("module_name")
            or next(
                (
                    mod.get("label")
                    for mod in module_catalog
                    if mod.get("code") == module_number
                ),
                original_paper.name,
            )
        )
        saqa_id = (
            source_assessment.saqa_id
            if source_assessment and source_assessment.saqa_id
            else ""
        )

        if not module_name:
            randomized_paper.delete()
            messages.error(
                request, "Module name is required before randomizing this paper."
            )
            return redirect("review_saved_selector")
        if not module_number:
            randomized_paper.delete()
            messages.error(
                request, "Module letter is required before randomizing this paper."
            )
            return redirect("review_saved_selector")

        try:
            pool_result = build_randomized_structure_from_pool(
                original_paper,
                randomized_paper,
                module_name,
                module_number,
                base_extractor=(
                    source_assessment.extractor_paper
                    if source_assessment and source_assessment.extractor_paper_id
                    else None
                ),
            )
        except RandomizationPoolError as exc:
            randomized_paper.delete()
            messages.error(request, str(exc))
            return redirect("review_saved_selector")
        total_marks = pool_result.get("total_marks", 0)
        pool_used = pool_result.get("used_pool", False)
        pool_size = pool_result.get("pool_size", 0)

        qual_name = (
            qualification.name
            if qualification
            else (
                original_paper.qualification.name
                if original_paper.qualification
                else None
            )
        )
        module_meta = get_module_meta(qual_name, module_number) if qual_name else {}
        rand_status = (
            randomization_status(qual_name, module_number)
            if qual_name
            else "unconfigured"
        )
        random_letters = allowed_letters(qual_name, module_number) if qual_name else []
        cover_heading = cover_title(qual_name, module_number) if qual_name else None

        meta_payload = randomized_paper.structure_json or {}
        meta_payload["randomization"] = {
            "module_number": module_number,
            "status": rand_status,
            "allowed_letters": random_letters,
            "base_paper_id": original_paper.id,
            "snapshot_pool_used": pool_used,
            "snapshot_pool_size": pool_size,
            "source": "snapshot_pool" if pool_used else "clone",
        }
        if qual_name:
            meta_payload["randomization"]["qualification"] = qual_name
        if cover_heading:
            meta_payload["randomization"]["cover_title"] = cover_heading
        if module_meta:
            meta_payload["randomization"]["meta"] = module_meta
        if module_name:
            meta_payload["randomization"]["module_name"] = module_name
        if source_assessment:
            meta_payload["randomization"]["base_assessment_id"] = source_assessment.id
            if source_assessment.extractor_paper_id:
                meta_payload["randomization"][
                    "base_extractor_id"
                ] = source_assessment.extractor_paper_id

        randomized_paper.total_marks = total_marks
        randomized_paper.structure_json = meta_payload
        randomized_paper.save(update_fields=["total_marks", "structure_json"])

        if request.user.is_authenticated:
            created_by = request.user
        elif source_assessment and source_assessment.created_by:
            created_by = source_assessment.created_by
        else:
            created_by = original_paper.created_by

        random_comment = f"Randomized from {original_paper.name}"
        if rand_status and rand_status != "unconfigured":
            random_comment += f" | status: {rand_status}"
        if pool_used:
            random_comment += " | source: snapshot pool"

        randomized_assessment = Assessment.objects.create(
            eisa_id=f"EISA-{uuid.uuid4().hex[:8].upper()}",
            qualification=qualification,
            paper=randomized_paper.name,
            module_number=module_number,
            module_name=module_name,
            saqa_id=saqa_id,
            comment=random_comment,
            created_by=created_by,
            paper_link=randomized_paper,
            status="draft",
            paper_type="randomized",
            forward_to_moderator=False,
        )

        if source_assessment and source_assessment.file:
            with source_assessment.file.open("rb") as src_file:
                randomized_assessment.file.save(
                    os.path.basename(source_assessment.file.name),
                    File(src_file),
                    save=False,
                )
        if source_assessment and source_assessment.memo:
            with source_assessment.memo.open("rb") as src_memo:
                randomized_assessment.memo.save(
                    os.path.basename(source_assessment.memo.name),
                    File(src_memo),
                    save=False,
                )
        randomized_assessment.save()

        extractor = None
        if randomized_assessment.file:
            sync_assessment_paper_bank(randomized_assessment, force=True)
            extractor = rebuild_extractor_from_bank(randomized_assessment, force=True)

        messages.success(
            request,
            f"Randomized paper created ({randomized_assessment.eisa_id}) and queued for assessor developer review.",
        )
        if extractor:
            return redirect("assessor_developer_paper", paper_id=extractor.id)
        return redirect("assessor_developer")

    except Exception as e:
        messages.error(request, f"Failed to randomize paper: {str(e)}")
        return redirect("review_saved_selector")


@login_required
def save_blocks_view(request, paper_id):
    """Save extracted blocks to database"""
    try:
        paper = get_object_or_404(Paper, id=paper_id)
        blocks = request.session.get("extracted_blocks")

        if not blocks:
            messages.error(request, "No blocks found to save")
            return JsonResponse({"status": "error", "message": "No blocks found"})

        success = save_nodes_to_db(blocks, paper)

        if success:
            messages.success(request, "Blocks saved successfully")
            return JsonResponse(
                {
                    "status": "success",
                    "redirect": reverse("view_paper", args=[paper_id]),
                }
            )
        else:
            messages.error(request, "Failed to save blocks")
            return JsonResponse({"status": "error", "message": "Save failed"})

    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return JsonResponse({"status": "error", "message": str(e)})


@login_required
def write_exam(request, assessment_id):
    """
    Render the write page for a single assessment.
    - Old pipeline: GeneratedQuestion
    - New pipeline: ExamNode-based paper
    """
    assessment = get_object_or_404(
        Assessment.objects.select_related("paper_link", "qualification"),
        id=assessment_id,
    )

    # Check if user already submitted this exam
    has_submitted = ExamSubmission.objects.filter(
        assessment=assessment, 
        student=request.user
    ).exists()
    
    if has_submitted:
        messages.warning(request, "You have already submitted this exam.")
        return redirect("student_dashboard")

    # Optional safety: only show assessments matching the learner's qualification
    if (
        request.user.qualification
        and assessment.qualification_id != request.user.qualification_id
    ):
        messages.error(
            request, "This assessment is not assigned to your qualification."
        )
        return redirect("student_dashboard")

    paper = assessment.paper_link  # may be None for old pipeline
    generated_qs = assessment.generated_questions.all().order_by(
        "id"
    )  # old pipeline support

    node_sequence: list[dict] = []
    question_nodes: list[dict] = []
    if paper:
        raw_node_tree, node_stats = _preview_structure_for_paper(paper)
        node_sequence = deepcopy(raw_node_tree or [])

        def _assign_ids(nodes, prefix=""):
            for index, node in enumerate(nodes, 1):
                identifier = (
                    node.get("template_id")
                    or node.get("id")
                    or node.get("uuid")
                    or node.get("pk")
                )
                if not identifier:
                    identifier = f"{prefix}{index}"
                identifier = str(identifier)
                node["template_id"] = identifier
                children = node.get("children") or []
                if children:
                    _assign_ids(children, f"{identifier}-")

        _assign_ids(node_sequence)

        def _collect_questions(nodes, bucket):
            for node in nodes:
                if (node.get("node_type") or "").lower() == "question":
                    bucket.append(node)
                children = node.get("children") or []
                if children:
                    _collect_questions(children, bucket)

        question_nodes = []
        _collect_questions(node_sequence, question_nodes)
    else:
        node_stats = {
            "total": 0,
            "questions": 0,
            "tables": 0,
            "images": 0,
            "instructions": 0,
        }
        node_sequence = []

    random_meta = {}
    if paper and isinstance(paper.structure_json, dict):
        random_meta = paper.structure_json.get("randomization", {}) or {}

    context = {
        "assessment": assessment,
        "paper": paper,
        "generated_qs": generated_qs,
        "node_sequence": node_sequence,
        "node_stats": node_stats,
        "attempt_number": 1,  # demo; bump if you track attempts
        "random_meta": random_meta,
        "question_nodes": question_nodes,
        "is_randomized": assessment.paper_type == "randomized",
        "assessment_id": assessment_id,  # Add this for JavaScript
    }
    return render(request, "core/student/write_exam.html", context)

@login_required
@require_http_methods(["POST"])
def submit_exam(request, assessment_id):
    """
    Accepts both regular and AJAX submissions
    - answer_<GeneratedQuestion.id>          (old pipeline)
    - answer_node_<ExamNode.uuid>           (new pipeline)

    Old: creates ExamAnswer rows (your existing flow).
    New: collects node answers; optionally save into StudentWrittenPaper if present.
    """
    assessment = get_object_or_404(Assessment, id=assessment_id)

    # Check if already submitted (prevent duplicate submissions)
    existing_attempts = ExamSubmission.objects.filter(
        assessment=assessment, student=request.user
    ).count()
    
    if existing_attempts > 0:
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'success': False,
                'message': 'You have already submitted this exam.'
            })
        messages.warning(request, "You have already submitted this exam.")
        return redirect("student_dashboard")

    paper = assessment.paper_link

    student_number = (getattr(request.user, "student_number", None) or "").strip()
    if not student_number or getattr(request.user, "role", "") != "learner":
        messages.error(
            request,
            "This submission route is reserved for learners. You're signed in as staff/superuser, so no attempt was recorded.",
        )
        return redirect("student_dashboard")

    attempt_number = existing_attempts + 1

    # --- OLD PIPELINE: GeneratedQuestion -> ExamAnswer ---
    saved_count_old = 0
    for key, value in request.POST.items():
        if not key.startswith("answer_"):
            continue
        # ignore the new pipeline prefix
        if key.startswith("answer_node_"):
            continue

        # key is "answer_<gq_id>"
        try:
            gq_id = int(key.split("_", 1)[1])
        except (IndexError, ValueError):
            continue

        gq = GeneratedQuestion.objects.filter(id=gq_id, assessment=assessment).first()
        if not gq:
            continue

        try:
            with transaction.atomic():
                ExamAnswer.objects.create(
                    user=request.user,
                    question=gq,
                    answer_text=(value or "").strip(),
                    attempt_number=attempt_number,
                )
                saved_count_old += 1
        except IntegrityError:
            # If unique_together blocks dupes, update instead
            ans = ExamAnswer.objects.get(
                user=request.user, question=gq, attempt_number=attempt_number
            )
            ans.answer_text = (value or "").strip()
            ans.save(update_fields=["answer_text"])

    # --- NEW PIPELINE: ExamNode -> collect answers ---
    node_answers: dict[str, str] = {}
    answers_data = []
    for key, value in request.POST.items():
        if not key.startswith("answer_node_"):
            continue
        node_uuid = key.replace("answer_node_", "", 1).strip()
        answer_text = (value or "").strip()
        node_answers[node_uuid] = answer_text

        if not paper:
            continue

        position = len(answers_data) + 1
        try:
            exam_node = ExamNode.objects.get(id=node_uuid, paper=paper)
            answers_data.append(
                {
                    "node": exam_node,
                    "answer": answer_text,
                    "marks": exam_node.marks,
                    "question_number": exam_node.number or f"Q{position}",
                }
            )
        except ExamNode.DoesNotExist:
            answers_data.append(
                {
                    "question_text": f"Question {node_uuid[:8]}...",
                    "answer": answer_text,
                    "marks": None,
                    "question_number": f"Q{position}",
                }
            )

    saved_new = len(node_answers)

    if paper and (saved_new or saved_count_old):
        context = {
            "paper": paper,
            "assessment": assessment,
            "attempt_number": attempt_number,
            "student": request.user,
            "student_number": request.user.student_number,
            "student_name": f"{request.user.first_name} {request.user.last_name}",
            "answers": answers_data,
            "submission_date": timezone.now(),
            "total_questions": len(node_answers),
        }

        exam_submission = ExamSubmission.objects.create(
            student=request.user,
            student_number=request.user.student_number,
            student_name=f"{request.user.first_name} {request.user.last_name}",
            paper=paper,
            assessment=assessment,
            attempt_number=attempt_number,
            submitted_at=timezone.now(),
        )

        try:
            html_string = render_to_string(
                "core/student/exam_pdf_template.html", context
            )
            pdf_buffer = BytesIO()
            pisa_status = pisa.CreatePDF(html_string, dest=pdf_buffer)

            if pisa_status.err:
                raise Exception("PDF generation failed")

            file_name = f"{request.user.student_number}_{paper.name.replace(' ', '_')}_Attempt_{attempt_number}.pdf"
            pdf_buffer.seek(0)
            exam_submission.pdf_file.save(file_name, ContentFile(pdf_buffer.getvalue()))
            exam_submission.save()
        except Exception as exc:  # pragma: no cover
            messages.warning(
                request,
                "Your answers were saved, but the PDF could not be generated. Please contact support.",
            )

    # OPTIONAL: persist node answers if you've added StudentWrittenPaper
    # if node_answers:
    #     if paper:
    #         swp, _ = StudentWrittenPaper.objects.get_or_create(
    #             learner=request.user,
    #             paper=paper,
    #         )
    #         existing = swp.answers_json or {}
    #         existing.update(node_answers)
    #         swp.answers_json = existing
    #         swp.save(update_fields=['answers_json', 'last_updated'])

    total_saved = saved_count_old + saved_new
    
    # Handle AJAX requests (for auto-save)
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({
            'success': True,
            'message': 'Auto-saved successfully',
            'saved_count': total_saved
        })
    
    # Regular form submission
    if total_saved:
        messages.success(request, "Your answers were submitted successfully.")
    else:
        messages.warning(
            request,
            "We could not detect any answers. Please ensure your responses are filled in and submit again.",
        )
    return redirect("student_dashboard")


def papers_demo(request):
    """
    Simple demo page: list ALL papers from the DB in a table.
    No status/qualification filtering (for now).
    """
    papers = Paper.objects.select_related("qualification").order_by("-created_at")
    return render(
        request,
        "core/student/papers_demo.html",
        {
            "papers": papers,
            "user": request.user,
        },
    )


# @login_required
# def write_paper_simple(request, paper_id):
#     """View for students to write/attempt a paper"""
#     paper = get_object_or_404(Paper.objects.select_related('qualification'), id=paper_id)

#     # Check if the paper belongs to an assessment that's released to students
#     assessment = get_object_or_404(
#         Assessment.objects.filter(
#             status="Released to students",
#             qualification=request.user.qualification,
#             paper_link=paper
#         )
#     )

#     # Build the questions tree for the paper
#     questions = build_questions_tree_for_paper(paper)

#     # Calculate attempt number (current count + 1)
#     attempt_count = ExamAnswer.objects.filter(
#         question__assessment=assessment,
#         user=request.user
#     ).values('attempt_number').distinct().count()

#     attempt_number = attempt_count + 1


#     return render(request, 'core/student/write_paper_simple.html', {
#         'paper': paper,
#         'assessment': assessment,  # Pass assessment context as well
#         'questions': questions,
#         'attempt_number': attempt_number,
#     })
@login_required
def write_paper_simple(request, paper_id):
    """View for students to write/attempt a paper"""

    # Check if student has completed registration (has student number)
    if (
        not request.user.student_number
        or not request.user.first_name
        or not request.user.last_name
    ):
        # Redirect to registration page
        return redirect("student_registration", paper_id=paper_id)

    paper = get_object_or_404(
        Paper.objects.select_related("qualification"), id=paper_id
    )

    # Check if the paper belongs to an assessment that's released to students
    assessment = get_object_or_404(
        Assessment.objects.filter(
            status="Released to students",
            qualification=request.user.qualification,
            paper_link=paper,
        )
    )

    # Build the questions tree for the paper
    questions = build_questions_tree_for_paper(paper)

    # Calculate attempt number
    attempt_count = ExamSubmission.objects.filter(
        assessment=assessment, student=request.user
    ).count()

    if attempt_count > 0:
        messages.info(request, "You have already submitted this exam.")
        return redirect("student_dashboard")

    attempt_number = attempt_count + 1

    return render(
        request,
        "core/student/write_paper_simple.html",
        {
            "paper": paper,
            "assessment": assessment,
            "questions": questions,
            "attempt_number": attempt_number,
        },
    )


@login_required
@require_http_methods(["POST"])
def submit_paper_simple(request, paper_id):
    """
    Accepts answer_node_<ExamNode.uuid> fields, generates PDF using xhtml2pdf, and saves submission.
    """
    paper = get_object_or_404(Paper, id=paper_id)
    assessment_id = request.POST.get("assessment_id")

    # Get the assessment
    assessment = get_object_or_404(Assessment, id=assessment_id, paper_link=paper)

    existing_attempts = ExamSubmission.objects.filter(
        assessment=assessment, student=request.user
    ).count()
    attempt_number = existing_attempts + 1

    # Gather answers and build context for PDF
    node_answers = {}
    answers_data = []

    for key, val in request.POST.items():
        if key.startswith("answer_node_"):
            node_uuid = key.replace("answer_node_", "", 1)
            answer_text = (val or "").strip()
            node_answers[node_uuid] = answer_text

            # Try to get the exam node for context
            try:
                exam_node = ExamNode.objects.get(id=node_uuid, paper=paper)
                answers_data.append(
                    {
                        "node": exam_node,
                        "answer": answer_text,
                        "marks": exam_node.marks,
                        "question_number": exam_node.number
                        or f"Q{len(answers_data) + 1}",
                    }
                )
            except ExamNode.DoesNotExist:
                answers_data.append(
                    {
                        "question_text": f"Question {node_uuid[:8]}...",
                        "answer": answer_text,
                        "marks": None,
                        "question_number": f"Q{len(answers_data) + 1}",
                    }
                )

    # Prepare context for PDF generation
    context = {
        "paper": paper,
        "assessment": assessment,
        "attempt_number": attempt_number,
        "student": request.user,
        "student_number": request.user.student_number,
        "student_name": f"{request.user.first_name} {request.user.last_name}",
        "answers": answers_data,
        "submission_date": timezone.now(),
        "total_questions": len(node_answers),
    }

    # Create the submission record before attempting PDF generation
    exam_submission = ExamSubmission.objects.create(
        student=request.user,
        student_number=request.user.student_number,
        student_name=f"{request.user.first_name} {request.user.last_name}",
        paper=paper,
        assessment=assessment,
        attempt_number=attempt_number,
        submitted_at=timezone.now(),
    )
    # Generate PDF with xhtml2pdf
    try:
        html_string = render_to_string("core/student/exam_pdf_template.html", context)

        # Create PDF in memory
        pdf_buffer = BytesIO()
        pisa_status = pisa.CreatePDF(html_string, dest=pdf_buffer)

        if pisa_status.err:
            raise Exception("PDF generation failed")

        # Save PDF to ExamSubmission model
        file_name = f"{request.user.student_number}_{paper.name.replace(' ', '_')}_Attempt_{attempt_number}.pdf"

        # Save the PDF file
        pdf_buffer.seek(0)
        exam_submission.pdf_file.save(file_name, ContentFile(pdf_buffer.getvalue()))
        exam_submission.save()

    except Exception as e:
        messages.warning(
            request,
            "Your answers were saved, but the PDF could not be generated. Please contact support.",
        )
        # Continue with submission even if PDF fails

    # Save individual answers to ExamAnswer model
    # for node_uuid, answer_text in node_answers.items():
    #     try:
    #         exam_node = ExamNode.objects.get(id=node_uuid, paper=paper)
    #         ExamNode.objects.create(
    #             user=request.user,
    #             question=exam_node,
    #             assessment=assessment,
    #             attempt_number=attempt_number,
    #             answer_text=answer_text,
    #             submission_time=timezone.now()
    #         )
    #     except ExamNode.DoesNotExist:
    #         continue

    messages.success(
        request,
        f"Successfully submitted {len(node_answers)} answers for '{paper.name}'. Attempt #{attempt_number} completed.",
    )
    return redirect("student_dashboard")


# end of demo views
####New views from billie's code#################################


def assessment_center_view(request):
    qualification = request.user.qualification
    
    # Filter batches by user's qualification if they have one
    if qualification:
        batches = Batch.objects.filter(
            submitted_to_center=True,
            qualification=qualification
        )
    else:
        # If no qualification assigned, show all batches (or none, depending on your needs)
        batches = Batch.objects.filter(submitted_to_center=True)
    
    return render(
        request,
        "core/assessment-center/assessment_center.html",
        {
            "batches": batches,
            "qualification": qualification,
            "user": request.user,
        },
    )


# @login_required
# def assessor_maker_dashboard(request):
#     return render(request, "core/assessor/marker_analytics.html")

# views.py in Marker app


@login_required
def assessor_maker_dashboard(request):
    """Dashboard view for markers to see student registry and exam submissions"""

    # Check if user is a marker
    if request.user.role not in ["assessor_marker", "internal_mod", "external_mod"]:
        return redirect("access_denied")

    # Get all registered students (learners with student numbers)
    students = (
        CustomUser.objects.filter(role="learner", student_number__isnull=False)
        .select_related("qualification")
        .order_by("-created_at")
    )

    # Get all exam submissions with related assessment and memo data
    submissions = ExamSubmission.objects.select_related(
        "paper", "assessment", "student", "offline_student", "graded_by"
    ).prefetch_related(
        'assessment__paper_link'
    ).order_by("-submitted_at")

    # Get offline students
    offline_students = OfflineStudent.objects.select_related("qualification").order_by(
        "-created_at"
    )

    # Get offline submissions only
    offline_submissions = (
        ExamSubmission.objects.filter(is_offline=True)
        .select_related("paper", "assessment", "offline_student", "graded_by")
        .prefetch_related('assessment__paper_link')
        .order_by("-submitted_at")
    )

    # Get papers and assessments for filters/forms
    papers = Paper.objects.all()
    assessments = Assessment.objects.all()

    return render(
        request,
        "core/assessor/mark_exam.html",
        {
            "students": students,
            "offline_students": offline_students,
            "submissions": submissions,
            "offline_submissions": offline_submissions,
            "papers": papers,
            "assessments": assessments,
            "user": request.user,
        },
    )

@login_required
def student_registration(request, paper_id):
    """Registration form before taking exam"""
    paper = get_object_or_404(Paper, id=paper_id)

    if request.method == "POST":
        form = StudentRegistrationForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save()
            # Redirect to the exam
            return redirect("write_paper_simple", paper_id=paper_id)
    else:
        form = StudentRegistrationForm(instance=request.user)

    return render(
        request, "core/student/registration.html", {"form": form, "paper": paper}
    )


# views.py


@login_required
def upload_student_register(request):
    """View to handle CSV upload for offline students"""

    # Get qualifications to show in the template
    qualifications = Qualification.objects.all()

    if request.method == "POST":
        register_file = request.FILES.get("register_file")

        if not register_file:
            messages.error(request, "Please select a CSV file to upload.")
            return render(
                request,
                "core/assessment-center/upload_register.html",
                {"qualifications": qualifications},
            )

        # Check if the uploaded file is a CSV
        if not register_file.name.endswith(".csv"):
            messages.error(request, "Please upload a CSV file.")
            return render(
                request,
                "core/assessment-center/upload_register.html",
                {"qualifications": qualifications},
            )

        try:
            # Read and decode file, handling UTF-8 BOM
            file_content = register_file.read().decode(
                "utf-8-sig"
            )  # Use utf-8-sig to remove BOM
            decoded_file = file_content.splitlines()

            # Detect delimiter automatically (comma or semicolon)
            sample = "\n".join(decoded_file[:5])
            dialect = csv.Sniffer().sniff(sample, delimiters=[",", ";"])
            reader = csv.DictReader(decoded_file, dialect=dialect)

            # Clean field names by stripping whitespace, lowercase, and removing BOM
            reader.fieldnames = [
                h.strip().lower().replace("\ufeff", "") for h in reader.fieldnames
            ]

            students_created = 0
            errors = []

            for row_num, row in enumerate(
                reader, start=2
            ):  # Start at 2 to account for header
                try:
                    # Validate required fields
                    required_fields = [
                        "student_number",
                        "first_name",
                        "last_name",
                        "email",
                        "qualification",
                        "status",
                    ]
                    for field in required_fields:
                        if field not in row or not row[field]:
                            errors.append(
                                f"Row {row_num}: Missing required field '{field}'"
                            )
                            continue

                    # Validate status
                    if row["status"].lower() not in ["present", "absent"]:
                        errors.append(
                            f"Row {row_num}: Status must be 'present' or 'absent'"
                        )
                        continue

                    # Check if student already exists
                    if OfflineStudent.objects.filter(
                        student_number=row["student_number"]
                    ).exists():
                        errors.append(
                            f"Row {row_num}: Student {row['student_number']} already exists"
                        )
                        continue

                    # Get or create qualification
                    qualification_name = row["qualification"].strip()
                    try:
                        qualification = Qualification.objects.get(
                            name=qualification_name
                        )
                    except Qualification.DoesNotExist:
                        errors.append(
                            f"Row {row_num}: Qualification '{qualification_name}' does not exist in the system"
                        )
                        continue

                    # Create new offline student
                    student = OfflineStudent(
                        student_number=row["student_number"],
                        first_name=row["first_name"],
                        last_name=row["last_name"],
                        email=row["email"],
                        qualification=qualification,
                        status=row["status"].lower(),
                        created_by=request.user,
                    )
                    student.save()
                    students_created += 1

                except Exception as e:
                    errors.append(f"Row {row_num}: Error - {str(e)}")

            # Show success message with stats
            if students_created > 0:
                messages.success(
                    request,
                    f"Successfully imported {students_created} offline students.",
                )

            # Show errors if any
            if errors:
                error_msg = f"Completed with {len(errors)} error(s). First few errors: {', '.join(errors[:3])}"
                if len(errors) > 3:
                    error_msg += f"... and {len(errors) - 3} more"
                messages.warning(request, error_msg)

        except Exception as e:
            messages.error(request, f"Error processing CSV file: {str(e)}")

        return redirect("upload_student_register")

    return render(
        request,
        "core/assessment-center/upload_register.html",
        {"qualifications": qualifications},
    )


@login_required
def download_register_template(request):
    """View to download CSV template for offline student registration"""
    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = (
        'attachment; filename="offline_students_template.csv"'
    )

    writer = csv.writer(response)

    # Header row with required fields
    writer.writerow(
        [
            "student_number",
            "first_name",
            "last_name",
            "email",
            "qualification",
            "status",
        ]
    )

    # Get existing qualifications for examples
    qualifications = Qualification.objects.all()[
        :3
    ]  # Get first 3 qualifications for examples
    qual_names = (
        [q.name for q in qualifications]
        if qualifications
        else ["BSc Computer Science", "BSc Information Technology", "BSc Data Science"]
    )

    # Example rows
    writer.writerow(
        [
            "STD1001",
            "John",
            "Smith",
            "john.smith@example.com",
            qual_names[0] if qual_names else "BSc Computer Science",
            "present",
        ]
    )
    writer.writerow(
        [
            "STD1002",
            "Jane",
            "Doe",
            "jane.doe@example.com",
            qual_names[1] if len(qual_names) > 1 else "BSc Information Technology",
            "absent",
        ]
    )

    # Add note about valid status values
    response.write(
        '\n\n# Note: Status must be either "present" or "absent" (case insensitive)'
    )
    response.write("\n# Qualification must match existing qualifications in the system")

    return response


# views.py in Assessment Center app
# views.py - Simplified view
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from datetime import datetime


@login_required
def upload_offline_exam_submission(request):
    """View for Assessment Center to upload offline exam submissions"""

    # Check if user has permission (assessment center role)
    if request.user.role not in ["assessment_center", "admin"]:
        return redirect("access_denied")

    # Get only present offline students
    offline_students = (
        OfflineStudent.objects.filter(status="present")
        .select_related("qualification")
        .order_by("student_number")
    )

    # Get all uploaded exam submissions (offline only) with related data
    exam_submissions = (
        ExamSubmission.objects.filter(is_offline=True)
        .select_related("offline_student", "offline_student__qualification")
        .prefetch_related("graded_by", "internal_graded_by", "external_graded_by")
        .order_by("-submitted_at")
    )

    if request.method == "POST":
        try:
            student_id = request.POST.get("offline_student")
            attempt_number = request.POST.get("attempt_number", 1)
            pdf_file = request.FILES.get("pdf_file")

            # Validate required fields
            if not all([student_id, pdf_file]):
                messages.error(request, "Student selection and PDF file are required.")
                return render(
                    request,
                    "core/assessment-center/upload_offline_exam.html",
                    {
                        "offline_students": offline_students,
                        "exam_submissions": exam_submissions,
                    },
                )

            # Validate file type
            if not pdf_file.name.lower().endswith(".pdf"):
                messages.error(request, "Only PDF files are allowed.")
                return render(
                    request,
                    "core/assessment-center/upload_offline_exam.html",
                    {
                        "offline_students": offline_students,
                        "exam_submissions": exam_submissions,
                    },
                )

            # Validate file size (10MB limit)
            if pdf_file.size > 10 * 1024 * 1024:  # 10MB in bytes
                messages.error(request, "File size must be less than 10MB.")
                return render(
                    request,
                    "core/assessment-center/upload_offline_exam.html",
                    {
                        "offline_students": offline_students,
                        "exam_submissions": exam_submissions,
                    },
                )

            # Get student object
            offline_student = OfflineStudent.objects.get(id=student_id)

            # Check if this attempt already exists for this student
            existing_attempt = ExamSubmission.objects.filter(
                offline_student=offline_student, attempt_number=attempt_number
            ).exists()

            if existing_attempt:
                messages.warning(
                    request,
                    f"Attempt {attempt_number} already exists for {offline_student.student_number}.",
                )
                return render(
                    request,
                    "core/assessment-center/upload_offline_exam.html",
                    {
                        "offline_students": offline_students,
                        "exam_submissions": exam_submissions,
                    },
                )

            # Create submission record - let the model's save method handle the student details
            submission = ExamSubmission(
                offline_student=offline_student,
                attempt_number=attempt_number,
                # paper and assessment can be set to None for offline submissions
                # or you can add logic to determine the appropriate paper/assessment
                paper=None,
                assessment=None,
                is_offline=True,
            )

            # Save the file - the model will auto-set student_number and student_name in save()
            file_name = f"offline_{offline_student.student_number}_attempt_{attempt_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            submission.pdf_file.save(file_name, pdf_file)

            messages.success(
                request,
                f"Exam submission for {offline_student.student_number} (Attempt {attempt_number}) uploaded successfully!",
            )
            return redirect("upload_offline_exam_submission")

        except OfflineStudent.DoesNotExist:
            messages.error(request, "Selected student not found.")
        except Exception as e:
            messages.error(request, f"Error uploading file: {str(e)}")

    return render(
        request,
        "core/assessment-center/upload_offline_exam.html",
        {
            "offline_students": offline_students,
            "exam_submissions": exam_submissions,
        },
    )


@login_required
@require_POST
def quick_grade_submission(request, submission_id):
    """AJAX view for quick grading"""

    # Check if user is a marker
    if request.user.role not in ["assessor_marker", "internal_mod", "external_mod"]:
        return JsonResponse({"success": False, "message": "Permission denied"})

    submission = get_object_or_404(ExamSubmission, id=submission_id)

    try:
        marks = request.POST.get("marks")
        total_marks = request.POST.get("total_marks", 100)
        feedback = request.POST.get("feedback", "")

        # Validate marks
        if not marks:
            return JsonResponse({"success": False, "message": "Marks are required"})

        # Update submission with grades
        submission.marks = Decimal(marks)
        submission.total_marks = Decimal(total_marks)
        submission.feedback = feedback
        submission.graded_by = request.user
        submission.graded_at = timezone.now()
        submission.save()

        return JsonResponse(
            {
                "success": True,
                "message": "Graded successfully",
                "marks": float(submission.marks),
                "total_marks": float(submission.total_marks),
                "status": submission.status,
                "graded_by": (
                    submission.graded_by.get_full_name()
                    if submission.graded_by
                    else "Unknown"
                ),
                "graded_at": (
                    submission.graded_at.strftime("%b %d, %Y")
                    if submission.graded_at
                    else ""
                ),
            }
        )

    except (ValueError, InvalidOperation):
        return JsonResponse(
            {"success": False, "message": "Please enter valid numbers for marks"}
        )
    except Exception as e:
        return JsonResponse({"success": False, "message": f"Error: {str(e)}"})

        # ----------------- DASHBOARD -----------------


@login_required
def internal_moderator_dashboard(request):
    if request.user.role != "internal_mod":
        return redirect("access_denied")

    # Show submissions that have been graded by marker but not internally moderated
    pending_submissions = (
        ExamSubmission.objects.filter(
            marks__isnull=False,  # Graded by marker
            internal_marks__isnull=True,  # Not yet internally moderated
        )
        .select_related(
            "student",
            "offline_student",
            "paper",
            "assessment",
            "graded_by",
            "internal_graded_by",
        )
        .order_by("-submitted_at")
    )

    # Show submissions that have been internally moderated
    moderated_submissions = (
        ExamSubmission.objects.filter(
            marks__isnull=False,  # Graded by marker
            internal_marks__isnull=False,  # Already internally moderated
        )
        .select_related(
            "student",
            "offline_student",
            "paper",
            "assessment",
            "graded_by",
            "internal_graded_by",
        )
        .order_by("-internal_graded_at")
    )

    papers = Paper.objects.all()
    assessments = Assessment.objects.all()

    return render(
        request,
        "core/moderator/internal_moderator_marker.html",
        {
            "pending_submissions": pending_submissions,
            "moderated_submissions": moderated_submissions,
            "pending_count": pending_submissions.count(),
            "moderated_count": moderated_submissions.count(),
            "papers": papers,
            "assessments": assessments,
            "user": request.user,
        },
    )


@login_required
@require_POST
def internal_grade_submission(request, submission_id):
    if request.user.role != "internal_mod":
        return JsonResponse({"success": False, "message": "Permission denied"})

    submission = get_object_or_404(ExamSubmission, id=submission_id)

    try:
        marks = request.POST.get("marks")
        total_marks = request.POST.get("total_marks", 100)
        feedback = request.POST.get("feedback", "")
        marked_paper = request.FILES.get("internal_marked_paper")

        if not marks:
            return JsonResponse({"success": False, "message": "Marks are required"})

        submission.internal_marks = Decimal(marks)
        submission.internal_total_marks = Decimal(total_marks)
        submission.internal_feedback = feedback
        submission.internal_graded_by = request.user
        submission.internal_graded_at = timezone.now()

        # Handle marked paper upload
        if marked_paper:
            # Validate file type
            if not marked_paper.name.lower().endswith(".pdf"):
                return JsonResponse(
                    {"success": False, "message": "Only PDF files are allowed"}
                )
            submission.internal_marked_paper = marked_paper

        submission.save()

        return JsonResponse(
            {
                "success": True,
                "message": "Internal grade saved successfully",
                "marks": float(submission.internal_marks),
                "total_marks": float(submission.internal_total_marks),
                "status": submission.status,
                "graded_by": (
                    submission.internal_graded_by.get_full_name()
                    if submission.internal_graded_by
                    else "Unknown"
                ),
                "graded_at": (
                    submission.internal_graded_at.strftime("%b %d, %Y")
                    if submission.internal_graded_at
                    else ""
                ),
            }
        )

    except (ValueError, InvalidOperation):
        return JsonResponse(
            {"success": False, "message": "Please enter valid numbers for marks"}
        )
    except Exception as e:
        return JsonResponse({"success": False, "message": f"Error: {str(e)}"})


@login_required
def external_moderator_dashboard(request):
    if request.user.role != "external_mod":
        return redirect("access_denied")

    # Show submissions that have been internally moderated but not externally moderated
    pending_submissions = (
        ExamSubmission.objects.filter(
            internal_marks__isnull=False,  # Internally moderated
            external_marks__isnull=True,  # Not yet externally moderated
        )
        .select_related(
            "student",
            "offline_student",
            "paper",
            "assessment",
            "graded_by",
            "internal_graded_by",
            "external_graded_by",
        )
        .order_by("-submitted_at")
    )

    # Show submissions that have been externally moderated
    moderated_submissions = (
        ExamSubmission.objects.filter(
            internal_marks__isnull=False,  # Internally moderated
            external_marks__isnull=False,  # Already externally moderated
        )
        .select_related(
            "student",
            "offline_student",
            "paper",
            "assessment",
            "graded_by",
            "internal_graded_by",
            "external_graded_by",
        )
        .order_by("-external_graded_at")
    )

    papers = Paper.objects.all()
    assessments = Assessment.objects.all()

    return render(
        request,
        "core/moderator/external_moderator_marker.html",
        {
            "pending_submissions": pending_submissions,
            "moderated_submissions": moderated_submissions,
            "pending_count": pending_submissions.count(),
            "moderated_count": moderated_submissions.count(),
            "papers": papers,
            "assessments": assessments,
            "user": request.user,
        },
    )


@login_required
@require_POST
def external_grade_submission(request, submission_id):
    if request.user.role != "external_mod":
        return JsonResponse({"success": False, "message": "Permission denied"})

    submission = get_object_or_404(ExamSubmission, id=submission_id)

    try:
        marks = request.POST.get("marks")
        total_marks = request.POST.get("total_marks", 100)
        feedback = request.POST.get("feedback", "")
        marked_paper = request.FILES.get("external_marked_paper")

        if not marks:
            return JsonResponse({"success": False, "message": "Marks are required"})

        submission.external_marks = Decimal(marks)
        submission.external_total_marks = Decimal(total_marks)
        submission.external_feedback = feedback
        submission.external_graded_by = request.user
        submission.external_graded_at = timezone.now()

        # Handle marked paper upload
        if marked_paper:
            # Validate file type
            if not marked_paper.name.lower().endswith(".pdf"):
                return JsonResponse(
                    {"success": False, "message": "Only PDF files are allowed"}
                )
            submission.external_marked_paper = marked_paper

        submission.save()

        return JsonResponse(
            {
                "success": True,
                "message": "External moderation saved successfully",
                "marks": float(submission.external_marks),
                "total_marks": float(submission.external_total_marks),
                "status": submission.status,  # This will now return "Finalized" automatically
                "graded_by": (
                    submission.external_graded_by.get_full_name()
                    if submission.external_graded_by
                    else "Unknown"
                ),
                "graded_at": (
                    submission.external_graded_at.strftime("%b %d, %Y")
                    if submission.external_graded_at
                    else ""
                ),
            }
        )
    except (ValueError, InvalidOperation):
        return JsonResponse(
            {"success": False, "message": "Please enter valid numbers for marks"}
        )
    except Exception as e:
        return JsonResponse({"success": False, "message": f"Error: {str(e)}"})


from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .models import ExamSubmission, Assessment


@login_required
def student_results(request):
    # Get the current student
    student = request.user

    # Get all submissions for this student, ordered by most recent
    submissions = (
        ExamSubmission.objects.filter(student=student)
        .select_related("paper", "assessment")
        .order_by("-submitted_at")
    )

    # Calculate final marks and status for each submission
    results = []
    for submission in submissions:
        # Determine the final marks based on grading progression
        if submission.external_marks is not None:
            final_marks = submission.external_marks
            final_total = submission.external_total_marks
            status = "Finalized"
        elif submission.internal_marks is not None:
            final_marks = submission.internal_marks
            final_total = submission.internal_total_marks
            status = "Reviewed"
        elif submission.marks is not None:
            final_marks = submission.marks
            final_total = submission.total_marks
            status = "Graded by Marker"
        else:
            final_marks = None
            final_total = submission.total_marks
            status = "Pending"

        # Calculate percentage if marks are available
        percentage = None
        if final_marks is not None and final_total:
            percentage = (final_marks / final_total) * 100

        results.append(
            {
                "submission": submission,
                "final_marks": final_marks,
                "final_total": final_total,
                "percentage": percentage,
                "status": status,
                "passed": (
                    percentage >= 50 if percentage else False
                ),  # Assuming 50% pass mark
            }
        )

    context = {
        "results": results,
        "student": student,
    }
    return render(request, "core/student/results.html", context)


# New view############
@login_required
def upload_marked_paper(request, submission_id):
    """Handle marked paper upload for markers"""
    if request.user.role not in ["assessor_marker", "internal_mod", "external_mod"]:
        return JsonResponse({"success": False, "message": "Access denied"})

    try:
        submission = ExamSubmission.objects.get(id=submission_id)

        if request.method == "POST":
            # Check if user has permission to upload marked paper
            if request.user.role == "assessor_marker":
                # For markers, they can upload marked paper when grading
                marked_paper = request.FILES.get("marked_paper")
                marks = request.POST.get("marks")
                total_marks = request.POST.get("total_marks", 100)
                feedback = request.POST.get("feedback", "")

                if marked_paper:
                    # Validate file type
                    if not marked_paper.name.lower().endswith(".pdf"):
                        return JsonResponse(
                            {"success": False, "message": "Only PDF files are allowed"}
                        )

                    submission.marked_paper = marked_paper

                if marks:
                    submission.marks = marks
                    submission.total_marks = total_marks
                    submission.feedback = feedback
                    submission.graded_by = request.user
                    submission.graded_at = timezone.now()

                submission.save()
                return JsonResponse(
                    {"success": True, "message": "Marked paper uploaded successfully"}
                )

            # We'll add internal and external moderator logic later
            return JsonResponse(
                {"success": False, "message": "Invalid user role for this action"}
            )

    except ExamSubmission.DoesNotExist:
        return JsonResponse({"success": False, "message": "Submission not found"})
    except Exception as e:
        return JsonResponse({"success": False, "message": str(e)})


# login waiting activation view
def waiting_activation(request):
    return render(request, "core/login/awaiting_activation.html")


from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.utils import timezone
# views.py - With URL parameter
@login_required
def toggle_student_status(request, student_id):
    if request.method == 'POST':
        try:
            # Get the student
            student = get_object_or_404(CustomUser, id=student_id)
            
            # Verify permissions (optional but recommended)
            if (hasattr(request.user, 'assessment_centre') and 
                student.assessment_centre != request.user.assessment_centre):
                messages.error(request, 'You do not have permission to modify this student.')
                return redirect('create_student_by_assessment_center')
            
            # Toggle the active status
            if student.is_active:
                student.is_active = False
                student.deactivated_at = timezone.now()
                message = f'Student {student.get_full_name()} has been deactivated.'
            else:
                student.is_active = True
                student.deactivated_at = None
                message = f'Student {student.get_full_name()} has been activated.'
            
            student.save()
            messages.success(request, message)
            
        except CustomUser.DoesNotExist:
            messages.error(request, 'Student not found.')
        except Exception as e:
            messages.error(request, f'Error updating student: {str(e)}')
    
    return redirect('create_student_by_assessment_center')

######new views####################################################################################
