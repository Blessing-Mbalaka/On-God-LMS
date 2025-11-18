#!/usr/bin/env python3
"""
robustexamextraction.py — ONE FILE to rule exam paper extraction

Design
------------
• Parse any .docx at the XML layer to preserve the exact order of content.
• Extract paragraphs, tables, images/figures, page/section breaks, equations placeholders.
• Detect questions of many styles (1, 1.1, 1.1.1, Q1, A1, 1), tail marks like "(10 Marks)".
• Attach trailing tables/figures/paragraphs to the active question until the next peer header.
• Parent/child relationship via dotted numbering (1.1 → parent 1), with robust backfill when numbers are missing or skipped.
• Gemini 2.5 Flash  for validation/backfill; Gemma 3  local fallback.
• Randomization‑ready output (structure_json) + console pretty print of tables & storage.
• Conservative ETA and verbose logging; tolerant of malformed docs.

Usage
-----
python robustexamextraction.py input.docx -o out_dir --gemini --gemma

Outputs in out_dir/
• structure_json.json  — canonical structure for DB
• media/               — extracted images (filenames referenced by nodes)
• preview.html         — quick human QA

Author: Blessing Mbalaka(2025‑08‑08)
"""
from __future__ import annotations

import tempfile 

import base64, mimetypes
from pathlib import Path
import os, io, re, json, time, math, uuid, zipfile, hashlib, argparse, traceback
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Optional, Tuple, Any
from xml.etree import ElementTree as ET




# -----------------------------
# Optional Django integration (safe to import even if Django not configured)
# -----------------------------
try:
    import django
    from django.conf import settings
    if not settings.configured:
        # Configure minimal Django settings for standalone use
        settings.configure(
            INSTALLED_APPS=[],
            DATABASES={},
            USE_TZ=True,
        )
    django.setup()
    
    # Now safe to import Django models
    from core.models import Paper, ExamNode
    from django.db import transaction
    DJANGO_AVAILABLE = True
except Exception:
    DJANGO_AVAILABLE = False
    # Create mock classes for standalone use
    class Paper:
        pass
    class ExamNode:
        pass
    def transaction():
        class atomic:
            def __enter__(self): return self
            def __exit__(self, *args): pass
        return atomic()

# -----------------------------
# Optional LLM clients (safe to import even if missing env)
# -----------------------------
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except Exception:
    GEMINI_AVAILABLE = False

import requests  # used for Gemma (Ollama) fallback — 
import filetype
import os

def guess_ext(b: bytes, path_hint: str = '') -> str:
    # Primary detection using filetype
    kind = filetype.guess(b)
    if kind:
        print(f"🔍 Detected via filetype: {kind.mime}")
        return f".{kind.extension}"

    # Heuristic byte signatures for fallback
    signatures = [
        (b'\x89PNG\r\n\x1a\n', '.png'),
        (b'\xff\xd8\xff', '.jpg'),
        (b'GIF87a', '.gif'),
        (b'GIF89a', '.gif'),
        (b'BM', '.bmp'),
        (b'II*\x00', '.tif'),
        (b'MM\x00*', '.tif'),
        (b'\x00\x00\x01\x00', '.ico'),
        (b'\x00\x00\x02\x00', '.cur'),
        (b'\xd7\xcd\xc6\x9a', '.wmf'),
        (b'\x01\x00\x09\x00', '.emf'),
        (b'\x50\x4B\x03\x04', '.docx'),  # ZIP: could be .docx/.xlsx/.pptx
        (b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1', '.doc'),  # Word 97–2003 binary
        (b'%PDF-', '.pdf'),
        (b'\x25\x21', '.eps'),
        (b'\x52\x49\x46\x46', '.webp'),  # RIFF header (check needed)
        (b'\x38\x42\x50\x53', '.psd'),   # Photoshop
        (b'<?xml', '.xml'),
    ]

    for sig, ext in signatures:
        if b.startswith(sig):
            print(f"🧠 Matched signature for {ext} → {sig[:4].hex()}")
            return ext

    # Fallback to extension from filename hint
    ext = os.path.splitext(path_hint)[1]
    if ext:
        print(f"⚠️ Fallback to path extension: {ext}")
        return ext

    # Final fallback
    print("⚠️ Unknown file type — defaulting to .bin")
    return '.bin'


# -----------------------------
# Logging / console helpers (no third‑party deps)
# -----------------------------
RESET = "\033[0m"; BOLD = "\033[1m"; DIM = "\033[2m"
GREEN = "\033[92m"; YELLOW = "\033[93m"; RED = "\033[91m"; CYAN = "\033[96m"

def info(msg):  print(f"{CYAN}ℹ{RESET} {msg}")

def ok(msg):    print(f"{GREEN}✔{RESET} {msg}")

def warn(msg):  print(f"{YELLOW}⚠{RESET} {msg}")

def err(msg):   print(f"{RED}✖{RESET} {msg}")

# -----------------------------
# Namespaces & regexes
# -----------------------------
NS = {
    'w':  'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r':  'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a':  'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/wordprocessingml/2006/wordprocessingDrawing',
    'pic':'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'pr': 'http://schemas.openxmlformats.org/package/2006/relationships',   
    'v':  'urn:schemas-microsoft-com:vml',                                   
}


# Add missing regex patterns that are referenced but not defined
PAGE_TRASH_RE = re.compile(r'^(page\s+\d+|©|\s*$)', re.IGNORECASE)
SECTION_HEADING_RE = re.compile(r'^(section|part|chapter)\s+[a-z0-9]+\s*$', re.IGNORECASE)
MARKS_RE = re.compile(r'\[(\d+)\]|\((\d+)\s*marks?\)', re.IGNORECASE)

# ENHANCED regex patterns for your specific question formats
QUESTION_HEADER_RE = re.compile(
    r"^(?:(?:QUESTION|Q)\s+)?(?P<number>\d+(?:[.\-]\d+)*[A-Z]?|\d+[A-Z]|\d+)(?:\s*[:).\-–]\s*)?(?P<title>.*)$",
    re.IGNORECASE
)

# Additional patterns for question variations - FIXED for your formats
QUESTION_VARIATIONS_RE = [
    # "1.1.1 Case Study", "2.1.2 Multiple Choice", etc. - This matches your format!
    re.compile(r"^(?P<number>\d+(?:\.\d+)*)\s+(?P<title>(?:case\s+study|multiple\s+choice|constructive\s+response|constructive\s+respond|essay|short\s+answer|true\s+false|matching).*?)$", re.IGNORECASE),
    
    # "Question 1.1.1", "QUESTION 2.1", etc.
    re.compile(r"^(?:QUESTION|Q)\s+(?P<number>\d+(?:\.\d+)*)\s*[:.]?\s*(?P<title>.*)$", re.IGNORECASE),
    
    # "1.1.8 Multiple Choice Questions" - Handles plural too
    re.compile(r"^(?P<number>\d+(?:\.\d+)*)\s+(?P<title>.*?(?:questions?|case|study|response|choice|essay|answer).*?)$", re.IGNORECASE),
    
    # Just standalone numbers that might be questions: "1", "2", "3"
    re.compile(r"^(?P<number>\d+(?:\.\d+)*)\s*$"),
    
    # Section patterns
    re.compile(r"^(?:section\s+[A-Z]|part\s+[IVX]+)\s+question\s+(?P<number>\d+(?:\.\d+)*)\s*[:.]?\s*(?P<title>.*)$", re.IGNORECASE),
]

class Heuristics:
    @staticmethod
    def is_question_header(text: str) -> Optional[Tuple[str, str]]:
        """Enhanced question header detection - FIXED for your formats"""
        t = text.strip()
        if not t:
            return None
            
        # Skip obvious non-questions
        if PAGE_TRASH_RE.match(t) or SECTION_HEADING_RE.match(t):
            return None
            
        print(f"🔍 Testing text: '{t}'")
            
        # Try primary regex first
        m = QUESTION_HEADER_RE.match(t)
        if m:
            number = (m.group('number') or '').strip()
            title = (m.group('title') or '').strip()
            # Guard against random dates like 20240501
            if len(number) > 8 and number.count('.') == 0:
                return None
            print(f"✅ Primary regex matched: Q{number}")
            return number, title
        
        # Try additional patterns - These should catch your formats!
        for i, pattern in enumerate(QUESTION_VARIATIONS_RE):
            m = pattern.match(t)
            if m:
                number = (m.group('number') or '').strip()
                title = (m.group('title') or '').strip() if 'title' in m.groupdict() else ''
                
                # Validate it looks like a real question number
                if number and len(number) <= 10:  # Reasonable length
                    print(f"🎯 Pattern {i+1} matched: '{t}' → Q{number}")
                    return number, title
        
        # Special handling for your exact formats
        special_patterns = [
            # "1.1.1 Case Study" - exact match for your format
            (r"^(\d+(?:\.\d+)*)\s+(case\s+study)$", lambda m: (m.group(1), m.group(2))),
            
            # "2.1.1 Multiple Choice Question" - exact match
            (r"^(\d+(?:\.\d+)*)\s+(multiple\s+choice\s+questions?)$", lambda m: (m.group(1), m.group(2))),
            
            # "1.1.5 Constructive Response" - exact match  
            (r"^(\d+(?:\.\d+)*)\s+(constructive\s+respon[ds]e?)$", lambda m: (m.group(1), m.group(2))),
            
            # Simple number patterns: "Q 1", "Q 2", etc.
            (r"^Q\s+(\d+(?:\.\d+)*)\s*(.*)$", lambda m: (m.group(1), m.group(2))),
        ]
        
        text_lower = t.lower()
        for pattern_str, extractor in special_patterns:
            pattern = re.compile(pattern_str, re.IGNORECASE)
            m = pattern.match(text_lower)
            if m:
                number, title = extractor(m)
                print(f"🎯 Special pattern matched: '{t}' → Q{number}")
                return number, title or t
        
        return None

    @staticmethod
    def find_question_in_table(table_rows: List[List[str]]) -> Optional[Tuple[str, str, str]]:
        """Scan table cells for a question header.

        Returns (number, title, source_cell_text) if found, else None.
        """
        try:
            for r_idx, row in enumerate(table_rows or []):
                for c_idx, cell in enumerate(row or []):
                    t = (cell or "").strip()
                    if not t:
                        continue
                    q = Heuristics.is_question_header(t)
                    if q:
                        number, title = q
                        print(f"dY` Found question in table at r{r_idx}c{c_idx}: Q{number}")
                        return number, title, t
        except Exception as e:
            warn(f"Error scanning table for question headers: {e}")
        return None

    @staticmethod  
    def extract_marks_from_table_rows(table_rows: List[List[str]]) -> Optional[int]:
        """Extract marks from table rows - ENHANCED for your specific table format"""
        if not table_rows:
            return None
            
        print(f"📊 Analyzing table with {len(table_rows)} rows for marks...")
        
        try:
            # Strategy 1: Look for "Total" row - THIS SHOULD CATCH YOUR TABLES
            for row_idx, row in enumerate(table_rows):
                if len(row) > 0:
                    first_cell = row[0].strip().lower()
                    
                    if 'total' in first_cell:
                        print(f"🎯 Found total row at index {row_idx}: {row}")
                        
                        # Look for number in any cell of this row
                        for cell_idx, cell in enumerate(row):
                            cell_text = cell.strip()
                            print(f"   Cell {cell_idx}: '{cell_text}'")
                            
                            # Direct number check
                            if cell_text.isdigit():
                                marks = int(cell_text)
                                print(f"✅ Found direct marks: {marks}")
                                return marks
                            
                            # Extract numbers from text like "Total: 10" or "10 marks"
                            numbers = re.findall(r'\d+', cell_text)
                            if numbers:
                                marks = int(numbers[-1])  # Take last number
                                print(f"✅ Extracted marks: {marks} from '{cell_text}'")
                                return marks
            
            # Strategy 2: Find "Marks" column and sum individual values
            marks_column_index = None
            
            # Find the marks column header
            for row_idx, row in enumerate(table_rows):
                for col_idx, cell in enumerate(row):
                    cell_lower = cell.strip().lower()
                    if any(word in cell_lower for word in ['marks', 'mark', 'points', 'score']):
                        marks_column_index = col_idx
                        print(f"📊 Found marks column at index {col_idx}: '{cell}'")
                        break
                if marks_column_index is not None:
                    break
            
            if marks_column_index is not None:
                total_marks = 0
                print("🧮 Summing marks from column:")
                
                for row_idx, row in enumerate(table_rows):
                    if len(row) > marks_column_index:
                        cell_text = row[marks_column_index].strip()
                        
                        # Skip header row and total row
                        first_cell = row[0].strip().lower() if len(row) > 0 else ""
                        if 'total' in first_cell or any(header in cell_text.lower() for header in ['marks', 'mark', 'points']):
                            continue
                        
                        # Try to extract number
                        if cell_text.isdigit():
                            marks_value = int(cell_text)
                            total_marks += marks_value
                            print(f"   Row {row_idx}: +{marks_value} marks")
                        else:
                            # Try extracting number from text
                            numbers = re.findall(r'\d+', cell_text)
                            if numbers:
                                marks_value = int(numbers[0])
                                total_marks += marks_value
                                print(f"   Row {row_idx}: +{marks_value} marks (from '{cell_text}')")
                
                if total_marks > 0:
                    print(f"✅ Calculated total marks: {total_marks}")
                    return total_marks
            
            print("❌ No marks found in table")
            return None
                        
        except Exception as e:
            warn(f"Error extracting marks from table: {e}")
            return None

    @staticmethod
    def repair_numbering(nodes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Basic repair of question numbering - placeholder implementation"""
        # Simple implementation - you can enhance this later
        question_nodes = [n for n in nodes if n.get('type') == 'question']
        
        for i, node in enumerate(question_nodes):
            if not node.get('number'):
                # Assign sequential number if missing
                node['number'] = str(i + 1)
                print(f"🔧 Assigned missing number: Q{i + 1}")
        
        return nodes

    # Add the missing find_marks_in_following_blocks method with proper implementation
    @staticmethod
    def find_marks_in_following_blocks(blocks: List[Block], current_index: int) -> Optional[int]:
        """Look for marks in following table blocks - ENHANCED"""
        
        print(f"🔍 Looking for marks after question at index {current_index}")
        
        # Look at the next 7 blocks for tables with marks
        for i in range(current_index + 1, min(current_index + 8, len(blocks))):
            next_block = blocks[i]
            
            # Stop if we hit another question to avoid cross-contamination
            if next_block.type == 'paragraph':
                q = Heuristics.is_question_header(next_block.text)
                if q:
                    print(f"   🛑 Stopped at next question: {q[0]}")
                    break
            
            # Check if this is a table with marks
            if next_block.type == 'table' and next_block.table:
                print(f"   🔍 Checking table at index {i}")
                marks = Heuristics.extract_marks_from_table_rows(next_block.table)
                if marks:
                    print(f"✅ Found {marks} marks in following table")
                    return marks
                else:
                    print(f"   ❌ No marks found in this table")
        
        print(f"❌ No marks found in following blocks")
        return None

    @staticmethod
    def extract_marks(text: str) -> Optional[int]:
        """Extract marks from text using regex"""
        if not text:
            return None
        m = MARKS_RE.search(text)
        if m:
            # Return whichever group matched (group 1 or 2)
            return int(m.group(1) or m.group(2))
        return None

    @staticmethod
    def group_blocks(blocks: List[Block]) -> List[Dict[str, Any]]:
        """Enhanced grouping with better question detection and marks extraction"""
        nodes: List[Node] = []
        current: Optional[Node] = None
        order = 0

        print(f"🔄 Grouping {len(blocks)} blocks into questions...")

        for i, b in enumerate(blocks):
            if b.type == 'paragraph':
                q = Heuristics.is_question_header(b.text)
                if q:
                    number, title = q
                    print(f"📝 Found question Q{number}: '{title[:50]}...'")
                    
                    # Try to extract marks from question text first
                    marks = Heuristics.extract_marks(b.text)
                    
                    # If no marks in text, look in following tables
                    if marks is None:
                        marks = Heuristics.find_marks_in_following_blocks(blocks, i)
                        if marks:
                            print(f"   Marks from following table: {marks}")
                    else:
                        print(f"   Marks from question text: {marks}")
                    
                    current = Node(
                        id=b.id, 
                        number=number, 
                        type='question', 
                        marks=marks,
                        parent_id=None, 
                        order=order,
                        content=[{'type':'question_text','text': b.text}]
                    )
                    nodes.append(current)
                    order += 1
                    continue

            # Handle non-question blocks
            payload: Dict[str, Any]
            if b.type == 'table':
                # Before attaching the table, check if it itself contains a question header
                q_in_table = Heuristics.find_question_in_table(b.table)
                if q_in_table:
                    number, title, source_text = q_in_table
                    print(f"dY? Found question in table: Q{number} '{title[:50]}...'")
                    # Prefer marks in the question cell text; fallback to table total/column
                    marks = Heuristics.extract_marks(source_text)
                    if marks is None:
                        marks = Heuristics.extract_marks_from_table_rows(b.table)
                        if marks:
                            print(f"   Marks from table analysis: {marks}")
                    else:
                        print(f"   Marks from table cell: {marks}")

                    current = Node(
                        id=b.id,
                        number=number,
                        type='question',
                        marks=marks,
                        parent_id=None,
                        order=order,
                        content=[
                            {'type': 'question_text', 'text': source_text},
                            {'type': 'table', 'rows': b.table},
                        ],
                    )
                    nodes.append(current)
                    order += 1
                    continue
                payload = {'type':'table','rows': b.table}
            elif b.type == 'image':
                payload = {'type':'figure','images': b.images, 'caption': b.text}
            elif b.type == 'pagebreak':
                payload = {'type':'pagebreak'}
            else:
                payload = {'type':'paragraph','text': b.text}

            if current is None:
                # preamble/rubric/instructions before first question
                nodes.append(Node(
                    id=b.id, 
                    number=None, 
                    type='instruction', 
                    marks=None,
                    parent_id=None, 
                    order=order, 
                    content=[payload]
                ))
                order += 1
            else:
                current.content.append(payload)

        # parent/child via dotted numbers (1.1.1 → 1.1)
        number_to_id = {n.number: n.id for n in nodes if n.number}
        for n in nodes:
            if n.number and '.' in n.number:
                parent = n.number.rsplit('.',1)[0]
                n.parent_id = number_to_id.get(parent)
                
        print(f"✅ Grouped into {len(nodes)} nodes ({sum(1 for n in nodes if n.type == 'question')} questions)")
        return [asdict(n) for n in nodes]

# -----------------------------
# Core data structures
# -----------------------------
@dataclass
class Block:
    """Raw content block from DOCX parsing"""
    id: str
    type: str  # 'paragraph', 'table', 'image', 'pagebreak'
    text: str = ''
    table: List[List[str]] = field(default_factory=list)
    images: List[str] = field(default_factory=list)  # filenames
    
@dataclass 
class Node:
    """Structured exam node after grouping"""
    id: str
    number: Optional[str]
    type: str  # 'question', 'instruction', 'table', 'image'
    marks: Optional[int]
    parent_id: Optional[str]
    order: int
    content: List[Dict[str, Any]] = field(default_factory=list)

# -----------------------------
# DOCX Parser - the missing piece!
# -----------------------------
class DocxParser:
    """Parse DOCX at XML level to extract all content in order"""
    
    def __init__(self, docx_path: str, output_dir: Optional[str] = None):
        self.docx_path = docx_path
        self.output_dir = output_dir or self._auto_output_dir()
        self.media_dir = os.path.join(self.output_dir, 'media')
        os.makedirs(self.media_dir, exist_ok=True)
        
    def _auto_output_dir(self) -> str:
        """Generate output directory based on input filename"""
        base = os.path.splitext(os.path.basename(self.docx_path))[0]
        return f"{base}_extract"
        
    def parse(self) -> List[Block]:
        """Main parsing entry point preserving exact order from the body."""
        blocks = []
        try:
            with zipfile.ZipFile(self.docx_path, 'r') as docx_zip:
                doc_xml = docx_zip.read('word/document.xml')
                root = ET.fromstring(doc_xml)

                image_rels = self._parse_image_relationships(docx_zip)

                body = root.find('.//w:body', NS)
                if body is None:
                    warn("No document body found.")
                    return blocks

                for elem in list(body):
                    tag = elem.tag
                    if tag.endswith('}p'):
                        block = self._parse_paragraph(elem, docx_zip, image_rels)
                        if block:
                            blocks.append(block)
                        # detect page breaks inside the paragraph (optional but handy)
                        for br in elem.findall('.//w:br', NS):
                            if br.get('{%s}type' % NS['w']) == 'page':
                                blocks.append(Block(id=str(uuid.uuid4()), type='pagebreak'))
                    elif tag.endswith('}tbl'):
                        block = self._parse_table(elem)
                        if block:
                            blocks.append(block)
                    # ignore other element types
        except Exception as e:
            warn(f"DOCX parsing failed: {e}")

        return blocks

        
    def _parse_image_relationships(self, docx_zip) -> Dict[str, str]:
        """Parse image relationships from DOCX (namespace-safe)."""
        image_rels = {}
        try:
            rels_xml = docx_zip.read('word/_rels/document.xml.rels')
            rels_root = ET.fromstring(rels_xml)

            # Properly query with the package rels namespace
            for rel in rels_root.findall('.//pr:Relationship', NS):
                rel_type = rel.get('Type', '')
                # match any image relationship
                if rel_type.endswith('/image'):
                    rel_id = rel.get('Id')
                    target = rel.get('Target') or ''
                    if rel_id and target:
                        # Normalize target like '../media/image1.png'
                        target = target.lstrip('./')
                        if not target.startswith('word/'):
                            target = os.path.normpath(os.path.join('word', target))
                        image_rels[rel_id] = target
                        print(f"🖼️ Found image relationship: {rel_id} → {target}")
        except Exception as e:
            warn(f"Image relationships parsing failed: {e}")

        return image_rels

        
    def _parse_paragraph(self, para_elem, docx_zip, image_rels) -> Optional[Block]:
        """Parse a paragraph: collect text, DrawingML images, and VML images."""
        block_id = str(uuid.uuid4())
        text_chunks = []
        images: list[str] = []
        captions: list[str] = []

        # -------- text collection (preserve tabs/linebreaks) --------
        for node in para_elem.iter():
            tag = node.tag
            # text runs
            if tag.endswith('}t'):
                if node.text:
                    # remove soft hyphen and normalize nbsp
                    txt = node.text.replace('\u00AD', '').replace('\u00A0', ' ')
                    text_chunks.append(txt)
            # tab
            elif tag.endswith('}tab'):
                text_chunks.append('\t')
            # line breaks and carriage returns
            elif tag.endswith('}br') or tag.endswith('}cr'):
                text_chunks.append('\n')

        # -------- DrawingML images (w:drawing) --------
        for drawing in para_elem.findall('.//w:drawing', NS):
            fname = self._extract_image(drawing, docx_zip, image_rels)
            if fname:
                images.append(fname)
                # try to capture alt/descr from docPr
                try:
                    dp = drawing.find('.//wp:docPr', NS)
                    if dp is not None:
                        alt = (dp.get('descr') or dp.get('name') or '').strip()
                        if alt:
                            captions.append(alt)
                except Exception:
                    pass

        # -------- VML images (legacy: w:pict/v:imagedata) --------
        for pict in para_elem.findall('.//w:pict', NS):
            for im in pict.findall('.//v:imagedata', NS):
                rid = im.get('{%s}id' % NS['r']) or im.get('{%s}link' % NS['r'])
                if rid and rid in image_rels:
                    try:
                        image_path = image_rels[rid]  # already normalized in rels parser
                        data = docx_zip.read(image_path)
                        ext = os.path.splitext(image_path)[1] or '.png'
                        filename = f"image_{hashlib.md5(data).hexdigest()[:8]}{ext}"
                        with open(os.path.join(self.media_dir, filename), 'wb') as f:
                            f.write(data)
                        images.append(filename)
                        # alt from Office namespace: o:title
                        alt = im.get('{urn:schemas-microsoft-com:office:office}title')
                        if alt:
                            captions.append(alt.strip())
                    except Exception as e:
                        warn(f"VML image extraction failed: {e}")

        # dedupe images while preserving order
        if images:
            images = list(dict.fromkeys(images))

        # assemble and normalize text
        text = ''.join(text_chunks)
        # collapse multiple spaces (keep tabs/newlines intact)
        text = re.sub(r'[ \u00A0]+', ' ', text).strip()

        # if no visible text but we do have captions, use them as a caption
        if not text and captions:
            text = ' '.join(dict.fromkeys(captions))

        if not text and not images:
            return None

        block_type = 'image' if images else 'paragraph'
        return Block(
            id=block_id,
            type=block_type,
            text=text,
            images=images
        )

        
    def _parse_table(self, table_elem) -> Optional[Block]:
        """Parse table element"""
        rows = []
        
        for row_elem in table_elem.findall('.//w:tr', NS):
            cells = []
            for cell_elem in row_elem.findall('.//w:tc', NS):
                cell_text_parts = []
                for text_elem in cell_elem.findall('.//w:t', NS):
                    if text_elem.text:
                        cell_text_parts.append(text_elem.text)
                cell_text = ''.join(cell_text_parts).strip()
                cells.append(cell_text)
            if cells:
                rows.append(cells)
                
        if not rows:
            return None
            
        return Block(
            id=str(uuid.uuid4()),
            type='table',
            table=rows
        )
        
    def _extract_image(self, drawing_elem, docx_zip, image_rels) -> Optional[str]:
        """Extract and save image from a w:drawing element with robust path + type handling."""
        try:
            blip = drawing_elem.find('.//a:blip', NS)
            if blip is None:
                return None

            rid = blip.get(f'{{{NS["r"]}}}embed') or blip.get(f'{{{NS["r"]}}}link')
            if not rid:
                return None

            target = image_rels.get(rid)
            if not target:
                return None

            # Skip external links (cannot read from zip)
            if target.startswith(('http://', 'https://')):
                warn(f"Skipping external image link: {target}")
                return None

            # Normalize to a zip path under word/
            # Relationship targets in document.xml.rels are relative to 'word/'
            target = target.lstrip('/')  # strip leading slash if present
            if not target.startswith('word/'):
                target = os.path.normpath(os.path.join('word', target)).replace('\\', '/')

            # Collapse any ../ and ensure we stay under word/
            target = os.path.normpath(target).replace('\\', '/')
            if not target.startswith('word/'):
                # last-ditch: drop to media folder with basename only
                target = f"word/media/{os.path.basename(target)}"

            # Read bytes from docx
            try:
                data = docx_zip.read(target)
            except KeyError:
                # Some writers store just 'media/x' (already under word/)
                alt_target = target.replace('word/word/', 'word/')
                if alt_target != target:
                    try:
                        data = docx_zip.read(alt_target)
                        target = alt_target
                    except KeyError:
                        warn(f"Image not found in zip: {target}")
                        return None
                else:
                    warn(f"Image not found in zip: {target}")
                    return None

            ext = guess_ext(data, target)

            # Name by content hash to dedupe across repeats
            hexd = hashlib.md5(data).hexdigest()[:16]
            filename = f"image_{hexd}{ext}"
            output_path = os.path.join(self.media_dir, filename)

            # Write once
            if not os.path.exists(output_path):
                os.makedirs(self.media_dir, exist_ok=True)
                with open(output_path, 'wb') as f:
                    f.write(data)

            return filename

        except Exception as e:
            warn(f"Image extraction failed: {e}")
            return None



# -----------------------------
# LLM adapters (Gemini / Gemma) — optional
# -----------------------------
class LLM:
    def __init__(self, enable_gemini: bool, enable_gemma: bool):
        self.enable_gemini = enable_gemini and GEMINI_AVAILABLE and bool(os.getenv('GEMINI_API_KEY'))
        if self.enable_gemini:
            genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
        self.gemini_model = 'models/gemini-2.5-flash'
        self.enable_gemma = enable_gemma
        # Support both legacy and common env var names
        self.gemma_endpoint = (
            os.getenv('GEMMA_OLLAMA_URL')
            or os.getenv('OLLAMA_BASE_URL')
            or os.getenv('OLLAMA_URL')
            or 'http://localhost:11434/api/generate'
        )
        self.gemma_model = os.getenv('GEMMA_OLLAMA_MODEL') or os.getenv('OLLAMA_MODEL') or 'gemma2:2b'

    def gemini_validate(self, nodes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        if not self.enable_gemini:
            return nodes
        prompt = (
            "You will receive an ordered list of exam paper nodes. "
            "Fix misattached figures/tables by moving them to the nearest preceding question, "
            "merge stray short paragraphs into their question, and ensure consistent numbering. "
            "Return ONLY the corrected JSON list."
        )
        try:
            model = genai.GenerativeModel(self.gemini_model)
            resp = model.generate_content([prompt, json.dumps(nodes)])
            raw = (resp.text or '').strip()
            cleaned = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
            return json.loads(cleaned)
        except Exception as e:
            warn(f"Gemini validation failed: {e}")
            return nodes

    def gemma_backfill(self, nodes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        if not self.enable_gemma:
            return nodes
        try:
            req = {
                "model": self.gemma_model,
                "prompt": (
                    "You are given an ordered JSON list of exam nodes extracted from a DOCX. "
                    "Each node has: type, optional number, optional marks, and content array of blocks (question_text, paragraph, table, figure, pagebreak).\n"
                    "Goal: Fix the structure automatically.\n"
                    "- If any paragraph or table cell text contains a question header (e.g., '1', '1.1', '1.1.1', 'Q 2', 'Question 3', '2A'), create a new node of type 'question' at that point.\n"
                    "- Move subsequent related content (paragraphs/tables/figures until the next question) under that question's content.\n"
                    "- Extract marks from text like '(10 marks)' or '[10]' or from tables (Totals/Marks columns) and set node.marks.\n"
                    "- Maintain parent_id relationship implicitly via dotted numbering (1.1 -> parent 1). Do not invent parents, just set children numbers.\n"
                    "- Preserve order. Do not drop content.\n"
                    "- If nothing needs fixing, return the original list.\n"
                    "Return ONLY the corrected JSON array of nodes (no prose).\n\n" + json.dumps(nodes)
                ),
                "stream": False
            }
            r = requests.post(self.gemma_endpoint, json=req, timeout=120)
            cleaned = re.sub(r"^```(?:json)?|```$", "", r.json().get('response',''), flags=re.MULTILINE).strip()
            return json.loads(cleaned)
        except Exception as e:
            warn(f"Gemma backfill failed: {e}")
            return nodes

# -----------------------------
# Orchestrator
# -----------------------------
class Storage:
    """Local filesystem storage for extracted papers to enable CLI randomization."""
    def __init__(self, root: str):
        self.root = root
        os.makedirs(self.root, exist_ok=True)

    def _safe_name(self, base: str) -> str:
        base = re.sub(r"[^A-Za-z0-9_.-]", "_", base)
        return base[:80]

    def save_manifest(self, manifest: Dict[str, Any]) -> str:
        ts = time.strftime("%Y%m%dT%H%M%S")
        name = self._safe_name(os.path.splitext(manifest.get('source','paper'))[0])
        paper_dir = os.path.join(self.root, f"{name}_{ts}")
        os.makedirs(paper_dir, exist_ok=True)
        # copy JSON
        with open(os.path.join(paper_dir, 'manifest.json'), 'w', encoding='utf-8') as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
        # copy media
        src_media = os.path.join(manifest['output_dir'], 'media')
        dst_media = os.path.join(paper_dir, 'media')
        if os.path.isdir(src_media):
            os.makedirs(dst_media, exist_ok=True)
            for fn in os.listdir(src_media):
                try:
                    s = os.path.join(src_media, fn)
                    d = os.path.join(dst_media, fn)
                    if os.path.isfile(s) and not os.path.exists(d):
                        with open(s, 'rb') as r, open(d, 'wb') as w:
                            w.write(r.read())
                except Exception as e:
                    warn(f"media copy failed for {fn}: {e}")
        ok(f"Saved to storage: {paper_dir}")
        return paper_dir

    def load_bank(self) -> List[Dict[str, Any]]:
        bank = []
        for item in os.listdir(self.root):
            p = os.path.join(self.root, item)
            if not os.path.isdir(p):
                continue
            mf = os.path.join(p, 'manifest.json')
            if os.path.isfile(mf):
                try:
                    with open(mf, 'r', encoding='utf-8') as f:
                        m = json.load(f)
                        m['__storage_dir'] = p
                        bank.append(m)
                except Exception as e:
                    warn(f"Failed reading {mf}: {e}")
        info(f"Loaded bank manifests: {len(bank)}")
        return bank

    def copy_images_into(self, manifest: Dict[str, Any], out_media_dir: str):
        os.makedirs(out_media_dir, exist_ok=True)
        src_media = os.path.join(manifest['__storage_dir'], 'media')
        if not os.path.isdir(src_media):
            return
        for fn in os.listdir(src_media):
            s = os.path.join(src_media, fn)
            d = os.path.join(out_media_dir, fn)
            try:
                if os.path.isfile(s) and not os.path.exists(d):
                    with open(s, 'rb') as r, open(d, 'wb') as w:
                        w.write(r.read())
            except Exception as e:
                warn(f"image copy failed {fn}: {e}")


def randomize_nodes(original: Dict[str, Any], bank: List[Dict[str, Any]], out_media_dir: str, *, require_same_top: bool=False, marks_tolerance: Optional[int]=None, required_tags: Optional[List[str]]=None) -> Dict[str, Any]:
    """Randomize each question's content by sampling another paper's same-number question.
    Keeps numbering/structure, swaps content. If none found, keeps original.
    Images from donor are copied into current media dir.
    """
    import random
    # Build index number -> list[(manifest, node)]
    index: Dict[str, List[Tuple[Dict[str, Any], Dict[str, Any]]]] = {}
    top_index: Dict[str, List[Tuple[Dict[str, Any], Dict[str, Any]]]] = {}
    for mf in bank:
        for n in mf.get('nodes', []):
            if n.get('type') == 'question' and n.get('number'):
                number = n['number']
                index.setdefault(number, []).append((mf, n))
                # build top-level index e.g., 1.1 -> 1
                top = number.split('.')[0]
                top_index.setdefault(top, []).append((mf, n))

    new_nodes: List[Dict[str, Any]] = []
    for n in original.get('nodes', []):
        if n.get('type') != 'question' or not n.get('number'):
            new_nodes.append(n)
            continue
        # choose candidates
        if require_same_top:
            key = n['number'].split('.')[0]
            candidates = top_index.get(key, [])
        else:
            candidates = index.get(n['number'], [])
        # exclude same manifest source if available
        candidates = [c for c in candidates if c[0].get('source') != original.get('source')] or candidates
        # filter by tags if requested
        if required_tags:
            req = {t.lower().strip() for t in required_tags}
            def has_tags(mf):
                tags = (mf.get('metadata',{}) or {}).get('tags', [])
                return bool(req.intersection({str(t).lower().strip() for t in tags}))
            candidates = [c for c in candidates if has_tags(c[0])] or candidates
        # marks tolerance filter
        if marks_tolerance is not None:
            try:
                want = int(n.get('marks') or 0)
                def within(dnode):
                    try:
                        got = int(dnode.get('marks') or 0)
                        return abs(got - want) <= int(marks_tolerance)
                    except Exception:
                        return False
                candidates = [c for c in candidates if within(c[1])] or candidates
            except Exception:
                pass
        if not candidates:
            new_nodes.append(n)
            continue
        donor_mf, donor_node = random.choice(candidates)
        # copy donor images
        Storage("").copy_images_into(donor_mf, out_media_dir)
        # construct swapped node: keep id/number/parent, replace content and marks
        swapped = dict(n)
        swapped['content'] = donor_node.get('content', [])
        swapped['marks'] = donor_node.get('marks', n.get('marks'))
        new_nodes.append(swapped)

    randomized = dict(original)
    randomized['nodes'] = new_nodes
    randomized['counts'] = dict(original.get('counts', {}))
    randomized['counts']['questions'] = sum(1 for x in new_nodes if x.get('type')=='question')
    randomized['source'] = f"Randomized_{original.get('source','paper')}"
    return randomized

# -----------------------------
# Orchestrator
# -----------------------------
class Extractor:
    def __init__(self, use_gemini: bool = True, use_gemma: bool = True):
        self.llm = LLM(use_gemini, use_gemma)

    @staticmethod
    def eta_minutes(blocks_count: int, images_count: int) -> float:
        return round(0.02*blocks_count + 0.08*images_count + 0.3, 2)

    def run(self, docx_path: str, output_dir: Optional[str] = None) -> Dict[str, Any]:
        t0 = time.time()
        parser = DocxParser(docx_path, output_dir)
        info("Parsing DOCX…")
        blocks = parser.parse()
        img_count = sum(len(b.images) for b in blocks)
        tbl_count = sum(1 for b in blocks if b.type == 'table')
        ok(f"Parsed {len(blocks)} blocks | images={img_count} tables={tbl_count}")
        info(f"ETA ~{self.eta_minutes(len(blocks), img_count)} minutes (conservative)")

        # Group into nodes
        info("Grouping into questions and attaching trailing content…")
        nodes = Heuristics.group_blocks(blocks)

        # Backfill/repair numbering/marks locally first
        nodes = Heuristics.repair_numbering(nodes)

        # LLM passes (optional)
        nodes = self.llm.gemma_backfill(nodes)
        nodes = self.llm.gemini_validate(nodes)

        # Optional metadata scaffold (filled via CLI flags/interactive)
        metadata = {
            'paper_name': None,
            'module': None,
            'tags': []
        }

        manifest = {
            'source': os.path.basename(docx_path),
            'output_dir': parser.output_dir,
            'generated_at': int(time.time()),
            'counts': {
                'blocks': len(blocks),
                'images': img_count,
                'tables': tbl_count,
                'questions': sum(1 for n in nodes if n['type']=='question')
            },
            'nodes': nodes,
            'metadata': metadata
        }

        # Save JSON + HTML preview
        os.makedirs(parser.output_dir, exist_ok=True)
        out_json = os.path.join(parser.output_dir, 'structure_json.json')
        with open(out_json, 'w', encoding='utf-8') as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
        ok(f"Wrote {out_json}")

        html_path = os.path.join(parser.output_dir, 'preview.html')
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(self._preview_html(manifest))
        ok(f"Wrote {html_path}")

        # Console pretty print — tables and storage
        self._console_summary(manifest)

        info(f"Total time: {round(time.time()-t0,2)}s")
        return manifest

    # ---------- console helpers ----------
    def _console_summary(self, manifest: Dict[str, Any]):
        print(f"\n{BOLD}=== STORAGE SUMMARY ==={RESET}")
        print(f"Source: {manifest['source']}")
        print(f"Output: {manifest['output_dir']}")
        print(f"Counts: {manifest['counts']}")

        print(f"\n{BOLD}=== QUESTION OVERVIEW ==={RESET}")
        for n in manifest['nodes']:
            if n['type'] == 'question':
                marks = n.get('marks')
                print(f"{GREEN}Q {n.get('number','?')}{RESET}  marks={marks if marks is not None else '-'}")
                for c in n.get('content', []):
                    if c['type'] == 'table':
                        self._print_table(c['rows'])
                    elif c['type'] == 'figure':
                        print(f"  [figure] images={len(c.get('images',[]))}")

    def _print_table(self, rows: List[List[str]]):
        print(f"  {DIM}-- table start --{RESET}")
        # compute column widths up to 80 chars per row
        if not rows:
            print("  (empty table)")
            return
        cols = max(len(r) for r in rows)
        widths = [0]*cols
        for r in rows:
            for i, cell in enumerate(r):
                widths[i] = min(40, max(widths[i], len((cell or '').split('\n')[0])))
        for r in rows[:20]:  # cap terminal spam
            line = " | ".join(((cell or '').replace('\n',' / ')[:widths[i]]).ljust(widths[i]) for i, cell in enumerate(r))
            print("   ", line)
        if len(rows) > 20:
            print(f"   … ({len(rows)} rows total)")
        print(f"  {DIM}-- table end --{RESET}")

    # ---------- preview ----------
   

    def _preview_html(self, manifest: Dict[str, Any], *, embed_images: bool=False) -> str:
        out_dir = manifest.get('output_dir') or '.'
        media_dir = os.path.join(out_dir, 'media')
        base_uri = Path(out_dir).resolve().as_uri() + '/'  # ensures 'media/...' resolves

        def img_src(fname: str) -> str:
            if not embed_images:
                return f"media/{fname}"
            # inline as data URI
            p = os.path.join(media_dir, fname)
            try:
                with open(p, 'rb') as f:
                    b64 = base64.b64encode(f.read()).decode('ascii')
                mime = mimetypes.guess_type(p)[0] or "application/octet-stream"
                return f"data:{mime};base64,{b64}"
            except Exception:
                # fallback to relative path if inlining fails
                return f"media/{fname}"

        buf = io.StringIO()
        buf.write("<html><head><meta charset='utf-8'><title>Preview</title>")
        buf.write(f"<base href='{base_uri}'>")  # make relative img URLs work
        buf.write("</head><body>")
        buf.write(f"<h2>Preview — {manifest.get('source','')}</h2>")

        for n in manifest.get('nodes', []):
            label = n.get('number') or n.get('type', '')
            buf.write(f"<h3>{label}</h3>")
            for c in n.get('content', []):
                t = c.get('type')
                if t in ('question_text', 'paragraph'):
                    buf.write(f"<p>{c.get('text','')}</p>")
                elif t == 'table':
                    rows = c.get('rows') or []
                    buf.write("<table border='1' cellspacing='0' cellpadding='4'>")
                    for row in rows:
                        buf.write('<tr>')
                        for cell in row:
                            buf.write(f"<td>{(cell or '').replace('\n','<br/>')}</td>")
                        buf.write('</tr>')
                    buf.write('</table>')
                elif t == 'figure':
                    images = c.get('images') or []
                    caption = c.get('caption') or ''
                    for img in images:
                        src = img_src(img)
                        buf.write("<div>")
                        buf.write(f"<img src='{src}' style='max-width:600px'>")
                        if caption:
                            buf.write("<br/><em>{}</em>".format(caption))
                        buf.write("</div>")
                elif t == 'pagebreak':
                    buf.write("<hr/>")

        buf.write("</body></html>")
        return buf.getvalue()

# -----------------------------
# DOCX reconstruction (export)
# -----------------------------

def export_docx_from_manifest(manifest: Dict[str, Any], out_docx_path: str, media_dir: Optional[str] = None) -> bool:
    """Reconstruct a readable .docx from a manifest structure.
    - Writes headings for questions
    - Renders paragraphs, tables, and embeds figures from media_dir
    """
    try:
        from docx import Document  # python-docx
        from docx.shared import Inches
    except Exception as e:
        warn(f"python-docx not available: {e}")
        return False

    media_dir = media_dir or os.path.join(manifest.get('output_dir') or '', 'media')
    doc = Document()

    title = (manifest.get('metadata') or {}).get('paper_name') or manifest.get('source') or 'Reconstructed Paper'
    doc.add_heading(str(title), level=1)

    def add_table(rows: List[List[str]]):
        if not rows:
            return
        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=0, cols=cols)
        t.style = 'Table Grid'
        for r in rows:
            row = t.add_row().cells
            for i in range(cols):
                val = (r[i] if i < len(r) else '') or ''
                row[i].text = str(val)

    for node in manifest.get('nodes', []):
        ntype = node.get('type')
        number = node.get('number')
        if ntype == 'question':
            hdr = f"Question {number}" if number else "Question"
            marks = node.get('marks')
            if marks is not None:
                hdr += f" ({marks} marks)"
            doc.add_heading(hdr, level=2)
        elif ntype in ('instruction','rubric','heading'):
            # Make it stand out a bit
            doc.add_heading(ntype.capitalize(), level=3)
        # Content blocks
        for c in node.get('content', []):
            t = c.get('type')
            if t in ('question_text','paragraph'):
                txt = c.get('text') or ''
                if txt.strip():
                    doc.add_paragraph(txt)
            elif t == 'table':
                add_table(c.get('rows') or [])
            elif t == 'figure':
                imgs = c.get('images') or []
                for img in imgs:
                    p = os.path.join(media_dir, img)
                    try:
                        if os.path.isfile(p):
                            doc.add_picture(p)  # default size; users can resize in Word
                    except Exception as e:
                        warn(f"Failed to add image {img}: {e}")
            elif t == 'pagebreak':
                try:
                    doc.add_page_break()
                except Exception:
                    pass
        # spacing between nodes
        doc.add_paragraph('')

    try:
        os.makedirs(os.path.dirname(out_docx_path) or '.', exist_ok=True)
        doc.save(out_docx_path)
        ok(f"Reconstructed DOCX written: {out_docx_path}")
        return True
    except Exception as e:
        err(f"Failed to save reconstructed DOCX: {e}")
        return False

# -----------------------------
# Public API
# -----------------------------

def extract_docx(docx_path: str, out_dir: Optional[str] = None, use_gemini: bool = True, use_gemma: bool = True) -> Dict[str, Any]:
    return Extractor(use_gemini, use_gemma).run(docx_path, out_dir)


def extract_docx_bytes(fileobj: io.BytesIO, out_dir: Optional[str] = None, use_gemini: bool = True, use_gemma: bool = True) -> Dict[str, Any]:
    """Convenience wrapper for Django/Flask uploads: pass a file-like object.
    Saves to a temp .docx and runs the main extractor.
    """
    import tempfile
    fileobj.seek(0)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(fileobj.read())
        tmp_path = tmp.name
    try:
        return extract_docx(tmp_path, out_dir, use_gemini, use_gemma)
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

# -----------------------------
# GUI (optional, Tkinter)
# -----------------------------

def launch_gui():
    try:
        import threading
        import tkinter as tk
        from tkinter import filedialog, ttk, messagebox
    except Exception as e:
        err(f"GUI not available: {e}")
        return

    root = tk.Tk()
    root.title("Robust Exam Extraction")
    root.geometry("760x640")

    # --- helpers ---
    def browse_docx():
        path = filedialog.askopenfilename(filetypes=[("Word documents","*.docx"), ("All files","*.*")])
        if path:
            docx_var.set(path)
            if not out_var.get():
                out_var.set(os.path.splitext(path)[0] + "_extract")

    def browse_out():
        path = filedialog.askdirectory()
        if path:
            out_var.set(path)

    def browse_storage():
        path = filedialog.askdirectory()
        if path:
            storage_var.set(path)

    def log(msg):
        console.configure(state='normal')
        console.insert('end', msg + "\n")
        console.see('end')
        console.configure(state='disabled')
        root.update_idletasks()

    def run_extract(randomize_after: bool=False, randomize_only: bool=False):
        def worker():
            try:
                storage = Storage(storage_var.get() or 'extracted_papers')
                manifest = None
                if not randomize_only:
                    if not docx_var.get():
                        messagebox.showerror("Missing file", "Please choose a .docx file")
                        return
                    log("Starting extraction…")
                    manifest = extract_docx(docx_var.get(), out_var.get() or None, use_gemini=gemini_var.get(), use_gemma=gemma_var.get())
                    # inject metadata
                    md = manifest.get('metadata', {}) or {}
                    if paper_name_var.get():
                        md['paper_name'] = paper_name_var.get()
                    if module_var.get():
                        md['module'] = module_var.get()
                    if tags_var.get():
                        md['tags'] = [t.strip() for t in tags_var.get().split()]
                    manifest['metadata'] = md
                    if save_var.get():
                        storage.save_manifest(manifest)
                    if export_var.get():
                        out_dir = out_var.get() or manifest.get('output_dir')
                        media_dir = os.path.join(out_dir, 'media')
                        export_docx_from_manifest(manifest, os.path.join(out_dir, 'reconstructed.docx'), media_dir=media_dir)
                    log("Extraction done.")

                if randomize_after or randomize_only:
                    log("Loading bank for randomization…")
                    bank = storage.load_bank()
                    if not bank:
                        messagebox.showwarning("No bank", "No stored manifests found in storage folder.")
                        return
                    # pick target
                    target = manifest
                    if randomize_only or target is None:
                        items = [os.path.join(storage_var.get() or 'extracted_papers', d)
                                 for d in os.listdir(storage_var.get() or 'extracted_papers')
                                 if os.path.isdir(os.path.join(storage_var.get() or 'extracted_papers', d))]
                        if not items:
                            messagebox.showwarning("Empty storage", "Storage has no entries.")
                            return
                        latest = max(items, key=os.path.getmtime)
                        with open(os.path.join(latest, 'manifest.json'), 'r', encoding='utf-8') as f:
                            target = json.load(f)
                            target['output_dir'] = target.get('output_dir') or latest
                            target['__storage_dir'] = latest

                    out_dir = out_var.get() or (os.path.splitext(docx_var.get() or 'randomized')[0] + '_randomized')
                    os.makedirs(out_dir, exist_ok=True)
                    out_media = os.path.join(out_dir, 'media')
                    os.makedirs(out_media, exist_ok=True)

                    log("Randomizing…")
                    randomized = randomize_nodes(
                        target, bank, out_media,
                        require_same_top=same_top_var.get(),
                        marks_tolerance=(int(marks_tol_var.get()) if marks_tol_var.get().strip().isdigit() else None),
                        required_tags=[t.strip() for t in rule_tags_var.get().split()] if rule_tags_var.get().strip() else None
                    )
                    with open(os.path.join(out_dir, 'structure_json.json'), 'w', encoding='utf-8') as f:
                        json.dump(randomized, f, ensure_ascii=False, indent=2)
                    with open(os.path.join(out_dir, 'preview.html'), 'w', encoding='utf-8') as f:
                        f.write(Extractor()._preview_html(randomized))
                    if export_var.get():
                        export_docx_from_manifest(randomized, os.path.join(out_dir, 'reconstructed.docx'), media_dir=out_media)
                    log("Randomization done.")
                messagebox.showinfo("Done", "Operation completed.")
            except Exception as e:
                err_txt = f"Error: {e}\n{traceback.format_exc()}"
                log(err_txt)
                messagebox.showerror("Error", str(e))

        threading.Thread(target=worker, daemon=True).start()

    # --- layout ---
    frm = ttk.Frame(root, padding=12)
    frm.pack(fill='both', expand=True)

    # File selectors
    docx_var = tk.StringVar()
    out_var = tk.StringVar()
    storage_var = tk.StringVar(value='extracted_papers')

    ttk.Label(frm, text="DOCX file:").grid(row=0, column=0, sticky='w')
    ttk.Entry(frm, textvariable=docx_var, width=70).grid(row=0, column=1, sticky='we')
    ttk.Button(frm, text="Browse", command=browse_docx).grid(row=0, column=2, padx=4)

    ttk.Label(frm, text="Output dir:").grid(row=1, column=0, sticky='w')
    ttk.Entry(frm, textvariable=out_var, width=70).grid(row=1, column=1, sticky='we')
    ttk.Button(frm, text="Browse", command=browse_out).grid(row=1, column=2, padx=4)

    ttk.Label(frm, text="Storage dir:").grid(row=2, column=0, sticky='w')
    ttk.Entry(frm, textvariable=storage_var, width=70).grid(row=2, column=1, sticky='we')
    ttk.Button(frm, text="Browse", command=browse_storage).grid(row=2, column=2, padx=4)

    # Options
    gemini_var = tk.BooleanVar(value=False)
    gemma_var = tk.BooleanVar(value=False)
    save_var = tk.BooleanVar(value=False)
    randomize_var = tk.BooleanVar(value=False)
    export_var = tk.BooleanVar(value=True)

    row = 3
    ttk.Label(frm, text="Options:").grid(row=row, column=0, sticky='w')
    ttk.Checkbutton(frm, text="Use Gemini", variable=gemini_var).grid(row=row, column=1, sticky='w')
    ttk.Checkbutton(frm, text="Use Gemma", variable=gemma_var).grid(row=row, column=1, sticky='e')
    row += 1
    ttk.Checkbutton(frm, text="Save to storage", variable=save_var).grid(row=row, column=1, sticky='w')
    ttk.Checkbutton(frm, text="Randomize after extraction", variable=randomize_var).grid(row=row, column=1, sticky='e')
    row += 1
    ttk.Checkbutton(frm, text="Export DOCX", variable=export_var).grid(row=row, column=1, sticky='w')

    # Metadata
    row += 1
    ttk.Label(frm, text="Paper name:").grid(row=row, column=0, sticky='w')
    paper_name_var = tk.StringVar()
    ttk.Entry(frm, textvariable=paper_name_var, width=70).grid(row=row, column=1, columnspan=2, sticky='we'); row += 1

    ttk.Label(frm, text="Module:").grid(row=row, column=0, sticky='w')
    module_var = tk.StringVar()
    ttk.Entry(frm, textvariable=module_var, width=70).grid(row=row, column=1, columnspan=2, sticky='we'); row += 1

    ttk.Label(frm, text="Tags (space-separated):").grid(row=row, column=0, sticky='w')
    tags_var = tk.StringVar()
    ttk.Entry(frm, textvariable=tags_var, width=70).grid(row=row, column=1, columnspan=2, sticky='we'); row += 1

    # Rules
    ttk.Label(frm, text="Rules:").grid(row=row, column=0, sticky='w')
    same_top_var = tk.BooleanVar(value=False)
    ttk.Checkbutton(frm, text="Same top-level only (1.x with 1.x)", variable=same_top_var).grid(row=row, column=1, sticky='w')
    row += 1
    ttk.Label(frm, text="Marks tolerance:").grid(row=row, column=0, sticky='w')
    marks_tol_var = tk.StringVar()
    ttk.Entry(frm, textvariable=marks_tol_var, width=10).grid(row=row, column=1, sticky='w')
    ttk.Label(frm, text="(leave blank for no filter)").grid(row=row, column=2, sticky='w')
    row += 1
    ttk.Label(frm, text="Required tags:").grid(row=row, column=0, sticky='w')
    rule_tags_var = tk.StringVar()
    ttk.Entry(frm, textvariable=rule_tags_var, width=70).grid(row=row, column=1, columnspan=2, sticky='we'); row += 1

    # --- console output ---
    console = tk.Text(frm, height=10, wrap='word')
    console.grid(row=row, column=0, columnspan=3, sticky='nsew')
    frm.columnconfigure(0, weight=1)
    frm.rowconfigure(row, weight=1)

    # --- buttons ---
    row += 1
    ttk.Button(frm, text="Extract", command=lambda: run_extract(randomize_after=False)).grid(row=row, column=0, padx=4)
    ttk.Button(frm, text="Extract & Randomize", command=lambda: run_extract(randomize_after=True)).grid(row=row, column=1, padx=4)
    ttk.Button(frm, text="Cancel", command=root.quit).grid(row=row, column=2, padx=4)

    # ---------- run ----------
    if len(ap.parse_args().docx or '') > 0:
        docx_var.set(ap.parse_args().docx)
        out_var.set(ap.parse_args().out or '')
        gemini_var.set(ap.parse_args().gemini)
        gemma_var.set(ap.parse_args().gemma)
        storage_var.set(ap.parse_args().storage)
        save_var.set(ap.parse_args().save)
        randomize_var.set(ap.parse_args().randomize)
        export_var.set(ap.parse_args().export_docx)
        # metadata / rules
        paper_name_var.set(ap.parse_args().paper_name or '')
        module_var.set(ap.parse_args().module or '')
        tags_var.set(' '.join(ap.parse_args().tags or []))
        same_top_var.set(ap.parse_args().same_top)
        marks_tol_var.set(ap.parse_args().marks_tolerance or '')
        rule_tags_var.set(' '.join(ap.parse_args().required_tags or []))
        # start extraction in a thread
        run_extract(ap.parse_args().randomize, ap.parse_args().randomize_only)

    root.mainloop()
# Add this to utils.py or create a new file like robust_django_integration.py



def save_robust_extraction_to_db(docx_file, paper_name, qualification, user, use_gemini=True, use_gemma=True):
    """
    Use the robust extractor to process DOCX and save to Django database
    
    Args:
        docx_file: Django uploaded file object
        paper_name: Name for the paper
        qualification: Qualification model instance
        user: User who uploaded
        use_gemini: Whether to use Gemini API
        use_gemma: Whether to use local Gemma
    
    Returns:
        Paper instance if successful, None if failed
    """
    print("🚀 Starting robust extraction to database...")
    
    try:
        # Save uploaded file temporarily
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            for chunk in docx_file.chunks():
                tmp.write(chunk)
            temp_path = tmp.name

        # Extract using robust extractor
        print("📄 Extracting document structure...")
        manifest = extract_docx(
            temp_path, 
            out_dir=None,  # Let it auto-generate
            use_gemini=use_gemini, 
            use_gemma=use_gemma,
            media_dir=os.path.join(extract_dir, "media")
        )

        copy_images_to_media_folder(media_dir)
        #Clean up temp file
        os.unlink(temp_path)
        
        if not manifest:
            print("❌ Extraction failed")
            return None
            
        print(f"✅ Extracted {manifest['counts']['questions']} questions")
        
        # Save to database
        with transaction.atomic():
            # Create Paper
            paper = Paper.objects.create(
                name=paper_name,
                qualification=qualification,
                total_marks=sum(
                    node.get('marks', 0) or 0 
                    for node in manifest['nodes'] 
                    if node.get('type') == 'question'
                ),
                structure_json=manifest  # Store full manifest as JSON
            )
            
            # Create ExamNodes
            node_map = {}  # Track for parent-child relationships
            
            for order, node_data in enumerate(manifest['nodes']):
                # Create ExamNode
                exam_node = ExamNode.objects.create(
                    id=uuid.uuid4().hex,
                    paper=paper,
                    node_type=map_robust_type_to_db(node_data.get('type', 'paragraph')),
                    number=node_data.get('number', ''),
                    marks=str(node_data.get('marks', '') or ''),
                    text=extract_node_text(node_data),
                    content= normalize_content_and_copy_media(node_data.get('content', [])),
                    manifest_output_dir=manifest.get('output_dir', ''),
                    paper_id=paper.id,
                    order_index=order
                )
                
                # Store in map for parent relationships
                if exam_node.number:
                    node_map[exam_node.number] = exam_node
                    
                print(f"📝 Saved node: {exam_node.number or exam_node.node_type}")
            
            # Set up parent-child relationships
            for number, node in node_map.items():
                if '.' in number:
                    parent_number = '.'.join(number.split('.')[:-1])
                    if parent_number in node_map:
                        node.parent = node_map[parent_number]
                        node.save()
                        print(f"🔗 Linked {number} → {parent_number}")
            
            print(f"✅ Saved paper '{paper.name}' with {len(manifest['nodes'])} nodes")
            return paper
            
    except Exception as e:
        print(f"❌ Database save failed: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return None

def map_robust_type_to_db(robust_type):
    """Map robust extractor types to ExamNode types"""
    mapping = {
        'question': 'question',
        'table': 'table', 
        'figure': 'image',
        'instruction': 'instruction',
        'case_study': 'case_study',
        'paragraph': 'instruction',  # Default paragraphs to instructions
        'rubric': 'instruction',
        'heading': 'instruction'
    }
    return mapping.get(robust_type, 'instruction')

def extract_node_text(node_data):
    """Extract readable text from node content"""
    texts = []
    
    for content_item in node_data.get('content', []):
        if content_item.get('type') in ['question_text', 'paragraph']:
            text = content_item.get('text', '')
            if text:
                texts.append(text)
    
    return ' '.join(texts) if texts else ''

# -----------------------------
# CLI entry point
# -----------------------------

if __name__ == "__main__":
    # Create argument parser
    ap = argparse.ArgumentParser(description="Robust exam paper extraction tool")
    ap.add_argument("docx", nargs="?", help="Input .docx file")
    ap.add_argument("-o", "--out", help="Output directory")
    ap.add_argument("--storage", default="extracted_papers", help="Storage directory")
    ap.add_argument("--gemini", action="store_true", help="Use Gemini API")
    ap.add_argument("--gemma", action="store_true", help="Use local Gemma")
    ap.add_argument("--save", action="store_true", help="Save to storage")
    ap.add_argument("--randomize", action="store_true", help="Randomize after extraction")
    ap.add_argument("--randomize-only", action="store_true", help="Only randomize latest")
    ap.add_argument("--export-docx", action="store_true", help="Export as DOCX")
    ap.add_argument("--paper-name", help="Paper name metadata")
    ap.add_argument("--module", help="Module name metadata")
    ap.add_argument("--tags", nargs="*", help="Paper tags")
    ap.add_argument("--same-top", action="store_true", help="Only swap within same top level")
    ap.add_argument("--marks-tolerance", type=int, help="Marks difference tolerance")
    ap.add_argument("--required-tags", nargs="*", help="Required donor tags")

    launch_gui()
