import io
import re
import traceback 
from PyPDF2 import PdfReader
import docx
from docx import Document
from core.extractor_images import convert_emf_to_png
from core.gemmaAI_classification import classify_with_local_gemma
import json
import google.generativeai as genai
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from zipfile import BadZipFile
import io
import uuid
from django.utils.timezone import now
from django.db import transaction
from core.models import ExamNode, Paper, RegexPattern
#To fix local issue

import requests
import sys

# Ensure stdout/stderr can handle UTF-8 log output on Windows consoles.
for stream_name in ("stdout", "stderr"):
    stream = getattr(sys, stream_name, None)
    if stream and hasattr(stream, "reconfigure"):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass

W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def ensure_ids(blocks):
    """Ensure all blocks have IDs by generating UUIDs"""
    for block in blocks:
        if 'id' not in block:
            block['id'] = uuid.uuid4().hex
        
        # Recursively ensure IDs for child blocks
        if 'children' in block and block['children']:
            ensure_ids(block['children'])
            
    return blocks

def extract_text(file_obj, content_type):
    """
    f: InMemoryUploadedFile or file-like
    content_type: e.g. 'application/pdf' or docx mime
    """
    file_obj.seek(0)
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_obj)
    elif content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return extract_text_from_docx(file_obj)
    else:
        # fallback: try decoding
        return file_obj.read().decode("utf-8", errors="ignore")
    
def extract_text_from_pdf(f):
    reader = PdfReader(f)
    texts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            texts.append(t)
    return "\n\n".join(texts)

def extract_text_from_docx(f):
    doc = docx.Document(io.BytesIO(f.read()))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())



import re
import time
from django.db.models import Q
from core.extract_AI import extract_with_ai
from core.gemmaAI_classification import classify_with_local_gemma


MAX_ATTEMPTS = 5
MIN_MATCH_THRESHOLD = 0.5  # 50% match rate for regex acceptance



def extract_full_docx_structure(docx_file):
    """
    Ultimate adaptive DOCX extractor with:
    - Full component extraction
    - Adaptive regex pool
    - Gemini primary, Gemma 3 fallback
    - Multi-pass validation
    - Full AI parse fallback
    - AI backfill for missing question numbers/marks
    """
    print("\n🔍 Starting ultimate adaptive document extraction...")

    stats = {'questions': 0, 'tables': 0, 'content_blocks': 0, 'errors': 0}

    try:
        # 1. Load DOCX
        doc = validate_and_load_document(docx_file)
        if not doc:
            print("❌ Could not load DOCX file")
            return []
        
        print(f"📄 Found {len(doc.paragraphs)} paragraphs")
        print(f"📊 Found {len(doc.tables)} tables")

        # 2. Extract raw paragraphs
        raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not raw_paragraphs:
            print("⚠️ No text found in DOCX")
            return []

        format_signature = detect_format_signature(raw_paragraphs)
        print(f"📄 Detected format signature: {format_signature}")

        # 3. Try regex from pool
        structured = []
        best_pattern = RegexPattern.objects.order_by("-match_score").first()

        if best_pattern:
            print(f"📌 Using pool regex: {best_pattern.pattern}")
            structured = apply_regex_to_paragraphs(best_pattern.pattern, raw_paragraphs)
            match_rate = len(structured) / len(raw_paragraphs)
            print(f"📊 Pool regex match rate: {match_rate:.0%}")
            
            if match_rate < MIN_MATCH_THRESHOLD:
                print("⚠️ Pool regex weak — regenerating")
                structured = []

        # 4. Generate regex via AI if needed
        if not structured:
            pattern = generate_regex_with_ai(raw_paragraphs)
            structured = apply_regex_to_paragraphs(pattern, raw_paragraphs)
            match_rate = len(structured) / len(raw_paragraphs)
            print(f"📊 AI regex match rate: {match_rate:.0%}")

            if match_rate >= MIN_MATCH_THRESHOLD:
                print("💾 Saving new regex to pool")
                RegexPattern.objects.create(
                    pattern=pattern,
                    description="AI-generated",
                    match_score=match_rate,
                    format_signature=format_signature,
                    example_usage="\n".join(raw_paragraphs[:5])
                )
            else:
                print("⚠️ AI regex weak — going to classification")
                structured = []

        # 5. Classification fallback
        if not structured:
            structured = ai_classification_loop(raw_paragraphs, stats)
        else:
            structured = ai_classification_loop_from_blocks(structured, stats)

        # 6. Full AI parse as final fallback (give it the raw paragraphs)
        if not structured:
            print("🚨 All earlier methods failed — sending raw paragraphs to AI parse")
        try:
            # pack the raw paragraphs into JSON so AI sees everything
            payload = {"paragraphs": raw_paragraphs}
            ai_prompt = f"""
    You’re given a list of paragraphs extracted from an exam paper.
    Extract every question (with number, text, marks), case-study, table, figure, etc.
    Return a JSON list of:
    - id (UUID)
    - type
    - number (if any)
    - text
    - marks
    - parent_id (if sub-question)
    Here is the payload:
    {json.dumps(payload, indent=2)}
    """
            ai_output = genai.GenerativeModel("gemini-1.5-flash") \
                        .generate_content(ai_prompt) \
                        .text
            structured = json.loads(ai_output)
            print(f"✅ Full AI parse returned {len(structured)} blocks")
        except Exception as e:
            print(f"❌ Full AI parse failed: {e}")
            # now at least hand back the raw paragraphs so downstream can inspect
            structured = [{"type":"paragraph","text":p} for p in raw_paragraphs]

        # 7. AI backfill
        structured = ai_backfill_missing_numbers_marks(structured)

        # 8. Normalize
        structured = normalize_recursive(structured)
        print("✅ Structure normalized and ready")
        print(f"✅ Extracted {len(structured)} blocks")
        return structured

    except Exception as e:
        print(f"❌ Fatal extraction error: {str(e)}")
        return []

def auto_classify_blocks_with_gemini(blocks):
    """
    Temporary wrapper so extraction doesn't fail.
    In future, this will use real Gemini API classification.
    For now, it just uses Gemma classification for testing.
    """
    print("⚠️ [TEMP] auto_classify_blocks_with_gemini not implemented, using Gemma instead")
    return auto_classify_blocks_with_gemma(blocks)


    
def auto_classify_blocks_with_gemma(blocks):
    """
    Wrapper to classify blocks using the local Gemma model.
    Matches output structure from Gemini classification.
    Injects the `type` field into each block and logs results.
    """
    print(f"🌀 [Gemma] Starting classification for {len(blocks)} blocks...")

    # Capture original types before Gemma modifies them
    original_types = [b.get('type') for b in blocks]

    # Defensive pre-types placeholder
    pre_types = [None] * len(blocks)

    try:
        types = classify_with_local_gemma(blocks, pre_types)
        if not isinstance(types, list) or len(types) != len(blocks):
            raise ValueError(f"Invalid Gemma output: expected list of length {len(blocks)}, got {types}")

        # Inject types into blocks
        for i, t in enumerate(types):
            if t == 'question_header':
                t = 'question'
            blocks[i]['type'] = t

        # Log before/after classification
        print("\n📊 [Gemma] Type Changes:")
        for i, (before, after) in enumerate(zip(original_types, types), start=1):
            normalized_after = 'question' if after == 'question_header' else after
            print(f"  Block {i}: {before or 'None'}  →  {normalized_after}")

        print(f"\n✅ [Gemma] Final classified types: {types}")
        return blocks

    except Exception as e:
        print(f"❌ [Gemma] Classification failed: {e}")
        # Fallback to marking all as paragraph
        for b in blocks:
            b['type'] = 'paragraph'
        print("⚠️ [Gemma] Defaulted all block types to 'paragraph'")
        return blocks

def detect_format_signature(paragraphs):
    """Detect numbering style signature."""
    if any(re.match(r"^\d+\.\d+", p) for p in paragraphs):
        return "decimal_nested"
    elif any(re.match(r"^\d+\)", p) for p in paragraphs):
        return "paren_number"
    elif any(re.match(r"^Q\d+", p, re.IGNORECASE) for p in paragraphs):
        return "q_prefix"
    return "generic"

def apply_regex_to_paragraphs(pattern, paragraphs):
    """Apply regex to paragraphs and build blocks."""
    compiled = re.compile(pattern, re.IGNORECASE)
    blocks = []
    for p in paragraphs:
        m = compiled.match(p)
        if m:
            blocks.append({
                "type": "question",
                "number": m.groupdict().get("number", ""),
                "text": m.groupdict().get("text", ""),
                "marks": m.groupdict().get("marks", "0"),
                "content": {},
                "children": []
            })
    return blocks

def generate_regex_with_ai(paragraphs):
    """Generate regex pattern using Gemini → Gemma fallback."""
    sample_text = "\n".join(paragraphs[:10])
    try:
        prompt = f"""
        Analyze the following exam question lines and create a Python regex 
        with groups: number, text, marks (if any).
        Return only the regex.
        Sample:
        {sample_text}
        """
        return auto_generate_regex_with_gemini(prompt).strip()
    except Exception:
        print("⚠️ Gemini failed — trying Gemma 3")
        return auto_generate_regex_with_gemma(prompt).strip()
    
def ai_classification_loop(paragraphs, stats):
    structured_raw = [
        {"type": "paragraph", "text": p, "content": {}, "children": []}
        for p in paragraphs
    ]
    for attempt in range(1, MAX_ATTEMPTS + 1):
        print(f"🔄 Classification pass {attempt}/{MAX_ATTEMPTS}...")
        try:
            candidate = auto_classify_blocks_with_gemini(structured_raw)
        except Exception:
            print("⚠️ Gemini failed — trying Gemma 3")
            try:
                candidate = auto_classify_blocks_with_gemma(structured_raw)
            except Exception:
                candidate = structured_raw

        # ✅ Normalize immediately — ensure 'number', 'marks', etc. exist
        candidate = normalize_recursive(candidate)

        if verify_structure(candidate, stats):
            return candidate

        time.sleep(1)

    # ✅ Final fallback: normalized raw
    return normalize_recursive(structured_raw)


def ai_classification_loop_from_blocks(blocks, stats):
    for attempt in range(1, MAX_ATTEMPTS + 1):
        print(f"🔄 Block classification pass {attempt}/{MAX_ATTEMPTS}...")
        try:
            candidate = auto_classify_blocks_with_gemini(blocks)
        except Exception:
            print("⚠️ Gemini failed — trying Gemma 3")
            try:
                candidate = auto_classify_blocks_with_gemma(blocks)
            except Exception:
                candidate = blocks

        # ✅ Normalize immediately — ensure 'number', 'marks', etc. exist
        candidate = normalize_recursive(candidate)

        if verify_structure(candidate, stats):
            return candidate

        time.sleep(1)

    # ✅ Final fallback: normalized original
    return normalize_recursive(blocks)


def ai_backfill_missing_numbers_marks(blocks):
    """Use AI to fill in missing question numbers or marks."""
    try:
        prompt = f"""
        Some extracted exam questions are missing 'number' or 'marks'.
        Please fill them logically, based on context and sequence.

        Input JSON:
        {blocks}
        """
        return auto_backfill_with_gemini(prompt)
    except Exception:
        print("⚠️ Gemini backfill failed — trying Gemma 3")
        return auto_backfill_with_gemma(prompt)

def normalize_block(block, index=0):
    block.setdefault('content', {})
    block.setdefault('children', [])
    block.setdefault('marks', '0')
    block.setdefault('number', '')
    block.setdefault('text', '')
    block.setdefault('order_index', index)
    return block

def normalize_recursive(blocks):
    for idx, b in enumerate(blocks):
        normalize_block(b, index=idx)
        if b.get("children"):
            normalize_recursive(b["children"])
    return blocks




import json
import re
import requests
import google.generativeai as genai

def auto_backfill_with_gemini(blocks):
    """Use Gemini to fill missing question numbers and marks."""
    prompt = f"""
Some extracted exam blocks are missing 'number' or 'marks'.  
Given the full JSON payload below, fill in any missing fields logically.
Return **only** the updated JSON list of blocks.

INPUT:
{json.dumps(blocks, indent=2)}
"""
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        raw = response.text or ""
        cleaned = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
        return json.loads(cleaned)
    except Exception as e:
        print(f"❌ auto_backfill_with_gemini failed: {e}")
        return blocks  # fallback to original list

def auto_backfill_with_gemma(blocks):
    """Use local Gemma to fill missing question numbers and marks."""
    prompt = f"""
Some extracted exam blocks are missing 'number' or 'marks'.  
Given the full JSON payload below, fill in any missing fields logically.
Return **only** the updated JSON list of blocks.

INPUT:
{json.dumps(blocks, indent=2)}
"""
    try:
        res = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": "gemma3:latest", "prompt": prompt, "stream": False},
            timeout=60
        )
        raw = res.json().get("response", "") or ""
        cleaned = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
        return json.loads(cleaned)
    except Exception as e:
        print(f"❌ auto_backfill_with_gemma failed: {e}")
        return blocks  # fallback to original list

def auto_generate_regex_with_gemini(prompt: str) -> str:
    """Ask Gemini to generate a Python regex string."""
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        raw = (response.text or "").strip()

        # Remove code fences if present
        cleaned = re.sub(r"^```(?:python|regex)?|```$", "", raw, flags=re.MULTILINE).strip()
        return cleaned
    except Exception as e:
        print(f"❌ auto_generate_regex_with_gemini failed: {e}")
        return r"^\s*(?P<number>\d+(?:\.\d+)*)(?:\s*[-.)]\s*)(?P<text>.*?)(?:\s*\(\s*(?P<marks>\d+)\s*marks?\s*\))?\s*$"

def auto_generate_regex_with_gemma(prompt: str) -> str:
    """Ask local Gemma to generate a Python regex string."""
    try:
        res = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": "gemma3:latest", "prompt": prompt, "stream": False},
            timeout=60
        )
        raw = res.json().get("response", "").strip()
        cleaned = re.sub(r"^```(?:python|regex)?|```$", "", raw, flags=re.MULTILINE).strip()
        return cleaned
    except Exception as e:
        print(f"❌ auto_generate_regex_with_gemma failed: {e}")
        return r"^\s*(?P<number>\d+(?:\.\d+)*)(?:\s*[-.)]\s*)(?P<text>.*?)(?:\s*\(\s*(?P<marks>\d+)\s*marks?\s*\))?\s*$"


def validate_and_load_document(docx_file):
    """Validate and load the document"""
    try:
        docx_file.seek(0)
        return Document(io.BytesIO(docx_file.read()))
    except BadZipFile:
        print("❌ Invalid or corrupted DOCX file")
        return None
    except Exception as e:
        print(f"❌ Document load failed: {str(e)}")
        return None

def build_question_pattern():
    """Build enhanced question pattern"""
    return re.compile(r'''
        ^\s*
        (?:Question(?:\s+Header)?\s*[:\-\u2013]?\s*)?
        (?P<number>\d+(?:\.\d+)*)
        (?:\s*[-\.)]\s*)?
        (?P<text>.*?)
        (?:\s*\(\s*(?P<marks>\d+)\s*marks?\s*\))?
        \s*$
    ''', re.IGNORECASE | re.VERBOSE | re.DOTALL)

def validate_block(block):
    """Validate block has required fields"""
    required = {
        'question': ['number', 'marks', 'text'],
        'case_study': ['text'],
        'instruction': ['text'],
        'table': ['content'],
        'image': ['binary_content', 'content_type']
    }
    
    block_type = block.get('type', '')
    if block_type in required:
        missing = [f for f in required[block_type] if not block.get(f)]
        if missing:
            print(f"⚠️ Block missing required fields: {missing}")
            return False
    return True

def verify_structure(blocks, stats):
    """Verify overall structure integrity"""
    if not blocks:
        print("❌ No blocks found for verification")
        return False
        
    # Verify question numbering
    numbers = [b['number'] for b in blocks if b.get('type') == 'question']
    if not verify_question_sequence(numbers):
        print("⚠️ Question numbering appears out of sequence — continuing anyway")
        
    # Verify marks total
    total_marks = sum(int(b.get('marks', 0)) for b in blocks if b.get('type') == 'question')
    if total_marks == 0:
        print("⚠️ No marks found in paper")
        return False
        
    return True

def print_extraction_stats(stats):
    """Print detailed extraction statistics"""
    print("\n📊 Extraction Statistics:")
    print(f"Questions processed: {stats['questions']}")
    print(f"Tables processed: {stats['tables']}")
    print(f"Content blocks: {stats['content_blocks']}")
    print(f"Errors encountered: {stats['errors']}")

# Fixed serialize_node function
def serialize_node(obj):
    print(f"\n🔄 Serializing node...")
    
    if isinstance(obj, dict):
        node_type = obj.get('type', '')
        node_id = obj.get('id', '')
        print(f"📦 Dict node: {node_type} (ID: {node_id})")
        
        # Get children if they exist, otherwise empty list
        children = obj.get('children', [])
        
        result = {
            "id":       obj.get("id"),
            "type":     obj.get("type", ""),
            "number":   obj.get("number", ""),
            "marks":    obj.get("marks", ""),
            "text":     obj.get("text", ""),
            "parent_id": str(obj.get("parent_id")) if obj.get("parent_id") else None,
            "content":  obj.get("content", []),
            "children": [serialize_node(c) for c in children],
            **({"data_uri": obj.get("data_uri", "")} if obj.get("type") == "figure" else {})
        }
        print(f"✅ Serialized dict node: {node_type}")
        return result

    elif isinstance(obj, ExamNode):
        print(f"📦 ExamNode: {obj.node_type} (ID: {obj.id})")
        # ... existing ExamNode serialization ...
        # Make sure to handle children properly here too
        children = getattr(obj, 'children', [])
        return {
            "id": str(obj.id),
            "type": obj.node_type,
            "number": getattr(obj, 'number', ''),
            "marks": getattr(obj, 'marks', ''),
            "text": getattr(obj, 'text', ''),
            "parent_id": str(obj.parent_id) if obj.parent_id else None,
            "content": getattr(obj, 'content', []),
            "children": [serialize_node(c) for c in children],
            **({"data_uri": getattr(obj, 'data_uri', '')} if obj.node_type == "figure" else {})
        }

    else:
        print(f"❌ Invalid node type: {type(obj)}")
        raise TypeError("Unsupported object type")

#--------Start of saving Serialized nodes to DB----------------------------------------------->
from core.models import ExamNode,RegexPattern, Paper  # Remove old model imports
import uuid

def save_nodes_to_db(nodes, paper):
    """Save extracted blocks to database with proper relationships"""
    print(f"\n💾 Starting database save for paper {paper.id}")
    
    try:
        with transaction.atomic():
            # Clear existing nodes
            print("🗑️ Clearing existing nodes...")
            ExamNode.objects.filter(paper=paper).delete()
            
            # Track nodes by number for parent-child relationships
            node_map = {}
            block_count = 0
            
            def save_block(block, parent=None, order_index=0):
                """Helper function to recursively save blocks"""
                nonlocal block_count
                
                # Create node - removed data_uri
                node = ExamNode.objects.create(
                    id=uuid.uuid4().hex,
                    paper=paper,
                    parent=parent,
                    node_type=block.get('type', 'text'),
                    number=block.get('number', ''),
                    marks=block.get('marks', '0'),
                    text=block.get('text', ''),
                    content=block.get('content', {}),
                    order_index=order_index
                )
                block_count += 1
                
                # Store in map if it has a number
                if node.number:
                    node_map[node.number] = node
                    print(f"📝 Saved node {node.number}")
                
                # Process children recursively
                for idx, child in enumerate(block.get('children', [])):
                    save_block(child, parent=node, order_index=idx)
                
                return node
            
            # Save all top-level blocks
            for idx, block in enumerate(nodes):
                save_block(block, parent=None, order_index=idx)
            
            print(f"\n✅ Successfully saved {block_count} nodes")
            return True
            
    except Exception as e:
        print(f"❌ Database save failed: {str(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        return False

def auto_classify_blocks_with_gemini(blocks):
    print("🔥 [DEBUG] Using UPDATED Gemini classification function")

    # 1. Defensive type check
    if not isinstance(blocks, list) or not all(isinstance(b, dict) for b in blocks):
        raise TypeError("❌ Expected list of dicts, but got invalid structure.")

    # 2. Pre‑label obvious types
    pre_types = []
    for b in blocks:
        text = (b.get('text') or '').strip()
        if re.match(r'^\d+(?:\.\d+)+', text):
            pre_types.append('question_header')
        elif text.lower().startswith('case study'):
            pre_types.append('case_study')
        elif b.get('data_uri'):
            pre_types.append('figure')
        elif b.get('rows') is not None:
            pre_types.append('table')
        elif re.fullmatch(r'-{5,}', text):
            pre_types.append('instruction')
        else:
            pre_types.append(None)
    print("🧠 [LOG] Pre‑assigned types:", pre_types)

    # 3. System prompt
    system_prompt = """
You are a classification assistant. You will be given a list of blocks extracted from a question paper.

Some blocks have already been pre-classified as one of: question_header, case_study, table, or figure — leave those untouched.

Your job is to classify the remaining untyped blocks using ONLY one of the following labels:
- "paragraph": General text that doesn’t give specific instructions, often background or descriptive.
- "instruction": Any directive to the student (e.g. "Answer all questions", "Use only the booklet", "Show all working").
- "rubric": General rules or formatting notes that apply to the exam, typically near the top of the paper.
- "diagram": Describes a visual element without being a full figure block.

 DO NOT classify any non-question text as "question_header".

---
*Examples of proper classification:*
- "1.1 Define the term ‘quality assurance’. (5 Marks)" → question_header
- "Use only the supplied EISA booklets." → instruction
- "All questions are compulsory." → instruction
- "The purpose of quality assurance is to ensure..." → paragraph
- "Diagrams must be labelled clearly." → rubric

---
Your output should ONLY be this JSON format:
The "types" list must have the same number of elements as the input blocks.

{
  "types": ["paragraph", "instruction", "rubric", ...]
}
""".strip()

    try :
        # 4. Send to Gemini
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content([
            system_prompt,
            json.dumps({'blocks': blocks, 'partial_types': pre_types})
        ])

        raw = (response.text or '').strip()
        cleaned = re.sub(r'^```(?:json)?|```$', '', raw, flags=re.MULTILINE).strip()
        ai_types = json.loads(cleaned).get('types', [])
        if not isinstance(ai_types, list):
            raise ValueError("Invalid 'types' from Gemini")

        # 5. Merge and inject into blocks
        final = [pre or ai for pre, ai in zip(pre_types, ai_types)]
        for i, t in enumerate(final):
            if t == 'question_header':
                t = 'question'
            blocks[i]['type'] = t

        print("✅ [LOG] Final merged types injected:", final)
        return blocks

    except Exception as e:
        print("❌ [ERROR] Gemini classification failed:", e)
        print("🔁 [FALLBACK] Trying classification with local Gemma...")

        # 6. Fallback via local Gemma
        ai_types = classify_with_local_gemma(blocks, pre_types)
        final = [pre or ai for pre, ai in zip(pre_types, ai_types)]
        for i, t in enumerate(final):
            if t == 'question_header':
                t = 'question'
            blocks[i]['type'] = t

        print("✅ [LOG] Gemma fallback injected types:", final)
        return blocks
#----------------------------------------------end----------------------------------------------------------------
#responsible for AI filtering not classification only.
import os
import json
import re
import logging


# Configure Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    print("Gemini API configured")
else:
    print("Warning: GEMINI_API_KEY not found in environment")

# --- Gemini Classification Logic ---

def is_structural_noise(text):
    text = text.strip().lower()
    noise_patterns = [
        r'^students are only allowed',
        r'^question paper',
        r'^instructions',
        r'^section [a-z]',
        r'^answer all questions',
        r'^eisa rules',
        r'^quality controller',
        r'^page \d+',
        r'^external integrated summative assessment',
        r'^this question paper consists of',
    ]
    return any(re.match(pattern, text) for pattern in noise_patterns)

def classify_block_type(text_blocks):
    """
    Accepts a list of text strings and returns a list of predicted block types:
    e.g., ['question', 'rubric', 'case_study', 'table', ...]
    """
    if not GEMINI_API_KEY:
        return ["other"] * len(text_blocks)

    prompt = f"""
You are a document examiner for exam papers.

Classify each of the following blocks into one of the following types:

- 'question': a numbered question (e.g. '1.1 Define...')
- 'case_study': if it contains context or scenario
- 'instruction': if it tells the learner what to do
- 'rubric': if it describes how marks are allocated
- 'table': if the block is a table or looks like one
- 'heading': if it is a heading like 'SECTION A'
- 'noise': if it is unrelated or structural like 'Page 1', or 'EISA Rules'
- 'other': if unsure

Return a JSON list in the same order as input.

INPUT:
{json.dumps(text_blocks[:40], indent=2)}

Respond with only the JSON list.
    """.strip()

    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        raw = response.text.strip()

        logging.info("📨 Gemini raw response: \n%s", raw)

        parsed = json.loads(raw)
        return parsed

    except Exception as e:
        logging.error("❌ Gemini classification failed: %s", e)
        return ["other"] * len(text_blocks)


#Rebuild a tree structure  for ensuring the tree instead of flattened is stored in DB
def rebuild_tree(flat_nodes):
    """
    Given a flat list of blocks with `number` and `parent_id`,
    reconstruct a nested tree.
    """
    id_map = {n['id']: {**n, 'children': []} for n in flat_nodes}

    root_nodes = []
    for node in id_map.values():
        parent_id = node.get('parent_id')
        if parent_id and parent_id in id_map:
            id_map[parent_id]['children'].append(node)
        else:
            root_nodes.append(node)

    return root_nodes

#<----------------------------------------------------------------------------------------------------------------------------->


def rebuild_nested_structure(flat_nodes):
    """
    Rebuilds a nested structure: parents like '1.1' will include their children like '1.1.1', '1.1.2' in a `children` list.
    Assumes each node has: id, number, parent_id (optional), and content fields.
    """
    id_to_node = {}
    root_nodes = []

    # Step 1: Prepare nodes and mapping
    for node in flat_nodes:
        node_copy = {**node, "children": []}
        id_to_node[node["id"]] = node_copy

    # Step 2: Assign children to parents
    for node in flat_nodes:
        parent_id = node.get("parent_id")
        if parent_id and parent_id in id_to_node:
            id_to_node[parent_id]["children"].append(id_to_node[node["id"]])
        else:
            root_nodes.append(id_to_node[node["id"]])  # no parent → top-level

    # Step 3: Optional — sort top-level by number (1.1, 1.2, ...)
    root_nodes.sort(key=lambda n: n.get("number", ""))
    return root_nodes
#-------------------------------------------------------------------------------------------------->>>>>>>>>
import uuid
from core.models import ExamNode
import uuid


def populate_examnodes_from_structure_json(paper, structure_json):
    """Convert structure JSON to ExamNodes"""
    print(f"💾 Creating ExamNodes for paper {paper.id}")
    
    ExamNode.objects.filter(paper=paper).delete()
    
    for block in structure_json:
        # Removed data_uri from creation
        node = ExamNode.objects.create(
            paper=paper,
            node_type=block.get('type', ''),
            number=block.get('number', ''),
            text=block.get('text', ''),
            marks=block.get('marks', ''),
            content=block.get('content', [])
        )
    print(f"✅ Created nodes from structure")
    return True



def validate_and_reorder_structure(blocks):
    """Validates and reorders blocks using Gemma3 LLM"""
    print("\n🔄 Starting structure validation and reordering...")
    
    # Prompt for Gemma3
    prompt = """
    Analyze this exam paper structure and ensure:
    1. Questions are in correct numerical order (1.1 before 1.2, etc)
    2. Case studies appear before their related questions
    3. Instructions appear at appropriate positions
    4. Content blocks are properly nested under their questions
    5. Tables and figures are placed with their relevant questions

    Return the reordered structure maintaining all original content.
    """

    try:
        from core.gemmaAI_classification import process_with_gemma3
        reordered = process_with_gemma3(blocks, prompt)
        print(f"✅ Structure reordered ({len(reordered)} blocks)")
        return reordered
    except Exception as e:
        print(f"⚠️ Reordering failed: {str(e)}")
        return blocks

def enhanced_save_structure(paper, structure_json):
    """Enhanced saving with validation and proper relationships"""
    print("\n💾 Starting enhanced structure save...")

    # 1. Validate and reorder
    structure_json = validate_and_reorder_structure(structure_json)

    try:
        with transaction.atomic():
            # Clear existing
            ExamNode.objects.filter(paper=paper).delete()
            
            # Track nodes by number for parent-child relationships
            node_map = {}
            block_count = 0
            
            def save_block(block, parent=None, order_index=0, depth=0):
                """Helper function to recursively save blocks"""
                nonlocal block_count
                
                if depth > 50:  # Prevent infinite recursion
                    print(f"⚠️ Depth limit reached for block {block.get('number', '')}")
                    return None
                
                # Create node
                node = ExamNode.objects.create(
                    id=uuid.uuid4().hex,
                    paper=paper,
                    parent=parent,
                    node_type=block.get('type', 'text'),
                    number=block.get('number', ''),
                    marks=block.get('marks', '0'),
                    text=block.get('text', ''),
                    content=block.get('content', {}),
                    order_index=order_index,
                    is_active=True
                )
                block_count += 1
                
                # Store in map if it has a number
                if node.number:
                    node_map[node.number] = node
                    print(f"📝 Saved node {node.number}")
                
                # Process children recursively
                for idx, child in enumerate(block.get('children', [])):
                    save_block(child, parent=node, order_index=idx, depth=depth+1)
                
                return node
            
            # Save all top-level blocks - FIXED variable name
            for idx, block in enumerate(structure_json):
                save_block(block, parent=None, order_index=idx)
            
            print(f"\n✅ Successfully saved {block_count} nodes")
            return True
            
    except Exception as e:
        print(f"❌ Save failed: {str(e)}")
        return False

def save_nodes_to_db(nodes, paper):
    """Save extracted blocks to database with proper relationships"""
    print(f"\n💾 Starting database save for paper {paper.id}")
    
    try:
        with transaction.atomic():
            # Clear existing nodes
            print("🗑️ Clearing existing nodes...")
            ExamNode.objects.filter(paper=paper).delete()
            
            # Track nodes by number for parent-child relationships
            node_map = {}
            block_count = 0
            
            def save_block(block, parent=None, order_index=0):
                """Helper function to recursively save blocks"""
                nonlocal block_count
                
                # Create node - removed data_uri
                node = ExamNode.objects.create(
                    id=uuid.uuid4().hex,
                    paper=paper,
                    parent=parent,
                    node_type=block.get('type', 'text'),
                    number=block.get('number', ''),
                    marks=block.get('marks', '0'),
                    text=block.get('text', ''),
                    content=block.get('content', {}),
                    order_index=order_index
                )
                block_count += 1
                
                # Store in map if it has a number
                if node.number:
                    node_map[node.number] = node
                    print(f"📝 Saved node {node.number}")
                
                # Process children recursively
                for idx, child in enumerate(block.get('children', [])):
                    save_block(child, parent=node, order_index=idx)
                
                return node
            
            # Save all top-level blocks
            for idx, block in enumerate(nodes):
                save_block(block, parent=None, order_index=idx)
            
            print(f"\n✅ Successfully saved {block_count} nodes")
            return True
            
    except Exception as e:
        print(f"❌ Database save failed: {str(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        return False

def save_parent_child_tables_from_structure_json(paper, blocks):
    """Convert structure JSON into ExamNode hierarchy"""
    print(f"\n🔄 Converting structure to ExamNode hierarchy for paper {paper.id}")
    
    try:
        node_map = {}
        
        # First pass - create all nodes without data_uri
        for block in blocks:
            node = ExamNode.objects.create(
                paper=paper,
                node_type=block.get('type', ''),
                number=block.get('number', ''),
                text=block.get('text', ''),
                marks=block.get('marks', ''),
                content=block.get('content', []),
                payload={
                    'text': block.get('text', ''),
                    'content': block.get('content', []),
                }
            )
            node_map[block.get('number')] = node
            print(f"✅ Created node for {block.get('number')}")

        # Second pass - establish parent-child relationships
        for block in blocks:
            number = block.get('number', '')
            if '.' in number:
                parent_number = '.'.join(number.split('.')[:-1])
                if parent_number in node_map:
                    child_node = node_map[number]
                    child_node.parent = node_map[parent_number]
                    child_node.save()
                    print(f"🔗 Linked {number} to parent {parent_number}")

        print(f"✅ Successfully saved {len(node_map)} nodes with relationships")
        return True

    except Exception as e:
        print(f"❌ Error saving structure: {str(e)}")
        return False
    
def extract_all_tables(docx_file):
    docx_file.seek(0)
    match_tables = extract_match_column_tables(docx_file)

    docx_file.seek(0)
    mcq_tables = extract_multiple_choice_tables(docx_file)

    docx_file.seek(0)
    generic_tables = extract_generic_tables(docx_file)

    return {
        "match_tables": match_tables,
        "mcq_tables": mcq_tables,
        "generic_tables": generic_tables,
    }

def process_element(element, pattern, registry, stats):
    """Process individual document elements and return structured blocks"""
    
    if element.tag.endswith('}p'):
        # Handle paragraphs
        text = "".join(run.text for run in element.iter(f'{{{W_NS}}}r')).strip()
        if not text:
            return None
            
        # Check for question pattern
        match = pattern.match(text)
        if match:
            # Question block
            number = match.group('number')
            
            # Validate question number
            if number in registry['numbers']:
                print(f"⚠️ Duplicate question number: {number}")
                return None
                
            registry['numbers'].add(number)
            stats['questions'] += 1
            
            block = {
                'type': 'question',
                'number': number,
                'text': match.group('text').strip(),
                'marks': match.group('marks') or '0',
                'content': [],
                'order_index': len(registry['numbers'])
            }
            
            # Update total marks
            registry['marks_total'] += int(block['marks'])
            print(f"📝 Found question {number} ({block['marks']} marks)")
            return block
            
        else:
            # Non-question block
            stats['content_blocks'] += 1
            return {
                'type': 'text',  # Will be classified later
                'text': text,
                'order_index': stats['content_blocks']
            }
            
    elif element.tag.endswith('}tbl'):
        # Handle tables
        print("📊 Processing table...")
        table_data = []
        
        for row in element.xpath('.//w:tr', namespaces={'w': W_NS}):
            cells = [
                "".join(t.text for t in cell.xpath('.//w:t', namespaces={'w': W_NS}))
                for cell in row.xpath('.//w:tc', namespaces={'w': W_NS})
            ]
            if any(cells):  # Skip empty rows
                table_data.append(cells)
        
        if table_data:
            stats['tables'] += 1
            return {
                'type': 'table',
                'content': table_data,
                'order_index': stats['tables']
            }
    
    return None

def extract_match_column_tables(docx_file):
    """Extract matching column type tables"""
    try:
        doc = Document(io.BytesIO(docx_file.read()))
        match_tables = []
        
        for table in doc.tables:
            # Check if it's a matching table (2 columns)
            if len(table.columns) == 2:
                columns = [[], []]
                for row in table.rows:
                    if len(row.cells) == 2:
                        columns[0].append(row.cells[0].text.strip())
                        columns[1].append(row.cells[1].text.strip())
                match_tables.append(columns)
                
        return match_tables
    except Exception as e:
        print(f"❌ Match table extraction failed: {str(e)}")
        return []

def extract_multiple_choice_tables(docx_file):
    """Extract MCQ style tables"""
    try:
        doc = Document(io.BytesIO(docx_file.read()))
        mcq_tables = []
        
        for table in doc.tables:
            # Check for MCQ structure (question + options)
            if any('a)' in cell.text.lower() for row in table.rows for cell in row.cells):
                mcq_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    mcq_data.append(row_data)
                mcq_tables.append(mcq_data)
                
        return mcq_tables
    except Exception as e:
        print(f"❌ MCQ table extraction failed: {str(e)}")
        return []

def extract_generic_tables(docx_file):
    try:
        doc = Document(io.BytesIO(docx_file.read()))
        generic_tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_data.append(row_data)
            if table_data:
                generic_tables.append(table_data)
        return generic_tables
    except Exception as e:
        print(f"❌ Generic table extraction failed: {e}")
        return []

def verify_question_sequence(numbers):
    """Verify question numbers are in correct sequence"""
    if not numbers:
        return False
        
    try:
        # Filter out empty strings first
        valid_numbers = [n for n in numbers if n.strip()]
        if not valid_numbers:
            return False
            
        # Convert to numeric for comparison
        numeric_parts = []
        for num in valid_numbers:
            parts = [int(p) for p in num.split('.')]
            numeric_parts.append(parts)
            
        # Check sequence
        for i in range(len(numeric_parts)-1):
            curr = numeric_parts[i]
            next = numeric_parts[i+1]
            if curr >= next:
                print(f"❌ Question numbering out of sequence: {valid_numbers[i]} >= {valid_numbers[i+1]}")
                return False
                
        return True
    except ValueError as e:
        print(f"❌ Invalid question number format: {str(e)}")
        return False

def post_process_blocks(blocks):
    """Post-process and validate block structure"""
    processed = []
    current_section = None
    
    for block in blocks:
        # Clean text
        if 'text' in block:
            block['text'] = block['text'].strip()
            
        # Validate marks
        if block.get('type') == 'question':
            try:
                marks = int(block.get('marks', '0'))
                block['marks'] = str(marks)
            except ValueError:
                block['marks'] = '0'
                
        processed.append(block)
        
    return processed





# import os
# import shutil
# from django.conf import settings

# def copy_images_to_media_folder(extract_media_dir):
#     """
#     Copies all images from the extraction media folder to Django's MEDIA_ROOT.
#     Converts .emf images to .png using fallback if needed.
#     """
#     target_dir = settings.MEDIA_ROOT
#     os.makedirs(target_dir, exist_ok=True)

#     for fname in os.listdir(extract_media_dir):
#         src = os.path.join(extract_media_dir, fname)
#         ext = os.path.splitext(fname)[1].lower()
#         dst = os.path.join(target_dir, fname)

#         if ext == ".emf":
#             # Convert EMF to PNG
#             png_name = fname.replace(".emf", ".png")
#             png_dst = os.path.join(target_dir, png_name)
#             success = convert_emf_to_png(src, png_dst)
#             if success:
#                 print(f"[✓] Converted {fname} to {png_name}")
#             else:
#                 print(f"[✗] Failed to convert {fname} to PNG")
#         else:
#             print(f"[✓] Copying {fname} to media folder")
#             print(f"   From: {src}")
#             print(f"   To: {dst}")
#             shutil.copy2(src, dst)



# def load_or_init_json(path: str):
#     m = re.search(r'upload_([a-f0-9]+)_extract', path, re.I)
#     upload_id = m.group(1) if m else "unknown"
#     default = {"id": upload_id, "status": "uploaded", "nodes": []}

#     os.makedirs(os.path.dirname(path), exist_ok=True)
#     if not os.path.exists(path) or os.path.getsize(path) == 0:
#         with open(path, "w", encoding="utf-8") as f:
#             json.dump(default, f, indent=2)
#         return default
#     try:
#         with open(path, "r", encoding="utf-8") as f:
#             return json.load(f)
#     except Exception:
#         with open(path, "w", encoding="utf-8") as f:
#             json.dump(default, f, indent=2)
#         return default


# # --- call the loader ---
# json_path = r"C:\Users\bjmba\CHIETA_LMS_fresh\upload_7e63da229b98466a862a494140c89d47_extract\structure_json.json"
# data = load_or_init_json(json_path)

# # --- media folder check ---
# media_folder = r"C:\Users\bjmba\CHIETA_LMS_fresh\media"
# os.makedirs(media_folder, exist_ok=True)   # make sure it exists
# actual_images = set(os.listdir(media_folder))

# # --- fix image references in nodes ---
# for node in data.get('nodes', []):
#     for block in node.get('content', []):
#         if block.get('type') == 'figure':
#             block['images'] = [
#                 img for img in block.get('images', []) 
#                 if img in actual_images
#             ]

# # --- save JSON back ---
# with open(json_path, 'w', encoding="utf-8") as f:
#     json.dump(data, f, indent=2)


import os
import shutil
import json
import re
from django.conf import settings

def get_upload_directory(upload_id=None):
    """Returns a writable directory for upload extraction files"""
    if upload_id:
        # Use MEDIA_ROOT for upload-specific files
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'uploads', f'upload_{upload_id}_extract')
    else:
        # Fallback to project data directory
        upload_dir = os.path.join(settings.BASE_DIR, 'data', 'uploads')
    
    os.makedirs(upload_dir, exist_ok=True)
    return upload_dir

def copy_images_to_media_folder(extract_media_dir):
    """
    Copies all images from the extraction media folder to Django's MEDIA_ROOT.
    Converts .emf images to .png using fallback if needed.
    """
    target_dir = settings.MEDIA_ROOT
    os.makedirs(target_dir, exist_ok=True)

    for fname in os.listdir(extract_media_dir):
        src = os.path.join(extract_media_dir, fname)
        ext = os.path.splitext(fname)[1].lower()
        dst = os.path.join(target_dir, fname)

        if ext == ".emf":
            # Convert EMF to PNG
            png_name = fname.replace(".emf", ".png")
            png_dst = os.path.join(target_dir, png_name)
            success = convert_emf_to_png(src, png_dst)
            if success:
                print(f"[✓] Converted {fname} to {png_name}")
            else:
                print(f"[✗] Failed to convert {fname} to PNG")
        else:
            print(f"[✓] Copying {fname} to media folder")
            print(f"   From: {src}")
            print(f"   To: {dst}")
            shutil.copy2(src, dst)

def load_or_init_json(path: str):
    """Load JSON data or initialize with default if file doesn't exist"""
    # Extract upload ID from path
    m = re.search(r'upload_([a-f0-9]+)_extract', path, re.I)
    upload_id = m.group(1) if m else "unknown"
    default = {"id": upload_id, "status": "uploaded", "nodes": []}

    try:
        # Create parent directories if they don't exist
        directory = os.path.dirname(path)
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)
        
        if not os.path.exists(path) or os.path.getsize(path) == 0:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(default, f, indent=2)
            return default
        
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
            
    except (PermissionError, OSError):
        # Fallback to a writable location if original path fails
        fallback_dir = get_upload_directory(upload_id)
        fallback_path = os.path.join(fallback_dir, 'structure_json.json')
        
        print(f"Permission denied for {path}, using fallback: {fallback_path}")
        
        os.makedirs(os.path.dirname(fallback_path), exist_ok=True)
        if not os.path.exists(fallback_path) or os.path.getsize(fallback_path) == 0:
            with open(fallback_path, "w", encoding="utf-8") as f:
                json.dump(default, f, indent=2)
            return default
        
        with open(fallback_path, "r", encoding="utf-8") as f:
            return json.load(f)
            
    except Exception as e:
        print(f"Error loading JSON: {e}")
        return default

def process_upload_data(upload_id):
    """Main function to process upload data with proper path handling"""
    # Get proper directory paths
    upload_dir = get_upload_directory(upload_id)
    json_filename = 'structure_json.json'
    json_path = os.path.join(upload_dir, json_filename)
    
    # Load or initialize JSON data
    data = load_or_init_json(json_path)
    
    # Ensure media folder exists
    media_folder = settings.MEDIA_ROOT
    os.makedirs(media_folder, exist_ok=True)
    actual_images = set(os.listdir(media_folder))

    # Fix image references in nodes
    for node in data.get('nodes', []):
        for block in node.get('content', []):
            if block.get('type') == 'figure':
                block['images'] = [
                    img for img in block.get('images', []) 
                    if img in actual_images
                ]

    # Save JSON back to proper location
    try:
        with open(json_path, 'w', encoding="utf-8") as f:
            json.dump(data, f, indent=2)
        print(f"Data successfully saved to: {json_path}")
    except PermissionError:
        # Fallback to project data directory if upload directory is not writable
        fallback_dir = os.path.join(settings.BASE_DIR, 'data', 'uploads')
        fallback_path = os.path.join(fallback_dir, f'upload_{upload_id}_extract', json_filename)
        os.makedirs(os.path.dirname(fallback_path), exist_ok=True)
        
        with open(fallback_path, 'w', encoding="utf-8") as f:
            json.dump(data, f, indent=2)
        print(f"Data saved to fallback location: {fallback_path}")

    return data

